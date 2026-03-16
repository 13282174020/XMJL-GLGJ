# -*- coding: utf-8 -*-
"""测试直接复制模板文档并修改内容"""
from docx import Document
import os

template_path = r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx'
output_path = r'e:\Qwen\xmjl\outputs\test_copy_template.docx'

# 直接复制模板文档
doc = Document(template_path)

# 统计原始样式
orig_styles = {}
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    orig_styles[style_name] = orig_styles.get(style_name, 0) + 1

print(f'原始模板样式分布（前 10 个）:')
for style, count in sorted(orig_styles.items(), key=lambda x: -x[1])[:10]:
    print(f'  {style}: {count}')

# 保存为新文档
os.makedirs(os.path.dirname(output_path), exist_ok=True)
doc.save(output_path)

# 验证新文档
new_doc = Document(output_path)
new_styles = {}
for para in new_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    new_styles[style_name] = new_styles.get(style_name, 0) + 1

print(f'\n复制后文档样式分布（前 10 个）:')
for style, count in sorted(new_styles.items(), key=lambda x: -x[1])[:10]:
    print(f'  {style}: {count}')

print(f'\n文档已保存到：{output_path}')
print(f'原始段落数：{len(doc.paragraphs)}, 复制后段落数：{len(new_doc.paragraphs)}')
