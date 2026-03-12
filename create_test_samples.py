#!/usr/bin/env python
# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_standard_sample(output_path):
    doc = Document()
    title = doc.add_heading('标准格式测试文档', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('这是一个使用标准 Heading 样式的测试文档。')
    
    doc.add_heading('第一章 项目概述', level=1)
    doc.add_paragraph('这是第一章的内容。')
    doc.add_heading('1.1 项目背景', level=2)
    doc.add_paragraph('项目背景描述。')
    doc.add_heading('1.1.1 行业现状', level=3)
    doc.add_paragraph('行业现状分析。')
    doc.add_heading('1.1.1.1 国内市场', level=4)
    doc.add_paragraph('国内市场描述。')
    doc.add_heading('1.1.1.2 国际市场', level=4)
    doc.add_paragraph('国际市场描述。')
    doc.add_heading('1.1.2 技术发展趋势', level=3)
    doc.add_paragraph('技术发展趋势。')
    doc.add_heading('1.2 项目目标', level=2)
    doc.add_paragraph('项目目标。')
    doc.add_heading('1.2.1 短期目标', level=3)
    doc.add_paragraph('短期目标。')
    doc.add_heading('1.2.2 长期目标', level=3)
    doc.add_paragraph('长期目标。')
    
    doc.add_heading('第二章 技术方案', level=1)
    doc.add_paragraph('第二章内容。')
    doc.add_heading('2.1 技术架构', level=2)
    doc.add_paragraph('技术架构。')
    doc.add_heading('2.1.1 前端架构', level=3)
    doc.add_paragraph('前端架构。')
    doc.add_heading('2.1.2 后端架构', level=3)
    doc.add_paragraph('后端架构。')
    doc.add_heading('2.1.2.1 服务层', level=4)
    doc.add_paragraph('服务层。')
    doc.add_heading('2.1.2.2 数据层', level=4)
    doc.add_paragraph('数据层。')
    doc.add_heading('2.2 关键技术', level=2)
    doc.add_paragraph('关键技术。')
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'已创建标准格式文档：{output_path}')

def create_nonstandard_sample(output_path):
    doc = Document()
    title = doc.add_paragraph('非标准格式测试文档')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0] if title.runs else title.add_run()
    title_run.bold = True
    title_run.font.size = Pt(18)
    doc.add_paragraph('这是一个使用非标准格式的测试文档。')
    
    p1 = doc.add_paragraph('第一章 项目概述')
    p1_run = p1.runs[0] if p1.runs else p1.add_run()
    p1_run.bold = True
    p1_run.font.size = Pt(14)
    doc.add_paragraph('这是第一章的内容。')
    
    p2 = doc.add_paragraph('1.1 项目背景')
    p2_run = p2.runs[0] if p2.runs else p2.add_run()
    p2_run.bold = True
    p2_run.font.size = Pt(12)
    doc.add_paragraph('项目背景描述。')
    
    p3 = doc.add_paragraph('1.1.1 行业现状')
    p3_run = p3.runs[0] if p3.runs else p3.add_run()
    p3_run.bold = True
    p3_run.font.size = Pt(11)
    doc.add_paragraph('行业现状分析。')
    
    p4 = doc.add_paragraph('1.1.1.1 国内市场')
    p4_run = p4.runs[0] if p4.runs else p4.add_run()
    p4_run.bold = True
    doc.add_paragraph('国内市场描述。')
    
    p4_2 = doc.add_paragraph('1.1.1.2 国际市场')
    p4_2_run = p4_2.runs[0] if p4_2.runs else p4_2.add_run()
    p4_2_run.bold = True
    doc.add_paragraph('国际市场描述。')
    
    p3_2 = doc.add_paragraph('1.1.2 技术发展趋势')
    p3_2_run = p3_2.runs[0] if p3_2.runs else p3_2.add_run()
    p3_2_run.bold = True
    p3_2_run.font.size = Pt(11)
    doc.add_paragraph('技术发展趋势。')
    
    p2_2 = doc.add_paragraph('1.2 项目目标')
    p2_2_run = p2_2.runs[0] if p2_2.runs else p2_2.add_run()
    p2_2_run.bold = True
    p2_2_run.font.size = Pt(12)
    doc.add_paragraph('项目目标。')
    
    p3_3 = doc.add_paragraph('1.2.1 短期目标')
    p3_3_run = p3_3.runs[0] if p3_3.runs else p3_3.add_run()
    p3_3_run.bold = True
    p3_3_run.font.size = Pt(11)
    doc.add_paragraph('短期目标。')
    
    p3_4 = doc.add_paragraph('1.2.2 长期目标')
    p3_4_run = p3_4.runs[0] if p3_4.runs else p3_4.add_run()
    p3_4_run.bold = True
    p3_4_run.font.size = Pt(11)
    doc.add_paragraph('长期目标。')
    
    p1_2 = doc.add_paragraph('第二章 技术方案')
    p1_2_run = p1_2.runs[0] if p1_2.runs else p1_2.add_run()
    p1_2_run.bold = True
    p1_2_run.font.size = Pt(14)
    doc.add_paragraph('第二章内容。')
    
    p2_3 = doc.add_paragraph('2.1 技术架构')
    p2_3_run = p2_3.runs[0] if p2_3.runs else p2_3.add_run()
    p2_3_run.bold = True
    p2_3_run.font.size = Pt(12)
    doc.add_paragraph('技术架构。')
    
    p3_5 = doc.add_paragraph('2.1.1 前端架构')
    p3_5_run = p3_5.runs[0] if p3_5.runs else p3_5.add_run()
    p3_5_run.bold = True
    p3_5_run.font.size = Pt(11)
    doc.add_paragraph('前端架构。')
    
    p3_6 = doc.add_paragraph('2.1.2 后端架构')
    p3_6_run = p3_6.runs[0] if p3_6.runs else p3_6.add_run()
    p3_6_run.bold = True
    p3_6_run.font.size = Pt(11)
    doc.add_paragraph('后端架构。')
    
    p4_3 = doc.add_paragraph('2.1.2.1 服务层')
    p4_3_run = p4_3.runs[0] if p4_3.runs else p4_3.add_run()
    p4_3_run.bold = True
    doc.add_paragraph('服务层。')
    
    p4_4 = doc.add_paragraph('2.1.2.2 数据层')
    p4_4_run = p4_4.runs[0] if p4_4.runs else p4_4.add_run()
    p4_4_run.bold = True
    doc.add_paragraph('数据层。')
    
    p2_4 = doc.add_paragraph('2.2 关键技术')
    p2_4_run = p2_4.runs[0] if p2_4.runs else p2_4.add_run()
    p2_4_run.bold = True
    p2_4_run.font.size = Pt(12)
    doc.add_paragraph('关键技术。')
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f'已创建非标准格式文档：{output_path}')

def main():
    base_dir = os.path.dirname(os.path.abspath(__file__))
    fixtures_dir = os.path.join(base_dir, 'tests', 'fixtures')
    print('=' * 60)
    print('开始生成测试样本文档...')
    print('=' * 60)
    
    standard_path = os.path.join(fixtures_dir, 'sample_standard.docx')
    create_standard_sample(standard_path)
    
    nonstandard_path = os.path.join(fixtures_dir, 'sample_nonstandard.docx')
    create_nonstandard_sample(nonstandard_path)
    
    print('=' * 60)
    print('生成完成!')
    print('=' * 60)
    print()
    print('验证结果:')
    for path in [standard_path, nonstandard_path]:
        if os.path.exists(path):
            size = os.path.getsize(path)
            print(f'  [OK] {path} ({size} 字节)')
        else:
            print(f'  [FAIL] {path} (未找到)')

if __name__ == '__main__':
    main()
