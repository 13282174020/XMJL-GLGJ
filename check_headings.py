from docx import Document

doc_path = '新模板_1773310314036 (1).docx'
doc = Document(doc_path)

print('=== Heading 1 样式 ===\n')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if p.style.name == 'Heading 1' and text:
        print(f'{i:3d}: {text}')

print('\n=== Heading 2 样式 ===\n')
count = 0
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if p.style.name == 'Heading 2' and text:
        count += 1
        if count <= 30:
            print(f'{i:3d}: {text}')

print(f'\nHeading 2 总计：{count} 个')
