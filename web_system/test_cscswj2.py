# -*- coding: utf-8 -*-
"""测试 cscswj2.docx 的完整流程"""
import sys
import logging
sys.path.insert(0, '.')

logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')

from template_analyzer import DocumentAnalyzer
from template_preprocessor import TemplatePreprocessor

template_path = 'cscswj2.docx'
output_path = 'outputs/test_cscswj2_processed.docx'

print('=' * 80)
print('步骤 1: 分析模板')
print('=' * 80)

analyzer = DocumentAnalyzer()
report = analyzer.analyze(template_path)

print(f'\n样式使用：{report.styles_usage}')
print(f'潜在标题样式：{report.potential_heading_styles}')

print('\n映射规则:')
rules = report.generate_mapping_rules()
for rule in rules:
    print(f"  {rule['source_style']} + {rule.get('pattern', 'None')} -> {rule['target_style']}")

print('\n' + '=' * 80)
print('步骤 2: 预处理文档')
print('=' * 80)

preprocessor = TemplatePreprocessor()
result = preprocessor.preprocess_with_analysis(template_path, output_path)

print(f'\n结果：{result.message}')
print(f'统计：{result.stats}')

if result.success:
    print('\n' + '=' * 80)
    print('步骤 3: 验证生成的文档')
    print('=' * 80)
    
    from docx import Document
    doc = Document(output_path)
    
    print(f'生成的文档段落数：{len(doc.paragraphs)}')
    
    styles = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else 'Normal'
        styles[style] = styles.get(style, 0) + 1
    
    print(f'\n样式分布（前 10 个）:')
    for style, count in sorted(styles.items(), key=lambda x: -x[1])[:10]:
        print(f'  {style}: {count}')
    
    print(f'\n前 30 个非空段落:')
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else 'Normal'
        print(f'  [{style}] {text[:60]}')
        count += 1
        if count >= 30:
            break
