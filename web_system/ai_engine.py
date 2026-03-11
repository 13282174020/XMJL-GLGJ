# -*- coding: utf-8 -*-
"""
AI 引擎服务 - 基础 AI 调用功能
调用百炼 API 生成章节内容

优化：支持字段分类
- 信息型字段：从需求文档提取（简短）
- 描述型字段：AI 详细生成（200-400 字）

V2.0 更新：集成数据点管理和需求分析
- 数据一致性：注入已确立的数据点
- 需求覆盖度：注入章节应回应的需求点
"""

import requests
import json
import re
from typing import Dict, Optional, Callable, Any

# 导入优化服务
from services.data_point_manager import DataPointManager
from services.requirement_analyzer import RequirementAnalyzer
from services.content_optimizer import ContentOptimizer
from model_config import get_model_config, ModelConfig


# 全局实例（用于保持跨章节的数据一致性）
_data_point_manager: Optional[DataPointManager] = None
_requirement_analyzer: Optional[RequirementAnalyzer] = None
_content_optimizer: Optional[ContentOptimizer] = None


def get_data_point_manager() -> DataPointManager:
    """获取数据点管理器实例（单例）"""
    global _data_point_manager
    if _data_point_manager is None:
        _data_point_manager = DataPointManager()
    return _data_point_manager


def get_requirement_analyzer() -> RequirementAnalyzer:
    """获取需求分析器实例（单例）"""
    global _requirement_analyzer
    if _requirement_analyzer is None:
        _requirement_analyzer = RequirementAnalyzer()
    return _requirement_analyzer


def get_content_optimizer() -> ContentOptimizer:
    """获取内容优化器实例（单例）"""
    global _content_optimizer
    if _content_optimizer is None:
        _content_optimizer = ContentOptimizer()
    return _content_optimizer


def extract_project_type(requirement_text: str) -> str:
    """从需求文档中提取项目类型
    
    Args:
        requirement_text: 需求文档内容
        
    Returns:
        项目类型字符串，如"未来社区"、"智慧园区"等
    """
    if not requirement_text:
        return ""
    
    # 项目类型关键词映射
    project_types = {
        '未来社区': ['未来社区', '未来小区'],
        '智慧社区': ['智慧社区', '智慧小区', '智慧住区'],
        '智慧园区': ['智慧园区', '产业园区', '科技园区', '工业园'],
        '智慧校园': ['智慧校园', '数字校园', '智慧学校', '数字化校园'],
        '智慧医院': ['智慧医院', '数字医院', '智慧医疗', '医院信息化'],
        '智慧政务': ['智慧政务', '数字政府', '政务信息化', '电子政务'],
        '智慧乡村': ['智慧乡村', '数字乡村', '乡村治理'],
        '智慧景区': ['智慧景区', '智慧旅游', '景区管理'],
    }
    
    # 统计每个项目类型的匹配次数
    type_scores = {}
    for project_type, keywords in project_types.items():
        score = sum(1 for keyword in keywords if keyword in requirement_text)
        if score > 0:
            type_scores[project_type] = score
    
    # 返回匹配度最高的项目类型
    if type_scores:
        return max(type_scores, key=type_scores.get)
    
    return ""


def extract_business_scenes(requirement_text: str) -> list:
    """从需求文档中提取业务场景
    
    Args:
        requirement_text: 需求文档内容
        
    Returns:
        业务场景列表
    """
    if not requirement_text:
        return []
    
    scenes = []
    
    # 业务场景关键词分类
    scene_keywords = {
        '智慧安防': ['智慧安防', '安防监控', '高空抛物', '门禁系统', '车辆识别', '平安码'],
        '人员管理': ['人员管理', '流动人口', '人口分析', '访客管理', '实名制'],
        '邻里商业': ['邻里街', '商业配套', '15 分钟生活圈', '商铺管理'],
        '共享空间': ['共享空间', '共享书房', '活动室', '公共空间', '文化空间'],
        '未来健康': ['未来健康', '健康管理', '健康咨询', '医疗服务'],
        '未来低碳': ['未来低碳', '低碳社区', '节能减排', '绿色社区', '垃圾分类'],
        '物业管理': ['物业管理', '物业服务', '物业费', '水电上报'],
        '停车管理': ['停车管理', '停车位', '智能停车', '车辆管理'],
        '环境监测': ['环境监测', '空气质量', '噪音监测', '环境监测'],
        '能源管理': ['能源管理', '能耗监测', '智能电表', '节能管理'],
        '社区治理': ['社区治理', '网格化管理', '基层治理', '民主协商'],
        '志愿服务': ['志愿服务', '志愿者', '公益活动', '社区活动'],
        '智慧养老': ['智慧养老', '养老服务', '居家养老', '老年服务'],
        '智慧教育': ['智慧教育', '社区教育', '培训服务'],
        '智慧物流': ['智慧物流', '快递柜', '物流配送', '最后 100 米'],
    }
    
    for scene_name, keywords in scene_keywords.items():
        for keyword in keywords:
            if keyword in requirement_text:
                if scene_name not in scenes:
                    scenes.append(scene_name)
                break
    
    return scenes


def reset_optimization_services():
    """重置优化服务（用于新文档生成任务）"""
    global _data_point_manager, _requirement_analyzer, _content_optimizer
    _data_point_manager = DataPointManager()
    _requirement_analyzer = RequirementAnalyzer()
    _content_optimizer = ContentOptimizer()
    print('[INFO] 优化服务已重置')


# 字段分类配置
INFO_FIELDS = [
    '项目名称', '项目建设单位', '负责人', '联系方式',
    '建设工期', '总投资', '资金来源', '编制单位',
    '项目建设单位与职能', '项目实施机构', '项目实施机构与职责'
]

# 信息提取正则表达式
INFO_PATTERNS = {
    '项目名称': [r'项目名称\s*[:：]\s*(.+)', r'项目名称为\s*[:：]\s*(.+)'],
    '项目建设单位': [r'建设单位\s*[:：]\s*(.+)', r'业主单位\s*[:：]\s*(.+)', r'项目建设单位\s*[:：]\s*(.+)'],
    '负责人': [r'负责人\s*[:：]\s*(.+)', r'联系人\s*[:：]\s*(.+)'],
    '联系方式': [r'联系方式\s*[:：]\s*(.+)', r'联系电话\s*[:：]\s*(.+)', r'电话\s*[:：]\s*(.+)'],
    '建设工期': [r'建设工期\s*[:：]\s*(.+)', r'工期\s*[:：]\s*(.+)', r'([\d]+)[个]?[年月天]'],
    '总投资': [r'总投资\s*[:：]\s*(.+)', r'投资估算\s*[:：]\s*(.+)', r'([\d\.]+)\s*万?元'],
    '资金来源': [r'资金来源\s*[:：]\s*(.+)', r'资金筹措\s*[:：]\s*(.+)'],
    '编制单位': [r'编制单位\s*[:：]\s*(.+)', r'可研编制\s*[:：]\s*(.+)'],
}


def _extract_value_from_path(data: Dict, path: str) -> Any:
    """根据路径从嵌套字典中提取值

    Args:
        data: 字典数据
        path: 路径，如 "choices.0.message.content"

    Returns:
        提取的值
    """
    keys = path.split('.')
    current = data
    for key in keys:
        if isinstance(current, dict) and key in current:
            current = current[key]
        elif isinstance(current, list) and key.isdigit():
            idx = int(key)
            if idx < len(current):
                current = current[idx]
            else:
                return None
        else:
            return None
    return current


def _extract_ai_response(result: Dict, model_config: ModelConfig) -> str:
    """提取 AI 响应内容

    支持多种响应格式：
    1. Ollama 原生格式：message.content
    2. 标准 OpenAI 格式：choices.0.message.content
    3. 百炼格式：output.text
    4. reasoning_content（思考过程，不使用，仅用于调试）

    重要：不返回 reasoning_content（思考过程），只返回最终的生成内容

    Args:
        result: API 返回的原始数据
        model_config: 模型配置

    Returns:
        提取的响应内容（不包含思考过程）
    """
    print(f'[DEBUG] API 返回原始数据：{json.dumps(result, ensure_ascii=False)[:500]}...')
    
    # Ollama 原生格式：message.content
    if model_config.provider_id == 'ollama':
        if 'message' in result and isinstance(result['message'], dict):
            content = result['message'].get('content', '')
            if content:
                print(f'[DEBUG] 从 Ollama message.content 提取成功：{content[:50]}...')
                return content
    
    # 首先尝试从配置的 response_path 提取
    content = _extract_value_from_path(result, model_config.response_path)

    # 如果提取到了内容且不为空，直接返回
    if content:
        print(f'[DEBUG] 从 response_path 提取成功：{content[:50]}...')
        return content

    # 尝试 output 字段（百炼格式）
    if "output" in result and "text" in result["output"]:
        return result["output"]["text"]

    # 优先返回实际内容 content，而不是 reasoning_content
    # 对于 GLM-4.7 等启用思考模式的模型，content 是实际生成内容
    content_field = _extract_value_from_path(result, "choices.0.message.content")
    if content_field:
        print(f'[DEBUG] 从 choices.0.message.content 提取成功：{content_field[:50]}...')
        return content_field
    
    # Ollama 兼容：检查是否有 choices 但结构略有不同
    if "choices" in result and len(result["choices"]) > 0:
        choice = result["choices"][0]
        if "message" in choice:
            msg = choice["message"]
            if isinstance(msg, dict):
                # 尝试多个可能的字段名
                for key in ["content", "text", "response"]:
                    if key in msg and msg[key]:
                        print(f'[DEBUG] 从 choices[0].message.{key} 提取成功')
                        return msg[key]
    
    # 尝试直接返回 text 字段（某些 API 直接返回）
    if "text" in result:
        return result["text"]

    # 注意：不再返回 reasoning_content（思考过程）
    # 如果 content 为空，返回空字符串或错误提示
    reasoning = _extract_value_from_path(result, "choices.0.message.reasoning_content")
    if reasoning:
        print(f'[WARN] content 字段为空，reasoning_content（思考过程）将被忽略')
        # 不返回思考过程，继续向下处理

    # 尝试直接返回错误信息（不包含原始 JSON，避免泄露思考过程）
    print(f'[ERROR] 无法从响应中提取内容，完整响应：{json.dumps(result, ensure_ascii=False)[:1000]}')
    return f"[API 返回格式异常] 无法提取内容，请检查 API 返回格式是否正确"


def call_ai_api(
    prompt: str,
    model_config: ModelConfig
) -> str:
    """通用 AI API 调用函数

    根据模型配置调用不同的 AI API

    Args:
        prompt: 提示词
        model_config: 模型配置

    Returns:
        AI 生成的文本内容
    """
    # Ollama 本地模型不需要 API Key
    if not model_config.api_key and model_config.provider_id != 'ollama':
        return "[API 错误] API Key 未配置"

    # 如果 base_url 为空，尝试从厂商配置获取
    base_url = model_config.base_url
    if not base_url:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        provider = manager.get_provider(model_config.provider_id)
        if provider:
            base_url = provider.base_url
            print(f'[DEBUG] 模型 {model_config.id} 的 base_url 为空，使用厂商默认值：{base_url}')
        else:
            return f"[API 错误] base_url 未配置且无法从厂商 {model_config.provider_id} 获取默认值"
    
    # Ollama 本地模型使用原生 API 端点
    if model_config.provider_id == 'ollama' and '/api/chat' not in base_url:
        base_url = "http://localhost:11434/api/chat"
        print(f'[DEBUG] Ollama 使用原生 API 端点：{base_url}')

    # 构建请求头
    headers = {
        'Content-Type': 'application/json'
    }
    headers.update(model_config.headers)
    
    # 根据请求格式构建 payload
    if model_config.provider_id == 'ollama':
        # Ollama 原生 API 格式
        payload = {
            'model': model_config.model,
            'messages': [
                {'role': 'user', 'content': prompt}
            ],
            'stream': False  # 非流式返回
        }
    elif model_config.request_format == "dashscope":
        # 阿里云百炼格式
        headers['Authorization'] = f'Bearer {model_config.api_key}'
        payload = {
            'model': model_config.model,
            'input': {
                'messages': [
                    {'role': 'user', 'content': prompt}
                ]
            },
            'parameters': {
                'max_tokens': model_config.max_tokens,
                'temperature': model_config.temperature
            }
        }
    elif model_config.request_format == "openai":
        # OpenAI 兼容格式
        # Ollama 本地模型不需要 API Key
        if model_config.provider_id != 'ollama':
            headers['Authorization'] = f'Bearer {model_config.api_key}'

        payload = {
            'model': model_config.model,
            'messages': [
                {'role': 'user', 'content': prompt}
            ],
            'max_tokens': model_config.max_tokens,
            'temperature': model_config.temperature
        }

        # Ollama 专用参数优化（如果使用 OpenAI 兼容接口）
        if model_config.provider_id == 'ollama':
            payload.update({
                'top_p': 0.85,
                'top_k': 40,
                'repeat_penalty': 1.15,
                'stop': ['\n\n\n', '###'],
                'num_predict': min(model_config.max_tokens, 800)
            })
            payload['temperature'] = min(model_config.temperature, 0.5)

        # 智谱 AI GLM-4.7 需要 thinking 参数
        # 注意：GLM-4.7-Flash 等快速模型默认启用 thinking，需要显式禁用
        if model_config.provider_id == "zhipu" and "glm-4" in model_config.model.lower():
            # GLM-4.7-Flash 默认启用 thinking，需要显式禁用
            if "flash" in model_config.model.lower():
                payload['thinking'] = {"type": "disabled"}
            # 只有非 Flash 版本且 max_tokens 足够大时才启用 thinking
            elif model_config.max_tokens > 200:
                payload['thinking'] = {"type": "enabled"}
                # GLM-4.7 支持更大的 max_tokens
                if model_config.max_tokens > 8000:
                    payload['max_tokens'] = min(model_config.max_tokens, 65536)
    elif model_config.request_format == "ollama":
        # Ollama 原生格式
        payload = {
            'model': model_config.model,
            'messages': [
                {'role': 'user', 'content': prompt}
            ],
            'stream': False,
            'options': {
                'temperature': model_config.temperature,
                'num_predict': model_config.max_tokens,
                'top_p': 0.85,
                'top_k': 40,
                'repeat_penalty': 1.15,
                'stop': ['\n\n\n', '###']
            }
        }
    else:
        # 自定义格式（待扩展）
        return f"[API 错误] 不支持的请求格式: {model_config.request_format}"

    try:
        # 使用默认超时时间
        timeout = model_config.timeout
        
        print(f'[DEBUG] 发送请求到：{base_url}')
        print(f'[DEBUG] 请求头：{headers}')
        print(f'[DEBUG] 请求体：{json.dumps(payload, ensure_ascii=False)[:500]}...')

        response = requests.post(
            base_url,
            headers=headers,
            json=payload,
            timeout=timeout
        )

        print(f'[DEBUG] 响应状态码：{response.status_code}')
        
        if response.status_code == 200:
            result = response.json()
            # 使用智能提取逻辑
            content = _extract_ai_response(result, model_config)
            if content and not content.startswith('[API'):
                # 检测并修复重复内容
                content = detect_and_fix_repetition(content)
                return content.strip()
            else:
                return content
        else:
            error_msg = response.text
            try:
                error_json = response.json()
                if 'error' in error_json:
                    error_msg = error_json['error'].get('message', error_msg)
                elif 'message' in error_json:
                    error_msg = error_json['message']
            except:
                pass
            return f"[API 调用失败] 状态码：{response.status_code}, 错误：{error_msg}"

    except requests.exceptions.Timeout:
        return "[API 调用超时] 请重试"
    except requests.exceptions.RequestException as e:
        return f"[API 调用异常] {str(e)}"


def call_bailian_api(
    prompt: str,
    api_key: str,
    model: str = 'qwen-max',
    max_tokens: int = 2000,
    temperature: float = 0.7
) -> str:
    """调用百炼 API 生成内容（兼容旧接口）

    Args:
        prompt: 提示词
        api_key: 百炼 API Key
        model: 模型名称（qwen-max/qwen-plus/qwen-turbo）
        max_tokens: 最大生成 token 数
        temperature: 温度参数（0-1）

    Returns:
        AI 生成的文本内容
    """
    # 尝试从配置中获取模型配置
    config = get_model_config(model)
    if config and config.api_key:
        # 使用配置中的 API Key
        return call_ai_api(prompt, config)
    
    # 使用传入的 API Key 创建临时配置
    from model_config import ModelProvider
    temp_config = ModelConfig(
        id=model,
        name=model,
        provider=ModelProvider.DASHSCOPE.value,
        model=model,
        api_key=api_key,
        base_url="https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation",
        max_tokens=max_tokens,
        temperature=temperature,
        request_format="dashscope",
        response_path="output.choices.0.message.content"
    )
    return call_ai_api(prompt, temp_config)


def extract_info_field(section_title: str, requirement_text: str, 
                       ai_call_func: Optional[Callable[[str], str]] = None,
                       model_config: Optional[ModelConfig] = None) -> Optional[str]:
    """从需求文档中提取信息型字段

    采用三级提取策略：
    1. 正则匹配（快速、无成本）
    2. AI 语义提取（精准、理解上下文）

    支持匹配包含编号的章节标题，如 "1.1 项目名称" 也能匹配 "项目名称"

    Args:
        section_title: 章节标题
        requirement_text: 需求文档内容
        ai_call_func: 可选的 AI 调用函数
        model_config: 模型配置（用于 AI 提取）

    Returns:
        提取的信息，如果未找到则返回 None
    """
    # 确定要匹配的字段名
    target_field = None
    if section_title in INFO_FIELDS:
        target_field = section_title
    else:
        # 尝试模糊匹配（章节标题可能包含编号，如 "1.1 项目名称"）
        for field_name in INFO_FIELDS:
            if field_name in section_title:
                target_field = field_name
                break
    
    if target_field:
        # 第 1 级：尝试正则匹配（快速、无成本）
        patterns = INFO_PATTERNS.get(target_field, [])
        for pattern in patterns:
            match = re.search(pattern, requirement_text, re.IGNORECASE)
            if match:
                result = match.group(1).strip()
                if target_field != section_title:
                    print(f'[INFO] 模糊匹配提取：{section_title} -> {target_field} = {result}')
                return result
        
        # 第 2 级：如果提供了 AI 函数，使用 AI 语义提取
        if ai_call_func and model_config:
            try:
                print(f'[INFO] 正则匹配失败，尝试 AI 提取：{section_title}')
                ai_result = _extract_info_field_with_ai(
                    target_field, requirement_text, ai_call_func, model_config
                )
                if ai_result:
                    print(f'[INFO] AI 提取成功：{section_title} = {ai_result[:50]}...')
                    return ai_result
            except Exception as e:
                print(f'[WARN] AI 提取失败：{e}，返回 None')
    
    return None


def _extract_info_field_with_ai(target_field: str, requirement_text: str,
                                 ai_call_func: Callable[[str], str],
                                 model_config: ModelConfig) -> Optional[str]:
    """使用 AI 提取信息型字段

    Args:
        target_field: 目标字段名（如"项目名称"）
        requirement_text: 需求文档内容
        ai_call_func: AI 调用函数
        model_config: 模型配置

    Returns:
        提取的值，如果未找到则返回 None
    """
    # 构建 AI 提取提示
    prompt = f"""请从以下需求文档中提取【{target_field}】的具体值。

【需求文档内容】
{requirement_text[:3000]}

【提取要求】
1. 只提取"{target_field}"的具体值，不要提取其他信息
2. 如果文档中没有明确提及"{target_field}"，请输出 null
3. 直接输出值即可，不要包含"{target_field}"字样
4. 如果同一信息有多个表述，选择最完整、最正式的一个

【输出格式】
直接输出"{target_field}"的值，如果未找到则输出 null

示例：
- 如果提取到"项目名称：杭海路未来社区建设项目"，输出：杭海路未来社区建设项目
- 如果未找到相关信息，输出：null
"""
    
    # 临时调整模型配置以获取简短回答
    original_max_tokens = model_config.max_tokens
    model_config.max_tokens = 50  # 限制输出长度
    
    try:
        result = ai_call_func(prompt, model_config)

        # 清理结果
        if result:
            result = result.strip()
            # 移除可能的引号
            if result.startswith('"') and result.endswith('"'):
                result = result[1:-1]
            # 检查是否为 null
            if result.lower() in ['null', 'none', '无', '未找到', '']:
                return None
            return result
        return None
    finally:
        # 恢复原始配置
        model_config.max_tokens = original_max_tokens


def generate_section_content_with_ai(
    section_title: str,
    requirement_text: str = '',
    template_text: str = '',
    user_instruction: str = '',
    api_key: str = '',
    model: str = 'qwen-max',
    extract_data_points: bool = True,
    fallback_to_template: bool = False,
    model_config: Optional[ModelConfig] = None
) -> str:
    """使用 AI 生成章节内容

    Args:
        section_title: 章节标题（如"项目概况"）
        requirement_text: 需求文档内容
        template_text: 模板文档内容
        user_instruction: 用户补充要求
        api_key: 百炼 API Key（已弃用，使用 model_config）
        model: 模型名称（已弃用，使用 model_config）
        extract_data_points: 是否提取数据点（默认 True）
        fallback_to_template: AI 失败时是否降级使用模板（默认 False，返回错误信息）
        model_config: 模型配置（推荐方式）

    Returns:
        AI 生成的章节内容，或错误信息（如果 fallback_to_template=False）
    """
    # 调试信息
    print(f'\n[DEBUG] ====== 生成章节: {section_title} ======')
    print(f'[DEBUG] requirement_text 长度: {len(requirement_text) if requirement_text else 0}')
    print(f'[DEBUG] template_text 长度: {len(template_text) if template_text else 0}')
    print(f'[DEBUG] user_instruction: {user_instruction[:50] if user_instruction else "无"}...')
    
    # 获取模型配置
    if model_config is None:
        # 尝试从配置管理器获取
        model_config = get_model_config(model)
        if model_config is None:
            # 创建临时配置（兼容旧接口）
            from model_config import ModelProvider
            model_config = ModelConfig(
                id=model,
                name=model,
                provider=ModelProvider.DASHSCOPE.value,
                model=model,
                api_key=api_key,
                base_url="https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation",
                max_tokens=2000,
                temperature=0.7,
                request_format="dashscope",
                response_path="output.choices.0.message.content"
            )
    
    print(f'[DEBUG] 使用模型: {model_config.name} ({model_config.id})')
    print(f'[DEBUG] 提供商: {model_config.provider_id}')
    print(f'[DEBUG] API Key 是否为空: {not model_config.api_key}')
    
    # 首次生成时，从需求文档提取初始数据点
    if extract_data_points and requirement_text:
        dp_manager = get_data_point_manager()
        initial_data = dp_manager.extract_from_text(requirement_text, "需求文档")
        if initial_data:
            conflicts = dp_manager.update(initial_data, "需求文档")
            if conflicts:
                print(f'[WARN] 需求文档中发现数据冲突: {conflicts}')
            print(f'[INFO] 从需求文档提取数据点: {list(initial_data.keys())}')
        else:
            print(f'[WARN] 未能从需求文档提取数据点')

    # 判断是否为信息型字段
    is_info_field = section_title in INFO_FIELDS
    # 检查是否包含信息型字段关键词
    contains_info_keyword = any(field in section_title for field in INFO_FIELDS)
    
    print(f'[DEBUG] 字段类型判断: section_title={section_title}')
    print(f'[DEBUG]   - 是否在INFO_FIELDS: {is_info_field}')
    print(f'[DEBUG]   - 是否包含信息型关键词: {contains_info_keyword}')
    
    if is_info_field or contains_info_keyword:
        # 尝试从需求文档提取（使用正则 + AI 两级提取）
        extracted = extract_info_field(
            section_title, 
            requirement_text,
            ai_call_func=call_ai_api,  # 传入 AI 调用函数
            model_config=model_config
        )
        if extracted:
            print(f'[INFO] 提取信息型字段：{section_title} = {extracted}')
            # 更新数据点
            if extract_data_points:
                dp_manager = get_data_point_manager()
                dp_manager.update({section_title: extracted}, section_title)
            return extracted

        # 提取失败，使用 AI 生成简短内容
        print(f'[INFO] 信息型字段提取失败，使用AI生成: {section_title}')
        prompt = build_info_field_prompt(section_title, requirement_text, template_text)
        max_tokens = 100  # 信息型字段限制字数
    else:
        # 描述型字段，详细生成（带数据注入）
        print(f'[INFO] 描述型字段，使用AI生成: {section_title}')
        prompt = build_desc_field_prompt(section_title, requirement_text, template_text, user_instruction)
        max_tokens = 800  # 描述型字段允许更多字数

    # 调用 AI API
    print(f'[DEBUG] 调用AI API: max_tokens={max_tokens}')
    # 更新模型配置的 max_tokens
    model_config.max_tokens = max_tokens
    content = call_ai_api(prompt, model_config)
    print(f'[DEBUG] AI返回内容长度: {len(content) if content else 0}')
    print(f'[DEBUG] AI返回内容前100字: {content[:100] if content else "空"}...')

    # 检查 AI 调用是否失败
    if content.startswith('[') and ('API' in content or 'Error' in content or 'error' in content):
        error_msg = f"[AI 生成失败] 章节 '{section_title}': {content}"
        print(f'[ERROR] {error_msg}')
        
        if fallback_to_template:
            # 降级使用模板
            print(f'[WARNING] 降级使用模板: {section_title}')
            from app import get_chapter_content_template
            content = get_chapter_content_template(section_title)
        else:
            # 返回错误信息，不降级
            print(f'[DEBUG] ====== 完成章节: {section_title} (失败) ======\n')
            return error_msg

    # 清理 AI 生成内容
    content = clean_ai_content(content, section_title)
    print(f'[DEBUG] 清理后内容长度: {len(content) if content else 0}')

    # 内容去重检测（新增功能）
    if content:
        optimizer = get_content_optimizer()
        duplicate_result = optimizer.check_duplicate(content, threshold=0.6)
        if duplicate_result["is_duplicate"]:
            print(f'[WARN] 检测到重复内容：与 {duplicate_result["duplicate_sections"][0]["section_title"]} 相似度 {duplicate_result["duplicate_sections"][0]["similarity"]:.1%}')
            # 记录但不阻止生成，由后续审校处理
        # 添加到历史记录
        optimizer.add_generated_content(section_title, content)
    

    # 从生成内容中提取数据点（增量更新）
    if extract_data_points and content:
        dp_manager = get_data_point_manager()
        new_data = dp_manager.extract_from_text(content, section_title)
        if new_data:
            conflicts = dp_manager.update(new_data, section_title)
            if conflicts:
                print(f'[WARN] 章节 [{section_title}] 数据冲突:')
                for c in conflicts:
                    print(f'  - {c}')
            print(f'[INFO] 从 [{section_title}] 提取新数据点: {list(new_data.keys())}')

    print(f'[DEBUG] ====== 完成章节: {section_title} ======\n')
    return content


def extract_requirements_from_text(requirement_text: str, 
                                    ai_call_func: Optional[Callable] = None) -> Dict:
    """从需求文档提取需求点
    
    Args:
        requirement_text: 需求文档内容
        ai_call_func: AI 调用函数（可选）
        
    Returns:
        需求点字典 {'pain_points': [...], 'requirements': [...], 'goals': [...]}
    """
    req_analyzer = get_requirement_analyzer()
    result = req_analyzer.extract(requirement_text, ai_call_func)
    print(f'[INFO] 提取需求点: {len(result["pain_points"])} 痛点, '
          f'{len(result["requirements"])} 需求, {len(result["goals"])} 目标')
    return result


def get_optimization_summary() -> Dict:
    """获取优化摘要（数据点和需求点）
    
    Returns:
        优化摘要字典
    """
    dp_manager = get_data_point_manager()
    req_analyzer = get_requirement_analyzer()
    
    return {
        'data_points': dp_manager.get_extraction_summary(),
        'requirements': req_analyzer.get_summary()
    }


def build_info_field_prompt(section_title: str, requirement_text: str, template_text: str) -> str:
    """构建信息型字段的 Prompt（智能提取或生成）"""
    return f"""你是一位专业的文档信息提取和生成专家。

【任务】为【{section_title}】提供内容

【需求文档】
{requirement_text[:1500] if requirement_text else '无'}

【要求】
1. 优先从需求文档中提取准确的{section_title}
2. 如果文档中没有明确的{section_title}，请根据上下文智能生成一个合理的、专业的{section_title}
3. **只输出{section_title}的内容，不要解释，不要输出分析过程、思考过程**
4. **直接输出结果，不要使用任何 Markdown 格式（如 **粗体**、## 标题等）**
5. 不超过 50 字
6. 内容要专业、符合可行性研究报告规范

【{section_title}】
"""


def extract_template_section(section_title: str, template_text: str) -> str:
    """从模板文档中提取指定章节的内容
    
    Args:
        section_title: 章节标题
        template_text: 模板文档全文
        
    Returns:
        该章节的模板内容
    """
    if not template_text or not section_title:
        return ""
    
    # 尝试找到章节标题位置
    # 支持多种格式："1.1 项目名称"、"项目名称"、"1.1项目名称"等
    patterns = [
        rf'{re.escape(section_title)}[\s\n]*',
        rf'{re.escape(section_title.split()[-1] if " " in section_title else section_title)}[\s\n]*',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, template_text, re.IGNORECASE)
        if match:
            start_pos = match.end()
            # 查找下一个章节标题位置（或文档结束）
            # 假设下一个章节以数字或"第"开头
            next_section_match = re.search(r'\n\s*(?:\d+\.|第[一二三四五六七八九十]+章|【)', template_text[start_pos:])
            if next_section_match:
                end_pos = start_pos + next_section_match.start()
                section_content = template_text[start_pos:end_pos].strip()
            else:
                section_content = template_text[start_pos:].strip()
            
            # 限制长度
            if len(section_content) > 1500:
                section_content = section_content[:1500] + "..."
            
            return section_content
    
    # 如果没找到特定章节，返回文档开头部分作为通用参考
    return template_text[:800] if template_text else ""


def build_desc_field_prompt(section_title: str, requirement_text: str, template_text: str, user_instruction: str) -> str:
    """构建描述型字段的 Prompt（详细论述，带数据注入、Few-shot 示例和类型识别）
    
    核心优化：
    1. 从需求文档提取项目类型和业务场景，明确告知 AI 应该写什么业务
    2. 强化指令，明确区分"内容来源"和"格式参考"
    3. 添加撰写原则，确保 AI 围绕需求文档的业务来写
    """
    # 调试输出
    print(f'[DEBUG] 生成描述型字段：{section_title}')
    print(f'[DEBUG] 需求文档长度：{len(requirement_text) if requirement_text else 0}')
    print(f'[DEBUG] 模板参考长度：{len(template_text) if template_text else 0}')
    
    # 从需求文档提取项目类型和业务场景
    project_type = extract_project_type(requirement_text)
    business_scenes = extract_business_scenes(requirement_text)
    print(f'[DEBUG] 识别项目类型：{project_type if project_type else "未识别"}')
    print(f'[DEBUG] 识别业务场景：{business_scenes if business_scenes else "未识别"}')

    # 获取数据点管理器和需求分析器
    dp_manager = get_data_point_manager()
    req_analyzer = get_requirement_analyzer()
    optimizer = get_content_optimizer()

    # 注入已确立的数据点
    data_points_text = dp_manager.get_formatted_prompt_text()

    # 注入本章应回应的需求点
    requirements_text = req_analyzer.get_requirements_text(section_title)

    # 提取该章节对应的模板内容
    template_section = extract_template_section(section_title, template_text)
    print(f'[DEBUG] 提取模板章节内容长度：{len(template_section)}')

    # 识别章节类型
    type_info = optimizer.identify_section_type(section_title)
    print(f'[DEBUG] 章节类型识别：{type_info["type"]}型 ({type_info["subtype"]})')

    # 获取 Few-shot 示例
    few_shot_prompt = optimizer.get_few_shot_prompt(section_title)
    if few_shot_prompt:
        print(f'[DEBUG] 已加载 Few-shot 示例')

    # 根据章节类型添加特定的格式指导
    format_guidance = get_format_guidance(section_title, template_section)

    # 构建类型指导
    type_guidance = ''
    if type_info['format_strategy']:
        type_guidance = f"\n【格式策略】{type_info['format_strategy']}"
    
    # 构建项目类型和业务场景提示
    project_type_hint = ''
    if project_type:
        project_type_hint = f"- **项目类型**：{project_type}\n"
    if business_scenes:
        scenes_text = '\n'.join(f"   - {scene}" for scene in business_scenes)
        project_type_hint += f"- **业务场景**：\n{scenes_text}\n"

    return f"""你是一位专业的可行性研究报告编写专家。

【任务】请为【{section_title}】这一章节撰写内容。

【重要提示】
- **内容来源**：必须严格基于【需求文档内容】中的业务信息，这是你撰写内容的唯一来源
- **模板用途**：【参考模板】仅用于学习格式、结构、论述方式，不要直接复用其中的具体业务信息
- **撰写原则**：围绕上述项目类型和业务场景进行撰写，确保内容贴合实际需求
- 只生成【{section_title}】这一章节的内容，不要涉及其他章节
- **严禁**在内容中重复或罗列其他章节的标题
- **重要**：直接输出章节正文内容，不要输出任何分析过程、思考过程、元信息
- **重要**：不要使用 Markdown 格式（如 **粗体**、## 标题等），直接输出纯文本

【项目信息】
{project_type_hint if project_type_hint else '- 项目类型：根据需求文档确定'}

【章节标题】{section_title}

{data_points_text}

【需求文档内容】（这是你撰写内容的唯一来源）
{requirement_text[:2000] if requirement_text else '无相关需求文档'}

{requirements_text}

【参考模板（仅学习格式，不要复制业务内容）】
{template_section if template_section else '无参考模板'}

【用户补充要求】
{user_instruction if user_instruction else '无特殊要求'}
{type_guidance}
{few_shot_prompt}
【输出要求】
1. **内容来源**：必须来自需求文档，根据项目类型和业务场景进行撰写
2. **格式参考**：参考模板的段落结构、列表形式、标准编号等格式特征
3. 内容要专业、准确、逻辑清晰，使用正式的公文语言
4. **重要**：引用数据时必须与【已确立的关键数据】保持完全一致
5. **重要**：必须回应【本章应回应的需求点】中的每一个要点
{format_guidance}
6. 结合需求文档中的具体场景和问题，不要泛泛而谈
7. **不要输出章节标题**（如"{section_title}"）
8. **不要生成总结性段落**（如"通过实施本项目..."）
9. **不要输出任何分析过程、思考过程、元信息**
10. 字数控制在 200-400 字（除非模板格式要求更多内容）

【请开始撰写，直接输出正文内容】
"""


def get_format_guidance(section_title: str, template_section: str) -> str:
    """根据章节类型获取特定的格式指导"""
    guidance = []
    
    # 检查章节标题关键词
    title_lower = section_title.lower()
    
    # 政策法规/技术规范类章节
    if any(keyword in title_lower for keyword in ['政策', '法规', '规范', '标准', '依据']):
        guidance.append("   - **本章节特殊要求**：请列出相关的法律法规、政策文件或技术标准")
        guidance.append("   - 使用列表形式，每条包含标准编号（如GB/T、DB33/T等）和标准名称")
        guidance.append("   - 参考模板的列举方式，保持格式一致")
    
    # 如果模板内容包含标准编号格式，明确要求保持
    if template_section:
        # 检查是否包含标准编号模式
        if re.search(r'[A-Z]+/T?\s*\d+[-–]\d{4}', template_section):
            guidance.append("   - **检测到标准编号格式**：请确保生成的内容包含相应的标准编号（如GB/T、DB、ISO等）")
        # 检查是否是列表格式
        if re.search(r'^\s*\d+\.', template_section, re.MULTILINE):
            guidance.append("   - **检测到列表格式**：请使用类似的数字列表格式（1. 2. 3.）")
    
    return '\n'.join(guidance) if guidance else ''


def extract_template_section(section_title: str, template_text: str) -> str:
    """从模板文档中提取指定章节的内容
    
    支持匹配:
    - "第三章 建设必要性"
    - "3.1 建设必要性"
    - "建设必要性"
    
    Args:
        section_title: 章节标题
        template_text: 模板文档全文
        
    Returns:
        该章节的模板内容
    """
    if not template_text or not section_title:
        return ""
    
    # 清理章节标题，提取核心部分
    core_title = re.sub(r'^第 [一二三四五六七八九十\d]+[章条节款]\s*', '', section_title)
    
    # 尝试多种模式匹配
    patterns = [
        rf'{re.escape(section_title)}\s*\n',
        rf'{re.escape(core_title)}\s*\n',
        rf'\d+[\.\s]+{re.escape(core_title)}\s*\n',
    ]
    
    for pattern in patterns:
        match = re.search(pattern, template_text, re.IGNORECASE)
        if match:
            start_pos = match.end()
            # 查找下一个章节标题位置
            next_section = re.search(
                r'\n\s*(?:第 [一二三四五六七八九十\d]+[章条节款]|\d+[\.\s])',
                template_text[start_pos:]
            )
            if next_section:
                end_pos = start_pos + next_section.start()
            else:
                end_pos = len(template_text)
            
            content = template_text[start_pos:end_pos].strip()
            # 限制长度，避免过长
            return content[:2000] + ('...' if len(content) > 2000 else '')
    
    # 未找到特定章节，返回空字符串
    return ""


def get_document_structure(template_text: str) -> str:
    """提取文档结构（章节标题列表）
    
    Args:
        template_text: 模板文档全文
        
    Returns:
        章节标题列表
    """
    structure = []
    patterns = [
        r'^第 ([一二三四五六七八九十\d]+)[章条节款]\s*(.+)$',
        r'^(\d+[\.\s]+)(.+)$',
    ]
    
    for line in template_text.split('\n'):
        line = line.strip()
        for pattern in patterns:
            match = re.match(pattern, line)
            if match:
                structure.append(line)
                break
    
    return '\n'.join(structure[:30])  # 最多 30 个章节


def clean_ai_content(content: str, section_title: str) -> str:
    """清理 AI 生成的内容

    Args:
        content: AI 生成的原始内容
        section_title: 章节标题

    Returns:
        清理后的内容
    """
    if not content:
        return ''

    # ===== 第一步：过滤 AI 思考过程 =====
    # AI 思考过程的典型标记
    thinking_markers = [
        r'^\*\*分析请求：\*\*',
        r'^\*\*目标：\*\*',
        r'^\*\*约束条件：\*\*',
        r'^\*\*输入数据：\*\*',
        r'^\*\*输出要求：\*\*',
        r'^\*\*请开始撰写\*\*',
        r'^\d+\.\s*分析请求',
        r'^\d+\.\s*分析需求',
        r'^\d+\.\s*确定.*内容',
        r'^\d+\.\s*起草内容',
        r'^\d+\.\s*起草策略',
        r'^\d+\.\s*优化内容',
        r'^\d+\.\s*格式检查',
        r'^角色：',
        r'^任务：',
        r'^约束条件：',
        r'^输入数据：',
        r'^输出要求：',
        r'^起草策略：',
        r'^优化内容：',
        r'^格式检查：',
        r'^尝试\s*\d+',
        r'^\*\*尝试\s*\d+\*\*',
        # GLM-4.7 思考过程标记
        r'^\*\*\s*1\.\s*分析请求\s*\*\*',
        r'^\*\*\s*2\.\s*分析需求\s*\*\*',
        r'^\*\*\s*3\.\s*确定.*\*\*',
        r'^\*\*\s*4\.\s*起草.*\*\*',
        r'^\*\*\s*5\.\s*优化.*\*\*',
        r'^\*\*\s*6\.\s*格式检查.*\*\*',
        r'^\*\*分析请求：\*\*',
        r'^\*\*目标：\*\*',
        r'^\*\*输入：\*\*',
        r'^\*\*输出：\*\*',
        r'^\*\*关键信息：\*\*',
        r'^\*\*提取结果：\*\*',
        r'^\*\*开始生成.*\*\*',
        r'^\*\*开始撰写.*\*\*',
        r'^【分析请求】',
        r'^【目标】',
        r'^【输入】',
        r'^【输出】',
        r'^\*\*1\.\*\*',
        r'^\*\*2\.\*\*',
        r'^\*\*3\.\*\*',
        r'^\*\*4\.\*\*',
        r'^\*\*5\.\*\*',
        r'^\*\*6\.\*\*',
    ]

    # 检测内容是否包含思考过程
    has_thinking = False
    for marker in thinking_markers:
        if re.search(marker, content, re.MULTILINE | re.IGNORECASE):
            has_thinking = True
            print(f'[INFO] 检测到 AI 思考过程，需要过滤')
            break

    if has_thinking:
        # 尝试提取实际内容（通常在 "尝试" 或数字编号之后）
        # 策略：找到第一个符合正文格式的行（如 "1. XXX" 或 "GB/T" 或 "第X条"）
        lines = content.split('\n')
        actual_content_start = 0

        for i, line in enumerate(lines):
            stripped = line.strip()
            if not stripped:
                continue

            # 检测思考过程标记
            is_thinking_line = False
            for marker in thinking_markers:
                if re.search(marker, stripped, re.IGNORECASE):
                    is_thinking_line = True
                    break

            # 检测思考过程中的元信息行
            if re.search(r'^(项目名称|建设单位|总人口|重点|核心|具体场景)', stripped):
                is_thinking_line = True

            if is_thinking_line:
                actual_content_start = i + 1
                continue

            # 检测到实际内容开始的标记
            # 1. 数字编号列表（如 "1. "、"1.1 "）
            # 2. 标准编号（如 "GB/T"、"DB33/T"）
            # 3. 书名号（如 "《XXX》"）
            if re.search(r'^\d+[\.\s]', stripped) or \
               re.search(r'^(GB/T|DB33/T|GB|DB|ISO|IEC)', stripped) or \
               stripped.startswith('《'):
                actual_content_start = i
                print(f'[INFO] 找到实际内容起始行 {i}: {stripped[:50]}...')
                break

        # 截取实际内容
        if actual_content_start > 0:
            content = '\n'.join(lines[actual_content_start:])
            print(f'[INFO] 已过滤思考过程，保留实际内容')

    # ===== 第二步：清理其他章节混入和格式 =====
    # 常见章节标题模式（用于检测混入的其他章节）
    other_section_patterns = [
        r'^\s*\d+[\.\s]+建设目标',
        r'^\s*\d+[\.\s]+建设规模',
        r'^\s*\d+[\.\s]+建设内容',
        r'^\s*\d+[\.\s]+建设周期',
        r'^\s*\d+[\.\s]+投资估算',
        r'^\s*\d+[\.\s]+资金筹措',
        r'^\s*\d+[\.\s]+项目效益',
        r'^\s*\d+[\.\s]+效益分析',
        r'^\s*\d+[\.\s]+风险分析',
        r'^\s*\d+[\.\s]+结论',
        r'^\s*\d+[\.\s]+建议',
        r'^\s*\d+[\.\s]+项目概况',
        r'^\s*\d+[\.\s]+项目背景',
        r'^\s*\d+[\.\s]+需求分析',
        r'^\s*\d+[\.\s]+技术方案',
        r'^\s*\d+[\.\s]+总体框架',
        r'^\s*第[一二三四五六七八九十]+章',
        r'^\s*第\s*\d+\s*章',
        r'^\s*\d+\.\d+\s+',  # 如 "1.1 项目名称"
        r'^\s*【[^】]+】',     # 如 "【建设目标】"
    ]

    lines = content.split('\n')
    cleaned_lines = []
    found_other_section = False

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 跳过当前章节标题行
        if stripped.startswith(section_title):
            continue
        if stripped.startswith('#'):
            continue

        # 检测是否出现其他章节标题
        for pattern in other_section_patterns:
            if re.search(pattern, stripped, re.IGNORECASE):
                # 如果检测到其他章节标题，截断后续内容
                print(f'[WARN] 检测到混入的其他章节标题，截断: {stripped[:50]}...')
                found_other_section = True
                break

        if found_other_section:
            break

        # 检测总结性段落的开头
        if re.search(r'^(通过实施本项目|综上所述|总之|本项目旨在|本次建设)', stripped):
            print(f'[WARN] 检测到总结性段落，截断: {stripped[:50]}...')
            break

        # 清理 Markdown 格式符号
        # 移除 ** 粗体符号
        stripped = re.sub(r'\*\*(.+?)\*\*', r'\1', stripped)
        # 移除 * 斜体符号
        stripped = re.sub(r'\*(.+?)\*', r'\1', stripped)
        # 移除 ` 代码符号
        stripped = re.sub(r'`(.+?)`', r'\1', stripped)
        # 移除 # 标题符号
        stripped = re.sub(r'^#+\s*', '', stripped)
        # 移除 - 列表符号（但保留内容）
        stripped = re.sub(r'^[-•●*]\s+', '', stripped)

        cleaned_lines.append(stripped)

    return '\n'.join(cleaned_lines)


def detect_and_fix_repetition(content: str, max_repeat: int = 3) -> str:
    """检测并修复 AI 生成内容中的重复问题
    
    Args:
        content: AI 生成的内容
        max_repeat: 最大允许重复次数
        
    Returns:
        修复后的内容
    """
    if not content:
        return content
    
    lines = content.split('\n')
    cleaned_lines = []
    repeat_count = {}
    
    for line in lines:
        stripped = line.strip()
        if not stripped:
            cleaned_lines.append(line)
            continue
        
        # 统计每行的重复次数
        repeat_count[stripped] = repeat_count.get(stripped, 0) + 1
        
        # 如果重复次数超过阈值，跳过
        if repeat_count[stripped] > max_repeat:
            print(f'[WARN] 检测到重复内容，跳过：{stripped[:50]}... (第{repeat_count[stripped]}次)')
            continue
        
        cleaned_lines.append(line)
    
    # 检查是否大部分内容都在重复
    if len(lines) > 5:
        unique_ratio = len(set(line.strip() for line in lines if line.strip())) / len([l for l in lines if l.strip()])
        if unique_ratio < 0.3:  # 唯一内容少于 30%
            print(f'[WARN] 内容重复度过高 ({unique_ratio:.1%})，可能生成失败')
            return f"[AI 生成质量异常] 内容重复度过高，请重试或更换模型"
    
    return '\n'.join(cleaned_lines)


# 测试函数
if __name__ == '__main__':
    # 简单测试
    test_api_key = 'sk-cc3bee3fa06c4987b52756db0abb7991'
    result = call_bailian_api('你好，请用一句话介绍你自己。', test_api_key)
    print(f'AI 回复：{result}')


def get_chapter_content_template(section_title, requirement_content=''):
    """根据章节标题返回预定义的内容模板
    优化：区分信息型字段和描述型字段

    信息型字段：简短、具体（如项目名称、总投资等）
    描述型字段：需要论述、分析（如项目概况、建设必要性等）
    """
    # 信息型字段 - 简短、具体
    info_templates = {
        '项目名称': '根据需求文档提取的项目名称',
        '项目建设单位': '根据需求文档提取的建设单位名称',
        '负责人': 'XXX',
        '联系方式': '电话：XXX-XXXXXXX',
        '建设工期': 'XX 个月',
        '总投资': 'XX 万元',
        '资金来源': '财政拨款/自筹',
        '编制单位': 'XX 数字科技有限公司',
    }

    # 描述型字段 - 需要论述、分析（每段 200-300 字）
    desc_templates = {
        '建设目标': '本项目的建设目标是根据实际需求，制定科学合理的建设方案，确保项目顺利实施并达到预期效果。通过本项目的实施，将有效提升相关领域的信息化水平，改善工作效率，优化服务质量，为业务发展提供有力支撑。项目建设遵循统筹规划、分步实施的原则，确保各项建设任务有序推进。',
        '项目概况': '本项目按照相关规范和要求进行编制，确保符合可行性研究报告的标准。项目立足于当前实际需求，采用先进的技术路线和科学的管理方法，力求在技术先进性、经济合理性和实施可行性之间取得最佳平衡。项目建设内容包括硬件设施、软件系统、数据资源等多个方面。',
        '项目建设单位概况': '项目建设单位的基本情况介绍，包括单位性质、主要职能、组织架构等内容。单位在相关领域具有丰富的经验和雄厚的技术实力，为本项目的顺利实施提供了坚实保障。单位现有人员配置合理，技术力量充足，能够满足项目建设和运营的需要。',
        '项目实施机构': '项目实施机构负责项目的具体实施工作，包括项目管理、协调推进、质量控制等职责。机构设置合理，人员配备充足，能够确保项目按计划高质量完成。实施机构建立了完善的管理制度和工作流程，为项目顺利实施提供了组织保障。',
        '项目建设的必要性': '从政策要求、实际需求、发展趋势等方面论证项目建设的必要性。当前，随着业务的不断发展和信息化水平的提升，现有系统已无法满足日益增长的需求，亟需通过本项目的建设来解决相关问题。项目建设对于提升工作效率、优化服务质量具有重要意义。',
        '需求分析': '对项目的需求进行全面分析，包括业务需求、功能需求、性能需求等。通过深入调研和分析，明确了项目的核心需求和关键指标，为后续的方案设计提供了重要依据。需求分析结果将为项目建设内容的确定和投资估算提供参考。',
        '总体思路': '项目建设的总体思路是坚持以需求为导向，以技术为支撑，确保项目建设的科学性和可行性。在实施过程中，注重技术创新与管理创新相结合，力求实现最佳的建设效果。项目建设将充分利用现有资源，避免重复建设，提高投资效益。',
        '总体框架': '项目总体框架包括架构设计、功能模块、技术路线等核心内容。框架设计遵循先进性、实用性、可扩展性的原则，能够满足当前需求并适应未来发展。总体框架的确定为后续详细设计和技术实现提供了指导。',
        '技术路线': '项目采用成熟可靠的技术路线，确保系统的稳定性、安全性和可扩展性。在技术选型上，充分考虑了当前技术发展趋势和项目实施风险，选择了最适合的技术方案。技术路线的确定将为项目实施提供技术保障。',
        '投资估算': '投资估算包括硬件设备、软件开发、系统集成、工程建设等费用。估算依据充分，计算方法科学，能够真实反映项目建设所需的资金投入。投资估算结果为项目资金筹措和使用计划提供了参考依据。',
        '资金筹措': '项目资金来源明确，筹措方案可行，确保项目顺利实施。资金安排合理，能够满足项目各阶段的资金需求。资金筹措方案的确定为项目顺利实施提供了资金保障。',
        '效益分析': '项目效益包括经济效益、社会效益、环境效益等多个方面。通过综合分析，项目具有良好的投资回报和显著的社会效益。效益分析结果为项目投资决策提供了重要参考。',
        '风险分析': '项目风险包括技术风险、管理风险、市场风险等，需制定相应的风险应对措施。通过风险识别和评估，制定了完善的风险管理方案。风险分析和应对措施的制定将为项目顺利实施提供保障。',
        '结论': '综合各方面分析，项目具备建设的必要性和可行性。项目建设条件成熟，实施方案可行，建议尽快启动实施。结论为项目投资决策提供了科学依据。',
        '建议': '建议尽快启动项目实施，并做好组织保障、资金保障、技术保障等工作。在实施过程中，要加强项目管理，确保项目按计划高质量完成。同时，要建立健全项目运营维护机制，确保项目长期稳定运行。',
    }

    # 先精确匹配
    if section_title in info_templates:
        return info_templates[section_title]
    if section_title in desc_templates:
        return desc_templates[section_title]

    # 再模糊匹配
    for key, content in info_templates.items():
        if key in section_title:
            return content
    for key, content in desc_templates.items():
        if key in section_title:
            return content

    # 默认返回
    return f'{section_title}的具体内容根据项目实际情况进行编制。'


def create_partial_document(task_id, chapters, template_type='future_community'):
    """
    根据章节内容创建临时文档（用于编辑后更新）
    
    Args:
        task_id: 任务 ID
        chapters: 章节列表（包含内容和标题）
        template_type: 模板类型
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import os

    print(f'[DEBUG] create_partial_document 被调用: task_id={task_id}, chapters={len(chapters)}')

    # 创建文档
    doc = Document()

    # 设置默认样式
    style = doc.styles['Normal']
    style.font.name = '仿宋'
    style.font.size = Pt(16)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')

    # 添加标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Cm(5)
    run = title.add_run('文档预览')
    run.font.name = '黑体'
    run.font.size = Pt(36)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.add_page_break()

    # 添加自动目录（TOC 域）
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    doc.add_paragraph()
    
    # 添加 TOC 域
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    run._r.append(instrText)
    
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    doc.add_page_break()

    # 添加正文内容
    for ch in chapters:
        level = ch.get('level', 1)
        number = ch.get('number', '')
        title_text = ch.get('title', '')
        content = ch.get('content', '')

        # 修复：title 已经包含编号，不需要再添加
        heading_text = title_text
        if level <= 3:
            # 使用 Word 的 Heading 样式（1-3级）
            heading = doc.add_heading(heading_text, level=level)
            # 设置中文字体
            for run in heading.runs:
                run.font.name = '黑体' if level == 1 else '楷体'
                run.font.size = Pt(22 if level == 1 else 16)
                run.font.bold = True
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体' if level == 1 else '楷体')
        else:
            # 4级及以上使用普通段落
            heading = doc.add_paragraph()
            run = heading.add_run(heading_text)
            run.font.name = '楷体'
            run.font.size = Pt(14)
            run.font.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
            heading.paragraph_format.first_line_indent = Cm(0.74)
        
        # 添加内容（如果有）
        if content:
            for para_text in content.split('\n'):
                if para_text.strip():
                    para = doc.add_paragraph()
                    para.paragraph_format.first_line_indent = Cm(0.74)
                    run = para.add_run(para_text.strip())
                    run.font.name = '仿宋'
                    run.font.size = Pt(16)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        
        doc.add_paragraph()  # 章节间距
    
    # 保存文档
    from task_manager import get_task_manager
    task_manager = get_task_manager()
    task_dir = task_manager._get_task_directory(task_id)
    filepath = os.path.join(task_dir, f'partial_{task_id}.docx')
    
    doc.save(filepath)
    print(f'[DEBUG] 临时文档已更新: {filepath}')
    
    return filepath

