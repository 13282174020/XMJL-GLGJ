# -*- coding: utf-8 -*-
"""模板文档分析器模块"""
import re
from dataclasses import dataclass, field
from typing import Dict, List, Set
from collections import defaultdict


@dataclass
class NumberingPattern:
    pattern: str
    suggested_level: int
    count: int = 0
    example_texts: List[str] = field(default_factory=list)


@dataclass
class StyleAnalysis:
    style_name: str
    total_count: int = 0
    is_heading: bool = False
    heading_level: int = 0
    numbering_patterns: Dict[str, NumberingPattern] = field(default_factory=dict)


@dataclass
class AnalysisReport:
    file_path: str
    total_paragraphs: int = 0
    styles_usage: Dict[str, int] = field(default_factory=dict)
    heading_styles: Dict[str, StyleAnalysis] = field(default_factory=dict)
    potential_heading_styles: List[str] = field(default_factory=list)

    def to_dict(self) -> dict:
        """转换为字典格式"""
        return {
            'file_path': self.file_path,
            'total_paragraphs': self.total_paragraphs,
            'styles_usage': self.styles_usage,
            'heading_styles': {
                k: {
                    'style_name': v.style_name,
                    'is_heading': v.is_heading,
                    'heading_level': v.heading_level,
                    'numbering_patterns': {
                        pk: {
                            'pattern': pv.pattern,
                            'suggested_level': pv.suggested_level,
                            'count': pv.count
                        } for pk, pv in v.numbering_patterns.items()
                    }
                } for k, v in self.heading_styles.items()
            },
            'potential_heading_styles': self.potential_heading_styles
        }

    def generate_mapping_rules(self) -> List[dict]:
        """生成样式映射规则"""
        rules = []
        
        # 规则 1: 标准标题样式保持不变
        for style_name, analysis in self.heading_styles.items():
            if analysis.is_heading and analysis.heading_level > 0:
                rules.append({
                    'source_style': style_name,
                    'pattern': None,
                    'target_style': f'Heading {analysis.heading_level}',
                    'rule_type': 'heading'
                })
        
        # 规则 2: 潜在标题样式根据编号格式映射
        # 注意：潜在标题样式可能已经在 heading_styles 中（is_heading=False）
        for style_name in self.potential_heading_styles:
            analysis = self.heading_styles.get(style_name)
            if analysis and not analysis.is_heading and analysis.numbering_patterns:
                # 统计该样式下各编号模式的数量，只生成最有效的规则
                total_numbered = sum(p.count for p in analysis.numbering_patterns.values())
                
                # 只有当该样式有足够多的编号段落时才生成规则（避免误判）
                if total_numbered >= 1:
                    for pattern_str, pattern_info in analysis.numbering_patterns.items():
                        rules.append({
                            'source_style': style_name,
                            'pattern': pattern_str,
                            'target_style': f'Heading {pattern_info.suggested_level}',
                            'rule_type': 'numbering'
                        })
        
        return rules


class DocumentAnalyzer:
    def __init__(self):
        # 编号正则表达式（按级别从高到低排序）
        # 修复：更精确的编号匹配，排除日期等
        self.numbering_regex_patterns = [
            # 数字编号格式（X.X.X.X 到 X.X）
            (r'^(\d+)\.(\d+)\.(\d+)\.(\d+)', 4),  # 4 级编号 X.X.X.X
            (r'^(\d+)\.(\d+)\.(\d+)', 3),          # 3 级编号 X.X.X
            (r'^(\d+)\.(\d+)', 2),                 # 2 级编号 X.X
            
            # 中文编号格式
            (r'^第 [一二三四五六七八九十百千\d]+章', 1),  # 第一章、第 1 章
            (r'^第 [一二三四五六七八九十百千\d]+[节条]', 2),  # 第一节、第 1 条
            (r'^第 [一二三四五六七八九十百千\d]+款', 3),  # 第一款
            
            # 中文数字编号
            (r'^[一二三四五六七八九十]+[、.]', 2),  # 一、或一.
            (r'^[一二三四五六七八九十]+', 1),  # 一、二、三
            
            # 括号编号
            (r'^\(([一二三四五六七八九十\d]+)\)', 3),  # (一)、(1)
            (r'^（([一二三四五六七八九十\d]+)）', 3),  # （一）、（1）
            
            # 带括号的数字编号
            (r'^\((\d+)\)', 3),  # (1)、(2)
            
            # 圆圈数字
            (r'^①', 3),
        ]
        
        # 需要排除的日期模式
        self.date_patterns = [
            r'^\d{4}年\d{1,2}月',  # 2024 年 01 月
            r'^\d{4}-\d{1,2}-\d{1,2}',  # 2024-01-01
            r'^\d{4}/\d{1,2}/\d{1,2}',  # 2024/01/01
        ]
    
    def analyze(self, file_path: str = None, file_stream = None) -> AnalysisReport:
        from docx import Document
        import io
        
        # 支持从文件路径或文件流加载
        if file_stream is not None:
            doc = Document(file_stream)
        elif file_path is not None:
            doc = Document(file_path)
        else:
            raise ValueError("必须提供 file_path 或 file_stream")
        
        styles_usage = defaultdict(int)
        heading_styles = {}
        potential_heading_styles = set()
        style_numbering = defaultdict(dict)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else 'Normal'
            styles_usage[style_name] += 1
            level = self._get_heading_level(style_name, text)

            if level > 0:
                if style_name not in heading_styles:
                    heading_styles[style_name] = StyleAnalysis(
                        style_name=style_name, is_heading=True, heading_level=level)
                heading_styles[style_name].total_count += 1
            else:
                # 检查是否是日期模式，如果是则跳过
                is_date = False
                for date_pattern in self.date_patterns:
                    if re.match(date_pattern, text):
                        is_date = True
                        break
                
                if not is_date:
                    # 检查编号模式
                    for pattern, suggested_level in self.numbering_regex_patterns:
                        if re.match(pattern, text):
                            potential_heading_styles.add(style_name)
                            if pattern not in style_numbering[style_name]:
                                style_numbering[style_name][pattern] = NumberingPattern(
                                    pattern=pattern, suggested_level=suggested_level)
                            pattern_info = style_numbering[style_name][pattern]
                            pattern_info.count += 1
                            if len(pattern_info.example_texts) < 3:
                                pattern_info.example_texts.append(text[:50])
                            break

        for style_name in potential_heading_styles:
            if style_name not in heading_styles:
                heading_styles[style_name] = StyleAnalysis(
                    style_name=style_name, is_heading=False, heading_level=0,
                    total_count=styles_usage.get(style_name, 0))
            heading_styles[style_name].numbering_patterns = style_numbering.get(style_name, {})

        return AnalysisReport(
            file_path=file_path, total_paragraphs=len(doc.paragraphs),
            styles_usage=dict(styles_usage), heading_styles=heading_styles,
            potential_heading_styles=list(potential_heading_styles))
    
    def _get_heading_level(self, style_name: str, text: str = '') -> int:
        if not style_name:
            return 0
        style_lower = style_name.lower()
        if 'heading' in style_lower:
            try:
                num_str = ''.join(filter(str.isdigit, style_lower.replace('heading', '')))
                if num_str:
                    level = int(num_str)
                    if 1 <= level <= 6:
                        return level
            except:
                pass
        if '标题' in style_lower:
            try:
                num_str = ''.join(filter(str.isdigit, style_lower.replace('标题', '')))
                if num_str:
                    level = int(num_str)
                    if 1 <= level <= 6:
                        return level
            except:
                pass
        return 0
