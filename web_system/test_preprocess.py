# -*- coding: utf-8 -*-
"""测试 web_system 的模板预处理功能"""
import sys
import logging
import os

sys.path.insert(0, '.')

# 配置日志
logging.basicConfig(level=logging.INFO, format='%(levelname)s - %(message)s')

from template_preprocessor import TemplatePreprocessor

# 使用英文文件名的测试文件
template_path = os.path.abspath('sample_standard.docx')
output_path = os.path.abspath('outputs/test_web_preprocess.docx')

print('=' * 80)
print('测试 web_system 模板预处理功能')
print('=' * 80)
print(f'模板路径：{template_path}')
print(f'输出路径：{output_path}')

# 检查文件是否存在
if not os.path.exists(template_path):
    print(f'错误：模板文件不存在！')
    sys.exit(1)

preprocessor = TemplatePreprocessor()
result = preprocessor.preprocess_with_analysis(template_path, output_path)

print(f'\n结果：{result.message}')
print(f'统计：{result.stats}')
print(f'输出文件：{result.output_path}')

# 验证生成的文档
if result.success and result.output_path:
    from docx import Document
    doc = Document(result.output_path)
    
    print(f'\n生成的文档：{len(doc.paragraphs)} 个段落')
    
    styles = {}
    heading_count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else 'Normal'
        styles[style_name] = styles.get(style_name, 0) + 1
        if 'Heading' in style_name:
            heading_count += 1
    
    print(f'\n样式分布:')
    for style, count in sorted(styles.items(), key=lambda x: -x[1])[:10]:
        print(f'  {style}: {count}')
    
    print(f'\n标题段落数：{heading_count}')
    
    print(f'\n前 20 个非空段落:')
    count = 0
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else 'Normal'
        print(f'  [{style_name}] {text[:50]}')
        count += 1
        if count >= 20:
            break
