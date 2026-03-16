# -*- coding: utf-8 -*-
"""模板预处理引擎模块"""
import re
import shutil
import tempfile
import os
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional, Dict, Tuple
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm


@dataclass
class PreprocessResult:
    success: bool
    output_path: Optional[str]
    message: str
    stats: Dict[str, int]

    def to_dict(self):
        return {'success': self.success, 'output_path': self.output_path,
                'message': self.message, 'stats': self.stats}


def remove_numbering_prefix(text: str) -> Tuple[str, str]:
    """
    删除文本开头的编号前缀
    
    Args:
        text: 原始文本
        
    Returns:
        (编号，纯文本) 元组
    """
    original_text = text
    
    # 定义编号前缀的正则表达式（按优先级排序）
    patterns = [
        # 数字编号：X.X.X.X 到 X.X（多级编号）
        (r'^(\d+\.\d+\.\d+\.\d+\.\d+\.\d+)\s*', '6 级数字'),
        (r'^(\d+\.\d+\.\d+\.\d+\.\d+)\s*', '5 级数字'),
        (r'^(\d+\.\d+\.\d+\.\d+)\s*', '4 级数字'),
        (r'^(\d+\.\d+\.\d+)\s*', '3 级数字'),
        (r'^(\d+\.\d+)\s*', '2 级数字'),
        
        # 中文章节：第 X 章、第 X 节（优先于单数字）
        (r'^(第 [一二三四五六七八九十百千\d]+[章条节款])\s*', '中文章节'),
        
        # 单数字编号：X. 后必须有空格，避免匹配 4.0 版本
        (r'^(\d+\.)\s+', '1 级数字'),
        
        # 中文数字：一、二、
        (r'^([一二三四五六七八九十]+[、.])\s*', '中文数字'),
        
        # 括号编号：（一）、(1) - 返回完整括号
        (r'^（[一二三四五六七八九十\d]+）\s*', '中文括号'),
        (r'^\(\d+\)\s*', '英文括号数字'),
        (r'^\([一二三四五六七八九十]+\)\s*', '英文括号中文'),
        
        # 圆圈数字：① ②
        (r'^(①|②|③|④|⑤|⑥|⑦|⑧|⑨|⑩)\s*', '圆圈数字'),
    ]
    
    for pattern, pattern_type in patterns:
        match = re.match(pattern, text)
        if match:
            numbering = match.group(1) if match.lastindex else match.group(0)
            pure_text = text[match.end():].strip()
            
            # 特殊处理：如果删除编号后文本太短（< 2 字），可能编号是文本的一部分
            if len(pure_text) < 2:
                continue
            
            # 特殊处理：避免删除版本号如 4.0
            if pattern_type == '1 级数字':
                # 检查原始文本是否是版本号格式（数字。数字，后面没有空格）
                # 如 "4.0 版本" 不删除，"4. 管理后台" 删除
                if re.match(r'^\d+\.\d', text.split()[0] if text.split() else text):
                    continue
            
            return numbering, pure_text
    
    # 没有匹配到编号
    return '', original_text


class TemplatePreprocessor:
    def __init__(self):
        self.standard_styles = ['Heading 1', 'Heading 2', 'Heading 3',
                                'Heading 4', 'Heading 5', 'Heading 6', 'Normal']

    def preprocess_with_analysis(self, input_path: str, output_path: str) -> PreprocessResult:
        """分析文档并根据目录结构生成新文档 - 保留正文内容"""
        import logging
        logging.info('[PREPROCESS] 开始分析并生成新文档')
        logging.info(f'[PREPROCESS] 输入文件：{input_path}')
        logging.info(f'[PREPROCESS] 输出文件：{output_path}')
        
        try:
            # 1. 分析文档
            from template_analyzer import DocumentAnalyzer
            analyzer = DocumentAnalyzer()
            report = analyzer.analyze(input_path)
            
            logging.info(f'[PREPROCESS] 分析完成：{len(report.heading_styles)} 种样式')
            logging.info(f'[PREPROCESS] 潜在标题样式：{report.potential_heading_styles}')
            
            # 2. 生成映射规则
            rules = report.generate_mapping_rules()
            logging.info(f'[PREPROCESS] 生成 {len(rules)} 条规则')
            for rule in rules:
                logging.info(f'[PREPROCESS]   规则：{rule["source_style"]} + {rule.get("pattern", "None")} -> {rule["target_style"]}')
            
            # 3. 复制模板文档并逐段处理（保留正文内容）
            logging.info('[PREPROCESS] 复制模板文档并转换样式...')
            doc = self._copy_template(input_path)
            
            # 4. 应用样式转换
            logging.info('[PREPROCESS] 应用样式转换...')
            stats = self._apply_style_conversions(doc, rules)
            
            # 5. 保存文档
            output_path_obj = Path(output_path)
            output_path_obj.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            logging.info(f'[PREPROCESS] 文档已保存：{output_path}')
            logging.info(f'[PREPROCESS] 统计：{stats}')
            
            return PreprocessResult(True, str(output_path),
                                   f'处理完成，{sum(stats.values())} 个段落已转换', stats)
        except Exception as e:
            import traceback
            logging.error(f'[PREPROCESS] 异常：{e}')
            logging.error(traceback.format_exc())
            return PreprocessResult(False, None, f'处理失败：{e}', {})

    def _copy_template(self, template_path: str) -> Document:
        """复制模板文档并确保 Heading 样式有正确的大纲级别"""
        import logging
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        # 创建临时文件复制模板
        temp_dir = tempfile.mkdtemp()
        temp_template = os.path.join(temp_dir, 'template.docx')
        shutil.copy2(template_path, temp_template)

        # 加载模板文档
        doc = Document(temp_template)

        logging.info(f'[PREPROCESS] 模板文档：{len(doc.paragraphs)} 个段落，{len(doc.tables)} 个表格')

        # 确保 Heading 样式有正确的大纲级别（用于多级列表）
        self._ensure_heading_outline_levels(doc)

        # 清理临时文件
        os.remove(temp_template)
        os.rmdir(temp_dir)

        return doc

    def _ensure_heading_outline_levels(self, doc: Document) -> None:
        """确保 Heading 1-9 样式有正确的大纲级别（用于多级列表）"""
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import logging
        
        logging.info('[PREPROCESS] 设置 Heading 样式的大纲级别...')
        
        for i in range(1, 10):
            style_name = f'Heading {i}'
            try:
                style = doc.styles[style_name]
                pPr = style._element.pPr
                if pPr is not None:
                    # 设置 outlineLvl
                    outlineLvl = pPr.outlineLvl
                    if outlineLvl is None:
                        outlineLvl = OxmlElement('w:outlineLvl')
                        pPr.insert(0, outlineLvl)
                    outlineLvl.set(qn('w:val'), str(i - 1))
                    logging.info(f'[PREPROCESS]   {style_name}: outlineLvl={i-1}')
            except Exception as e:
                logging.debug(f'[PREPROCESS]   {style_name}: 设置失败 - {e}')

    def _apply_style_conversions(self, doc: Document, rules: List[dict]) -> Dict[str, int]:
        """应用样式转换 - 保留正文内容，删除编号前缀"""
        import logging
        
        stats = {}
        removed_numbering_count = 0
        
        # 编译正则表达式规则
        compiled_rules = []
        for rule in rules:
            compiled_rule = rule.copy()
            if rule.get('pattern'):
                try:
                    compiled_rule['compiled_pattern'] = re.compile(rule['pattern'])
                except re.error as e:
                    logging.error(f'[PREPROCESS] 正则错误：{e}')
                    continue
            compiled_rules.append(compiled_rule)
        
        # 逐段处理
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            
            current_style = para.style.name if para.style else 'Normal'
            new_style = self._find_matching_style(current_style, text, compiled_rules)
            
            if new_style and new_style != current_style:
                # 应用新样式
                try:
                    para.style = new_style
                except:
                    # 如果样式不存在，尝试设置基本格式
                    self._apply_basic_heading_style(para, new_style)
                
                # 删除编号前缀（仅对 Heading 样式）
                if 'Heading' in new_style:
                    numbering, pure_text = remove_numbering_prefix(text)
                    if numbering:
                        # 更新段落文本
                        para.clear()
                        run = para.add_run(pure_text)
                        removed_numbering_count += 1
                        logging.debug(f'[PREPROCESS] 删除编号：{numbering} | {pure_text[:30]}')
                
                key = f'{current_style}->{new_style}'
                stats[key] = stats.get(key, 0) + 1
                logging.debug(f'[PREPROCESS] 段落 {i}: {current_style} -> {new_style} | {text[:30]}')
            else:
                stats['unchanged'] = stats.get('unchanged', 0) + 1
        
        stats['removed_numbering'] = removed_numbering_count
        logging.info(f'[PREPROCESS] 转换完成：{sum(v for k, v in stats.items() if k not in ["unchanged", "removed_numbering"])} 个段落已转换，{removed_numbering_count} 个编号前缀已删除')
        return stats

    def _find_matching_style(self, style_name: str, text: str, rules: List[dict]) -> Optional[str]:
        """查找匹配的样式"""
        for rule in rules:
            if rule['source_style'] != style_name:
                continue
            
            if rule.get('compiled_pattern'):
                if rule['compiled_pattern'].match(text):
                    return rule['target_style']
            elif rule.get('pattern') is None:
                return rule['target_style']
        
        return None

    def _apply_basic_heading_style(self, para, style_name: str) -> None:
        """应用基本标题样式（当样式不存在时）"""
        # 设置段落格式
        para.paragraph_format.space_before = Pt(12)
        para.paragraph_format.space_after = Pt(6)
        
        # 设置 run 格式
        for run in para.runs:
            if style_name == 'Heading 1':
                run.font.bold = True
                run.font.size = Pt(16)
            elif style_name == 'Heading 2':
                run.font.bold = True
                run.font.size = Pt(14)
            elif style_name == 'Heading 3':
                run.font.bold = True
                run.font.size = Pt(12)
            else:
                run.font.bold = True
                run.font.size = Pt(10.5)

    def preprocess(self, input_path: str, output_path: str,
                   rules: List[dict]) -> PreprocessResult:
        """原有的预处理方法（样式转换）"""
        try:
            doc = Document(input_path)
            stats = {}

            compiled_rules = []
            for rule in rules:
                compiled_rule = rule.copy()
                if rule.get('pattern'):
                    try:
                        compiled_rule['compiled_pattern'] = re.compile(rule['pattern'])
                    except re.error as e:
                        return PreprocessResult(False, None, f'正则错误：{e}', {})
                compiled_rules.append(compiled_rule)

            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                current_style = para.style.name if para.style else 'Normal'
                new_style = self._apply_rules(current_style, text, compiled_rules)

                if new_style and new_style != current_style:
                    para.style = new_style
                    key = f'{current_style}->{new_style}'
                    stats[key] = stats.get(key, 0) + 1
                else:
                    stats['unchanged'] = stats.get('unchanged', 0) + 1

            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))

            total = sum(v for k, v in stats.items() if k != 'unchanged')
            return PreprocessResult(True, str(output_path),
                                   f'处理完成，{total} 个段落已转换', stats)
        except Exception as e:
            return PreprocessResult(False, None, f'处理失败：{e}', {})

    def _apply_rules(self, style_name: str, text: str, rules: List[dict]) -> Optional[str]:
        for rule in rules:
            if rule['source_style'] != style_name:
                continue
            if rule.get('compiled_pattern'):
                if rule['compiled_pattern'].match(text):
                    return rule['target_style']
            elif rule.get('pattern') is None:
                return rule['target_style']
        return None
