# -*- coding: utf-8 -*-
"""
未来社区建设方案生成系统 - Flask 后端 (AI 增强版)
支持调用 Qwen API 生成定制化内容

V2.0 更新：集成 AI 内容优化
- 数据一致性管理
- 需求覆盖度检查
- 全文质量审校
"""

import os
import uuid
import json
import requests
import threading
import time
from enum import Enum
from functools import wraps
from flask import Flask, render_template, request, jsonify, send_file, make_response
from flask_cors import CORS
from werkzeug.utils import secure_filename
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import base64

# 导入 AI 内容优化服务
from services.data_point_manager import DataPointManager
from services.requirement_analyzer import RequirementAnalyzer
from services.quality_reviewer import QualityReviewer
from model_config import get_model_config_manager, ModelConfig, get_enabled_models, get_model_config
from task_manager import get_task_manager, TaskInfo, ChapterStatus, MAX_REGENERATE_COUNT
from ai_engine import get_chapter_content_template

app = Flask(__name__)
CORS(app)

app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['OUTPUT_FOLDER'] = 'outputs'
app.config['MAX_CONTENT_LENGTH'] = 50 * 1024 * 1024
app.config['SECRET_KEY'] = os.urandom(24).hex()

os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['OUTPUT_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'docx', 'doc', 'txt', 'pdf'}

# ========== 异步任务管理模块 ==========

class TaskStatus(Enum):
    """任务状态枚举"""
    PENDING = 'pending'              # 等待处理
    PROCESSING = 'processing'        # 处理中
    PARSING_FILE = 'parsing_file'    # 解析文件中
    GENERATING_AI = 'generating_ai'  # AI 推理生成中
    CREATING_DOC = 'creating_doc'    # 生成 Word 文档中
    COMPLETED = 'completed'          # 已完成
    FAILED = 'failed'                # 失败
    CANCELLED = 'cancelled'          # 已取消


class TaskManager:
    """任务管理器 - 单例模式"""
    _instance = None
    _lock = threading.Lock()

    def __new__(cls):
        if cls._instance is None:
            with cls._lock:
                if cls._instance is None:
                    cls._instance = super().__new__(cls)
                    cls._instance._initialized = False
        return cls._instance

    def __init__(self):
        if self._initialized:
            return
        self._initialized = True
        self.tasks = {}  # task_id -> task_info
        self._lock = threading.Lock()

    def create_task(self, task_type, user_prompt='', api_key='', model='qwen-max',
                    template_type='future_community', requirement_file=None, template_file=None):
        """创建新任务"""
        task_id = uuid.uuid4().hex
        with self._lock:
            self.tasks[task_id] = {
                'task_id': task_id,
                'template_type': template_type,
                'task_type': task_type,
                'user_prompt': user_prompt,
                'api_key': api_key,
                'model': model,
                'status': TaskStatus.PENDING.value,
                'progress': 0,
                'message': '任务已提交，等待处理',
                'created_at': datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
                'started_at': None,
                'completed_at': None,
                'output_filename': None,
                'error_message': None,
                'requirement_file': requirement_file,
                'template_file': template_file,
                'chapter_steps': [],  # 记录章节生成步骤
                'current_chapter': None,  # 当前正在生成的章节
                'completed_chapters': [],  # 已完成的章节列表
                'pending_confirmation': False,  # 是否等待确认
                'partial_filename': None  # 部分生成的文档文件名
            }
        return task_id

    def get_task_status(self, task_id):
        """获取任务状态"""
        with self._lock:
            return self.tasks.get(task_id)

    def update_task_progress(self, task_id, progress=None, message=None, status=None):
        """更新任务进度"""
        if not task_id:
            return False
        with self._lock:
            if task_id not in self.tasks:
                return False
            task = self.tasks[task_id]
            if progress is not None:
                task['progress'] = max(0, min(100, progress))
            if message is not None:
                task['message'] = message
            if status is not None:
                task['status'] = status
            return True

    def mark_task_completed(self, task_id, output_filename=None):
        """标记任务完成"""
        with self._lock:
            if task_id not in self.tasks:
                return False
            task = self.tasks[task_id]
            task['status'] = TaskStatus.COMPLETED.value
            task['progress'] = 100
            task['message'] = '文档生成成功'
            task['completed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            if output_filename:
                task['output_filename'] = output_filename
            return True

    def mark_task_failed(self, task_id, error_message):
        """标记任务失败"""
        with self._lock:
            if task_id not in self.tasks:
                return False
            task = self.tasks[task_id]
            task['status'] = TaskStatus.FAILED.value
            task['error_message'] = error_message
            task['message'] = f'生成失败：{error_message}'
            task['completed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            return True

    def cancel_task(self, task_id):
        """取消任务"""
        with self._lock:
            if task_id not in self.tasks:
                return False
            task = self.tasks[task_id]
            if task['status'] in [TaskStatus.COMPLETED.value, TaskStatus.FAILED.value]:
                return False  # 已完成或失败的任务不能取消
            task['status'] = TaskStatus.CANCELLED.value
            task['message'] = '任务已取消'
            task['completed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            return True

    def get_task_list(self, page=1, page_size=10, keyword=None, status_filter=None):
        """获取任务列表（分页）"""
        with self._lock:
            tasks = list(self.tasks.values())

            # 关键词过滤
            if keyword:
                tasks = [t for t in tasks if keyword.lower() in t.get('template_type', '').lower()
                        or keyword in t.get('task_id', '').lower()]

            # 状态过滤
            if status_filter:
                tasks = [t for t in tasks if t.get('status') == status_filter]

            # 按创建时间倒序
            tasks.sort(key=lambda x: x.get('created_at', ''), reverse=True)

            total = len(tasks)
            start = (page - 1) * page_size
            end = start + page_size

            return {
                'tasks': tasks[start:end],
                'total': total,
                'page': page,
                'page_size': page_size
            }

    def set_task_started(self, task_id):
        """标记任务开始处理"""
        with self._lock:
            if task_id in self.tasks:
                self.tasks[task_id]['started_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                self.tasks[task_id]['status'] = TaskStatus.PROCESSING.value

    def add_chapter_step(self, task_id, chapter_index, chapter_title, status='pending'):
        """添加章节生成步骤记录"""
        with self._lock:
            if task_id in self.tasks:
                step = {
                    'chapter_index': chapter_index,
                    'chapter_title': chapter_title,
                    'status': status,  # pending, generating, completed, failed
                    'started_at': None,
                    'completed_at': None,
                    'message': ''
                }
                self.tasks[task_id]['chapter_steps'].append(step)
                return True
            return False

    def update_chapter_step(self, task_id, chapter_index, status=None, message=None):
        """更新章节生成步骤状态"""
        with self._lock:
            if task_id in self.tasks:
                for step in self.tasks[task_id]['chapter_steps']:
                    if step['chapter_index'] == chapter_index:
                        if status:
                            step['status'] = status
                            if status == 'generating':
                                step['started_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                            elif status in ['completed', 'failed']:
                                step['completed_at'] = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                        if message:
                            step['message'] = message
                        return True
                return False
            return False

    def set_current_chapter(self, task_id, chapter_title):
        """设置当前正在生成的章节"""
        with self._lock:
            if task_id in self.tasks:
                self.tasks[task_id]['current_chapter'] = chapter_title
                return True
            return False

    def add_completed_chapter(self, task_id, chapter_title):
        """添加已完成的章节"""
        with self._lock:
            if task_id in self.tasks:
                self.tasks[task_id]['completed_chapters'].append(chapter_title)
                return True
            return False

    def set_pending_confirmation(self, task_id, pending):
        """设置等待确认状态"""
        with self._lock:
            if task_id in self.tasks:
                self.tasks[task_id]['pending_confirmation'] = pending
                return True
            return False

    def set_partial_filename(self, task_id, filename):
        """设置部分生成的文档文件名"""
        with self._lock:
            if task_id in self.tasks:
                self.tasks[task_id]['partial_filename'] = filename
                return True
            return False


# 全局任务管理器实例
task_manager = TaskManager()

# 应用配置（用于百炼应用 API）
APP_CONFIGS = {
    'default': {'app_id': '2641738c46274ac0b08ba4520cc2313b'}  # 阿里云百炼文本推理智能体
}

# 模型配置（用于 API Key 验证提示）
MODEL_CONFIGS = {
    'qwen-max': {'name': '通义千问-Max', 'provider': '阿里云'},
    'qwen-plus': {'name': '通义千问-Plus', 'provider': '阿里云'},
    'qwen-turbo': {'name': '通义千问-Turbo', 'provider': '阿里云'},
    'glm-4': {'name': 'GLM-4', 'provider': '智谱 AI'},
    'glm-3-turbo': {'name': 'GLM-3-Turbo', 'provider': '智谱 AI'},
    'kimi-latest': {'name': 'Kimi', 'provider': '月之暗面'},
    'mini-max': {'name': 'MiniMax', 'provider': 'MiniMax'}
}

# 章节分类配置
# 信息提取类：直接从需求文档中提取信息，不需要 AI 生成
# 内容撰写类：需要 AI 根据模板和需求撰写内容
CHAPTER_CATEGORIES = {
    'info_extract': [  # 信息提取类章节
        '项目名称', '项目建设单位', '负责人', '联系方式',
        '建设工期', '总投资', '资金来源', '编制单位',
        '项目概况', '建设单位', '项目负责', '实施机构',
        '工期', '投资估算', '资金筹措'
    ],
    'content_write': [  # 内容撰写类章节
        '必要性', '分析', '方案', '技术路线',
        '效益分析', '风险分析', '结论', '建议',
        '背景', '需求', '设计', '建设内容',
        '系统', '安全', '环保', '节能',
        '招标', '组织', '人员', '培训',
        '进度', '财务', '评价', '市场',
        '产品', '服务', '运营', '管理'
    ]
}


def is_info_chapter(section_title):
    """判断是否为信息提取类章节"""
    for keyword in CHAPTER_CATEGORIES['info_extract']:
        if keyword in section_title:
            return True
    return False


def is_content_chapter(section_title):
    """判断是否为内容撰写类章节"""
    for keyword in CHAPTER_CATEGORIES['content_write']:
        if keyword in section_title:
            return True
    return False


# 模板类型配置
TEMPLATE_TYPES = {
    'future_community': {
        'name': '未来社区可行性研究报告',
        'chapters': [
            ('1 项目概况', ['1.1 项目名称', '1.2 项目建设单位及负责人', '1.3 方案编制依据', '1.4 建设目标、规模、内容、周期', '1.5 项目总投资及资金来源', '1.6 项目效益分析', '1.7 主要结论和建议']),
            ('2 项目建设单位概况', ['2.1 项目建设单位与职能', '2.2 项目实施机构与职责']),
            ('3 项目建设的必要性', ['3.1 项目提出的背景', '3.2 是否纳入一本帐', '3.3 项目建设的必要性']),
            ('4 需求分析', ['4.1 与政务职能相关的社会问题和政务目标分析', '4.2 现况分析及差距']),
            ('5 项目建设方案', ['5.1 总体思路', '5.2 总体框架', '5.3 技术路线', '5.4 建设目标', '5.5 建设内容']),
            ('6 与数字化改革总体方案的关系', ['6.1 基础设施', '6.2 数据资源', '6.3 应用系统', '6.4 公共组件建设']),
            ('7 应用系统设计', ['7.1 系统总体设计', '7.2 功能模块设计', '7.3 硬件系统设计']),
            ('8 国产化改造和适配', ['8.1 产品选型及平台适配', '8.2 应用安全可靠构建及部署']),
            ('9 系统安全', ['9.1 网络安全等级保护方案', '9.2 信息系统安全管理方案', '9.3 安全防护软件和设备']),
            ('10 项目招标方案', ['10.1 招标范围', '10.2 招标方式', '10.3 招标组织形式']),
            ('11 环保、消防、职业安全和节能', ['11.1 环境影响和环保措施', '11.2 消防措施', '11.3 职业安全和卫生措施', '11.4 节能目标及措施']),
            ('12 项目组织机构和人员', ['12.1 项目领导、实施和运维机构', '12.2 人员配置', '12.3 人员培训']),
            ('13 项目实施进度', ['13.1 项目建设工期', '13.2 实施进度计划']),
            ('14 投资估算和资金筹措', ['14.1 投资估算的有关说明', '14.2 项目总投资估算', '14.3 资金来源与落实情况', '14.4 资金使用计划']),
            ('15 效益与风险分析', ['15.1 经济效益分析', '15.2 社会效益分析', '15.3 项目风险与风险对策']),
        ]
    },
    'general_project': {
        'name': '通用项目可行性研究报告',
        'chapters': [
            ('1 项目总论', ['1.1 项目概况', '1.2 编制依据', '1.3 研究范围', '1.4 主要结论']),
            ('2 项目背景与必要性', ['2.1 项目背景', '2.2 项目建设的必要性']),
            ('3 市场分析', ['3.1 市场现状', '3.2 市场需求预测', '3.3 竞争分析']),
            ('4 建设方案', ['4.1 建设目标', '4.2 建设内容', '4.3 技术方案']),
            ('5 投资估算与资金筹措', ['5.1 投资估算', '5.2 资金筹措']),
            ('6 效益分析', ['6.1 经济效益', '6.2 社会效益']),
            ('7 风险分析', ['7.1 风险识别', '7.2 风险对策']),
        ]
    },
    'smart_community': {
        'name': '智慧社区建设方案',
        'chapters': [
            ('1 项目概述', ['1.1 项目背景', '1.2 建设目标', '1.3 建设内容']),
            ('2 需求分析', ['2.1 业务需求', '2.2 功能需求', '2.3 性能需求']),
            ('3 总体设计', ['3.1 架构设计', '3.2 技术路线', '3.3 安全设计']),
            ('4 应用系统建设', ['4.1 社区治理', '4.2 物业服务', '4.3 便民服务']),
            ('5 基础设施建设', ['5.1 网络建设', '5.2 感知设备', '5.3 数据中心']),
            ('6 投资估算', ['6.1 预算编制', '6.2 资金安排']),
            ('7 实施计划', ['7.1 进度安排', '7.2 保障措施']),
        ]
    },
    'business_plan': {
        'name': '商业计划书',
        'chapters': [
            ('1 执行摘要', ['1.1 项目简介', '1.2 商业模式', '1.3 核心优势', '1.4 融资需求']),
            ('2 公司概述', ['2.1 公司基本情况', '2.2 发展历程', '2.3 组织架构', '2.4 核心团队']),
            ('3 产品与服务', ['3.1 产品介绍', '3.2 技术特点', '3.3 服务模式', '3.4 研发规划']),
            ('4 市场分析', ['4.1 行业现状', '4.2 市场规模', '4.3 目标客户', '4.4 竞争格局']),
            ('5 商业模式', ['5.1 盈利模式', '5.2 营销策略', '5.3 运营计划', '5.4 合作伙伴']),
            ('6 财务预测', ['6.1 收入预测', '6.2 成本分析', '6.3 利润预测', '6.4 现金流分析']),
            ('7 融资方案', ['7.1 融资计划', '7.2 资金用途', '7.3 退出机制', '7.4 投资回报']),
            ('8 风险分析', ['8.1 市场风险', '8.2 技术风险', '8.3 管理风险', '8.4 应对策略']),
        ]
    },
    'project_proposal': {
        'name': '项目建议书',
        'chapters': [
            ('1 项目总论', ['1.1 项目概况', '1.2 编制依据', '1.3 研究结论']),
            ('2 项目背景', ['2.1 政策背景', '2.2 市场背景', '2.3 建设必要性']),
            ('3 市场预测', ['3.1 市场现状', '3.2 需求预测', '3.3 价格预测']),
            ('4 建设规模', ['4.1 建设规模', '4.2 产品方案', '4.3 技术标准']),
            ('5 建设条件', ['5.1 选址方案', '5.2 资源条件', '5.3 配套条件']),
            ('6 技术方案', ['6.1 技术方案', '6.2 设备方案', '6.3 工程方案']),
            ('7 环保节能', ['7.1 环境影响', '7.2 节能措施', '7.3 劳动安全']),
            ('8 投资估算', ['8.1 投资估算', '8.2 资金筹措', '8.3 用款计划']),
            ('9 效益分析', ['9.1 经济效益', '9.2 社会效益']),
            ('10 结论建议', ['10.1 主要结论', '10.2 建议']),
        ]
    },
    'funding_application': {
        'name': '资金申请报告',
        'chapters': [
            ('1 项目概况', ['1.1 项目基本情况', '1.2 建设单位概况', '1.3 编制依据']),
            ('2 政策符合性', ['2.1 产业政策', '2.2 行业准入', '2.3 资金支持方向']),
            ('3 市场分析', ['3.1 市场现状', '3.2 市场需求', '3.3 竞争优势']),
            ('4 技术方案', ['4.1 技术来源', '4.2 工艺流程', '4.3 设备选型']),
            ('5 建设方案', ['5.1 建设内容', '5.2 建设工期', '5.3 实施进度']),
            ('6 投资估算', ['6.1 总投资估算', '6.2 资金申请', '6.3 配套资金']),
            ('7 财务评价', ['7.1 基础数据', '7.2 财务指标', '7.3 敏感性分析']),
            ('8 风险分析', ['8.1 风险识别', '8.2 风险程度', '8.3 防范措施']),
        ]
    },
    'social_stability': {
        'name': '社会稳定风险评估报告',
        'chapters': [
            ('1 总论', ['1.1 项目概况', '1.2 评估依据', '1.3 评估范围']),
            ('2 风险调查', ['2.1 调查内容', '2.2 调查方式', '2.3 调查结果']),
            ('3 风险识别', ['3.1 风险因素', '3.2 风险类型', '3.3 主要风险']),
            ('4 风险估计', ['4.1 风险概率', '4.2 影响程度', '4.3 风险等级']),
            ('5 风险防范', ['5.1 防范措施', '5.2 应急预案', '5.3 责任主体']),
            ('6 评估结论', ['6.1 风险等级', '6.2 主要结论', '6.3 建议']),
        ]
    }
}


def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def clean_ai_content(content, section_title):
    """清理 AI 生成的内容，移除可能的标题行和 Markdown 格式"""
    if not content:
        return []

    lines = content.split('\n')
    cleaned_lines = []

    for line in lines:
        stripped = line.strip()
        if not stripped:
            continue

        # 跳过可能是标题的行（包含小节标题的行）
        if stripped.startswith(section_title) or stripped.startswith('《' + section_title + '》'):
            continue
        if stripped.startswith('#') or stripped.startswith('##') or stripped.startswith('###'):
            continue

        # 移除 Markdown 格式
        cleaned = stripped
        cleaned = re.sub(r'\*\*(.+?)\*\*', r'\1', cleaned)  # 移除 ** 粗体
        cleaned = re.sub(r'\*(.+?)\*', r'\1', cleaned)      # 移除 * 斜体
        cleaned = re.sub(r'`(.+?)`', r'\1', cleaned)        # 移除 ` 代码
        cleaned = re.sub(r'^[-•●]\s*', '', cleaned)         # 移除列表符号

        if cleaned:
            cleaned_lines.append(cleaned)

    return cleaned_lines


def extract_info_from_requirement(section_title, requirement_content):
    """从需求文档中提取信息类章节的内容"""
    if not requirement_content:
        return None
    
    # 项目名称
    if '项目名称' in section_title:
        patterns = [
            r'项目名称 [：:]\s*([^\n]+)',
            r'项目名称为 [：:]\s*([^\n]+)',
            r'项目名称是 [：:]\s*([^\n]+)',
            r'称 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
        # 尝试从第一行提取
        first_line = requirement_content.strip().split('\n')[0]
        if len(first_line) < 50 and '项目' in first_line:
            return first_line
    
    # 建设单位
    if '建设单位' in section_title or '项目实施单位' in section_title:
        patterns = [
            r'建设单位 [：:]\s*([^\n]+)',
            r'业主单位 [：:]\s*([^\n]+)',
            r'实施单位 [：:]\s*([^\n]+)',
            r'项目单位 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 负责人
    if '负责人' in section_title or '联系人' in section_title:
        patterns = [
            r'负责人 [：:]\s*([^\n]+)',
            r'联系人 [：:]\s*([^\n]+)',
            r'项目负责 [人]?[：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 联系方式/电话
    if '联系' in section_title or '电话' in section_title or '方式' in section_title:
        patterns = [
            r'联系方式 [：:]\s*([^\n]+)',
            r'联系电话 [：:]\s*([^\n]+)',
            r'电话 [：:]\s*([^\n]+)',
            r'手机 [：:]\s*([^\n]+)',
            r'邮箱 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 建设工期
    if '工期' in section_title or '建设周期' in section_title:
        patterns = [
            r'建设工期 [：:]\s*([^\n]+)',
            r'工期 [：:]\s*([^\n]+)',
            r'建设周期 [：:]\s*([^\n]+)',
            r'周期 [：:]\s*([^\n]+)',
            r'([\d]+)[个]?[年月天]'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 总投资
    if '投资' in section_title or '资金' in section_title:
        patterns = [
            r'总投资 [：:]\s*([^\n]+)',
            r'投资估算 [：:]\s*([^\n]+)',
            r'项目总投 [：:]\s*([^\n]+)',
            r'([\d\.]+)\s*万?元'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 资金来源
    if '资金来源' in section_title or '资金筹措' in section_title:
        patterns = [
            r'资金来源 [：:]\s*([^\n]+)',
            r'资金筹措 [：:]\s*([^\n]+)',
            r'出资 [：:]\s*([^\n]+)',
            r'财政拨款|自筹|银行贷款|专项资金'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(0).strip()
    
    # 编制单位
    if '编制' in section_title or '可研单位' in section_title:
        patterns = [
            r'编制单位 [：:]\s*([^\n]+)',
            r'可研编制 [：:]\s*([^\n]+)',
            r'咨询单位 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 建设目标
    if '建设目标' in section_title or '项目目标' in section_title:
        patterns = [
            r'建设目标 [：:]\s*([^\n]+)',
            r'项目目标 [：:]\s*([^\n]+)',
            r'目标 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 建设规模
    if '建设规模' in section_title or '规模' in section_title:
        patterns = [
            r'建设规模 [：:]\s*([^\n]+)',
            r'项目规模 [：:]\s*([^\n]+)',
            r'规模 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    # 建设内容
    if '建设内容' in section_title or '建设任务' in section_title:
        patterns = [
            r'建设内容 [：:]\s*([^\n]+)',
            r'建设任务 [：:]\s*([^\n]+)',
            r'主要建设内容 [：:]\s*([^\n]+)'
        ]
        for pattern in patterns:
            match = re.search(pattern, requirement_content, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    
    return None


def extract_template_example(section_title, template_content):
    """从模板内容中提取该章节的参考示例"""
    if not template_content:
        return ""
    
    # 找到章节标题位置
    title_pos = template_content.find(section_title)
    if title_pos == -1:
        # 尝试模糊匹配
        for keyword in [section_title.split(' ')[-1] if ' ' in section_title else section_title]:
            title_pos = template_content.find(keyword)
            if title_pos != -1:
                break
    
    if title_pos == -1:
        return ""
    
    # 提取章节后面的内容（到下一个章节前）
    next_chapter_pos = len(template_content)
    
    # 查找下一个章节标记
    chapter_markers = ['\n\n', '\n', '。', '；']
    for marker in chapter_markers:
        pos = template_content.find(marker, title_pos + len(section_title))
        if pos != -1 and pos < next_chapter_pos:
            # 检查是否是下一个章节标题
            next_text = template_content[title_pos + len(section_title):pos].strip()
            if len(next_text) < 50 and not any(c in next_text for c in '。！？'):
                next_chapter_pos = title_pos + len(section_title) + len(next_text)
                break
    
    # 提取示例内容（取前 500 字符）
    example_start = title_pos + len(section_title)
    example_end = min(example_start + 500, next_chapter_pos)
    example = template_content[example_start:example_end].strip()
    
    # 清理示例内容，去掉过多的空白
    example = ' '.join(example.split())
    
    return example if len(example) > 10 else ""


def read_docx_text(file_path):
    """读取 Word 文档文本内容"""
    try:
        doc = Document(file_path)
        texts = []
        for para in doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        return '\n'.join(texts)
    except Exception as e:
        return f"读取失败：{str(e)}"


def read_txt_text(file_path):
    """读取 TXT 文件内容"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except:
        try:
            with open(file_path, 'r', encoding='gbk') as f:
                return f.read()
        except Exception as e:
            return f"读取失败：{str(e)}"


def get_heading_level(style_name):
    """从样式名获取标题级别"""
    if style_name and 'Heading' in style_name:
        try:
            level_str = style_name.replace('Heading', '').strip()
            level = int(''.join(filter(str.isdigit, level_str)) or '0')
            return level if level > 0 else 1
        except:
            return 1
    return 0


def generate_chapter_numbering(headings):
    """
    根据标题层级生成标准的中文编号
    L1: 第一章、第二章...
    L2: 1.1, 1.2...
    L3: 1.1.1, 1.1.2...
    L4: (1), (2)...
    """
    counters = {1: 0, 2: 0, 3: 0, 4: 0, 5: 0}
    chinese_num = [
        '零', '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
        '十一', '十二', '十三', '十四', '十五', '十六', '十七', '十八', '十九', '二十',
        '二十一', '二十二', '二十三', '二十四', '二十五', '二十六', '二十七', '二十八', '二十九', '三十'
    ]

    for h in headings:
        level = h['level']
        counters[level] += 1
        for l in range(level + 1, 6):
            counters[l] = 0

        if level == 1:
            h['numbering'] = f'第{chinese_num[counters[1]]}章'
        elif level == 2:
            h['numbering'] = f'{counters[1]}.{counters[2]}'
        elif level == 3:
            h['numbering'] = f'{counters[1]}.{counters[2]}.{counters[3]}'
        elif level == 4:
            h['numbering'] = f'({counters[4]})'
        elif level == 5:
            h['numbering'] = f'{counters[1]}.{counters[2]}.{counters[3]}.{counters[4]}'
    return headings


def build_chapter_tree(headings):
    """将扁平的标题列表转换为树形结构"""
    chapters = []
    for h in headings:
        level = h['level']
        node = {
            'number': h['numbering'],
            'title': h['text'],
            'level': level,
            'style': h['style'],
            'children': []
        }
        if level == 1:
            chapters.append(node)
        else:
            parent = find_parent_node(chapters, level - 1)
            if parent:
                parent['children'].append(node)
            else:
                chapters.append(node)
    return chapters


def find_parent_node(chapters, target_level):
    """递归查找指定层级的最后一个节点 - 修复空列表处理"""
    for chapter in reversed(chapters):
        if chapter['level'] == target_level:
            return chapter
        children = chapter.get('children')
        if children is not None and len(children) > 0:
            result = find_parent_node(children, target_level)
            if result:
                return result
    return None


def scan_template_styles(file_path):
    """
    扫描 Word 模板文件，提取章节目录结构
    严格基于 Heading 样式和 XML 编号信息提取（参考 extract_toc.py）
    """
    try:
        doc = Document(file_path)
        headings = []

        # 提取所有 Heading 样式的段落
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ''
            level = get_heading_level(style_name)

            if level > 0:
                headings.append({
                    'level': level,
                    'text': text,
                    'style': style_name,
                    'numbering': ''
                })

        # 生成编号
        headings = generate_chapter_numbering(headings)

        # 构建章节树
        chapters = build_chapter_tree(headings)

        return {
            'success': True,
            'chapters': chapters,
            'message': f'成功扫描 {len(chapters)} 个一级章节',
            'total_nodes': count_all_nodes(chapters)
        }
    except Exception as e:
        import traceback
        return {
            'success': False,
            'chapters': [],
            'message': f'扫描失败：{str(e)}',
            'detail': traceback.format_exc()
        }


def get_paragraph_numbering(para):
    """从段落的 XML 结构中获取编号信息（支持 WPS 多级编号）"""
    try:
        # 获取段落的 XML 元素
        p_element = para._element
        
        # 查找编号属性（w:numPr）
        numPr = p_element.find('.//w:numPr', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        
        if numPr is None:
            return None
        
        # 获取编号级别（w:ilvl）
        ilvl = numPr.find('w:ilvl', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if ilvl is None:
            return None
        
        level = int(ilvl.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')) + 1
        
        # 获取编号 ID（w:numId）
        numId = numPr.find('w:numId', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
        if numId is None:
            return None
        
        num_id = int(numId.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val'))
        
        # 获取编号文本（从段落样式或编号定义中）
        # 尝试从段落文本开头提取编号
        text = para.text.strip()
        
        # 匹配常见编号格式
        import re
        
        # 格式 1: 第一章 项目概况
        match = re.match(r'^第 ([一二三四五六七八九十\d]+) 章\s*(.+)$', text)
        if match:
            return {'number': f'第{match.group(1)}章', 'level': 1}
        
        # 格式 2: 第一节 项目名称
        match = re.match(r'^第 ([一二三四五六七八九十\d]+) 节\s*(.+)$', text)
        if match:
            return {'number': f'第{match.group(1)}节', 'level': 2}
        
        # 格式 3: 一、项目概况
        match = re.match(r'^([一二三四五六七八九十]+)[、.]\s*(.+)$', text)
        if match:
            return {'number': f'{match.group(1)}、', 'level': 1}
        
        # 格式 4: 1 项目概况 或 1.1 项目名称
        match = re.match(r'^(\d+(?:\.\d+)*)\s+(.+)$', text)
        if match:
            num = match.group(1)
            return {'number': num, 'level': len(num.split('.'))}
        
        # 如果无法从文本提取，返回编号级别
        return {'number': str(level), 'level': level}
        
    except Exception as e:
        return None


def add_to_parent_by_level(chapters, node, level):
    """将节点添加到合适的父节点下（按层级）"""
    if not chapters:
        chapters.append(node)
        return
    
    # 从后往前查找最近的层级为 level-1 的节点
    def find_parent(nodes, target_level):
        for n in reversed(nodes):
            if n.get('level', 1) == target_level:
                return n
            children = n.get('children')
            if children is not None and len(children) > 0:
                result = find_parent(children, target_level)
                if result:
                    return result
        return None
    
    parent = find_parent(chapters, level - 1)
    if parent:
        parent['children'].append(node)
    else:
        # 没有找到父节点，添加到最后一个一级章节
        chapters[-1]['children'].append(node)


def add_to_parent(chapters, node, number):
    """将节点添加到正确的父节点下（递归查找）"""
    parts = number.split('.')
    current_level = len(parts)  # 当前节点的层级
    
    if current_level == 1:
        # 一级章节直接添加到 chapters 列表
        chapters.append(node)
    else:
        # 查找父节点（层级为 current_level - 1 的最后一个节点）
        parent_parts = parts[:-1]
        parent_number = '.'.join(parent_parts)
        
        # 在 chapters 中递归查找父节点
        parent = find_node_by_number(chapters, parent_number)
        if parent:
            parent['children'].append(node)


def find_node_by_number(chapters, number):
    """递归查找指定编号的节点 - 修复空列表处理"""
    for chapter in chapters:
        if chapter['number'] == number:
            return chapter
        # 在子节点中递归查找
        children = chapter.get('children')
        if children is not None and len(children) > 0:
            result = find_node_by_number(children, number)
            if result:
                return result
    return None


def count_all_nodes(chapters):
    """递归统计所有节点数量 - 修复空列表处理"""
    count = 0
    for chapter in chapters:
        count += 1
        children = chapter.get('children')
        if children is not None and len(children) > 0:
            count += count_all_nodes(children)
    return count


def extract_paragraph_style(para):
    """提取段落的样式信息"""
    style_info = {
        'font_name': '',
        'font_size': 0,
        'bold': False,
        'italic': False,
        'underline': False,
        'color': '',
        'alignment': '',
        'line_spacing': 0,
        'space_before': 0,
        'space_after': 0,
        'first_line_indent': 0
    }
    
    try:
        # 获取段落格式
        para_format = para.paragraph_format
        
        # 对齐方式
        alignment_map = {
            0: 'left',
            1: 'center',
            2: 'right',
            3: 'justify'
        }
        style_info['alignment'] = alignment_map.get(para_format.alignment, 'left')
        
        # 间距和缩进
        if hasattr(para_format, 'line_spacing') and para_format.line_spacing:
            style_info['line_spacing'] = float(para_format.line_spacing)
        if hasattr(para_format, 'space_before') and para_format.space_before:
            style_info['space_before'] = float(para_format.space_before)
        if hasattr(para_format, 'space_after') and para_format.space_after:
            style_info['space_after'] = float(para_format.space_after)
        if hasattr(para_format, 'first_line_indent') and para_format.first_line_indent:
            style_info['first_line_indent'] = float(para_format.first_line_indent)
        
        # 获取第一个 run 的字体样式（通常段落内样式一致）
        if para.runs:
            run = para.runs[0]
            style_info['font_name'] = run.font.name or ''
            style_info['font_size'] = float(run.font.size) if run.font.size else 0
            style_info['bold'] = run.font.bold or False
            style_info['italic'] = run.font.italic or False
            style_info['underline'] = run.font.underline or False
            
            # 尝试获取中文字体名称
            try:
                from docx.oxml.ns import qn
                rFonts = run._element.rPr.rFonts
                if rFonts is not None and hasattr(rFonts, 'get'):
                    east_asia = rFonts.get(qn('w:eastAsia'))
                    if east_asia:
                        style_info['font_name'] = east_asia
            except:
                pass
    except Exception as e:
        print(f"提取样式失败：{e}")
    
    return style_info


def call_bailian_api(prompt, api_key, model='qwen-max'):
    """调用阿里云百炼平台 API 生成内容（支持多模型和应用 API）"""
    
    # 检测 API Key 类型：sk-sp- 开头的是百炼应用 API Key
    is_app_api = api_key.startswith('sk-sp-') or api_key.startswith('sk-app-')
    
    if is_app_api:
        # 百炼应用 API 调用
        return call_bailian_app_api(prompt, api_key, model)
    else:
        # DashScope 模型 API 调用
        return call_dashscope_api(prompt, api_key, model)


def call_bailian_app_api(prompt, api_key, model='qwen-max'):
    """调用百炼应用 API（使用 sk-sp- 开头的 API Key）"""
    # 百炼应用 API 需要 APP_ID，这里使用默认端点
    # 用户需要在 APP_CONFIGS 中配置 APP_ID，或者在 API Key 中附带 APP_ID
    
    # 检查 API Key 是否包含 APP_ID（格式：sk-sp-xxx#appid）
    app_id = None
    if '#' in api_key:
        api_key, app_id = api_key.split('#', 1)
    else:
        # 尝试从配置中获取默认 APP_ID
        app_id = APP_CONFIGS.get('default', {}).get('app_id')
    
    if not app_id:
        return "错误：百炼应用 API Key 需要配置 APP_ID。请在 APP_CONFIGS 中设置，或在 API Key 后添加 #APP_ID"
    
    url = f"https://dashscope.aliyuncs.com/api/v1/apps/{app_id}/completion"
    
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "input": {
            "prompt": prompt
        },
        "parameters": {
            "max_tokens": 4000
        }
    }
    
    try:
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        if response.status_code == 200:
            result = response.json()
            # 应用 API 响应格式
            if 'output' in result:
                if 'text' in result['output']:
                    return result['output']['text']
                elif 'choices' in result['output']:
                    return result['output']['choices'][0]['message']['content']
            elif 'choices' in result and len(result['choices']) > 0:
                return result['choices'][0]['message']['content']
            else:
                return f"API 返回格式异常：{result}"
        else:
            return f"API 调用失败 (状态码 {response.status_code}): {response.text}"
    except requests.exceptions.Timeout:
        return "API 调用超时，请重试"
    except Exception as e:
        return f"API 调用异常：{str(e)}"


def call_dashscope_api(prompt, api_key, model='qwen-max'):
    """调用 DashScope 模型 API（使用 sk- 开头的 API Key）"""
    
    # 不同模型的 API 端点
    api_endpoints = {
        'qwen-max': 'https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation',
        'qwen-plus': 'https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation',
        'qwen-turbo': 'https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation',
        'glm-4': 'https://open.bigmodel.cn/api/paas/v4/chat/completions',
        'glm-3-turbo': 'https://open.bigmodel.cn/api/paas/v4/chat/completions',
        'kimi-latest': 'https://api.moonshot.cn/v1/chat/completions',
        'mini-max': 'https://api.minimax.chat/v1/text/chatcompletion_v2',
    }
    
    url = api_endpoints.get(model, api_endpoints['qwen-max'])
    
    headers = {
        "Content-Type": "application/json"
    }
    
    # 根据模型构建不同的请求格式
    if model.startswith('qwen'):
        # 通义千问系列（DashScope 格式）
        headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "input": {
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            },
            "parameters": {
                "max_tokens": 4000,
                "temperature": 0.7,
                "top_p": 0.8
            }
        }
    elif model.startswith('glm'):
        # GLM 系列（智谱 AI / 百炼平台格式）
        headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 4000,
            "temperature": 0.7,
            "top_p": 0.8
        }
    elif model.startswith('kimi'):
        # Kimi（月之暗面格式）
        headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 4000,
            "temperature": 0.7,
            "top_p": 0.8
        }
    elif model.startswith('mini'):
        # MiniMax 格式
        headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "messages": [
                {"role": "user", "content": prompt}
            ],
            "max_tokens": 4000,
            "temperature": 0.7,
            "top_p": 0.8
        }
    else:
        # 默认使用通义千问格式
        headers["Authorization"] = f"Bearer {api_key}"
        payload = {
            "model": model,
            "input": {
                "messages": [
                    {"role": "user", "content": prompt}
                ]
            },
            "parameters": {
                "max_tokens": 4000,
                "temperature": 0.7,
                "top_p": 0.8
            }
        }

    try:
        response = requests.post(url, headers=headers, json=payload, timeout=120)
        if response.status_code == 200:
            result = response.json()
            # 不同模型的响应格式解析
            if 'output' in result and 'text' in result['output']:
                # 通义千问格式
                return result['output']['text']
            elif 'output' in result and 'choices' in result['output']:
                # 通义千问新格式
                return result['output']['choices'][0]['message']['content']
            elif 'choices' in result and len(result['choices']) > 0:
                # OpenAI 兼容格式（GLM、Kimi、MiniMax）
                return result['choices'][0]['message']['content']
            else:
                return f"API 返回格式异常：{result}"
        else:
            return f"API 调用失败 (状态码 {response.status_code}): {response.text}"
    except requests.exceptions.Timeout:
        return "API 调用超时，请重试"
    except Exception as e:
        return f"API 调用异常：{str(e)}"


def call_qwen_api(prompt, api_key, model='qwen-max'):
    """调用 AI API 生成内容（兼容旧接口）"""
    return call_bailian_api(prompt, api_key, model)


def generate_content_with_ai(requirement_content, template_content, user_prompt, section_title, api_key, model='qwen-max', progress_callback=None):
    """使用 AI 生成指定小节的内容（支持章节分类处理）"""
    
    # 判断章节类型
    if is_info_chapter(section_title):
        # 信息提取类章节：直接从需求文档中提取
        extracted_info = extract_info_from_requirement(section_title, requirement_content)
        if extracted_info:
            return extracted_info
        # 如果提取失败，降级为 AI 生成
    
    # 内容撰写类章节（或信息提取失败）：调用 AI 生成
    # 从模板中提取参考示例
    template_example = extract_template_example(section_title, template_content)
    
    # 构建精准 Prompt
    if template_example:
        prompt = f"""你是一位专业的可行性研究报告撰写专家。请根据以下材料，为"{section_title}"这一小节撰写专业、详细的正文内容。

【模板中的参考示例】（请参考此格式和风格）
{template_example}

【需求文档内容】
{requirement_content[:5000] if requirement_content else '暂无需求文档'}

【用户额外要求】
{user_prompt if user_prompt else '无特殊要求'}

【重要撰写要求】
1. 内容要专业、严谨，符合可行性研究报告的规范
2. 结合需求文档中的具体信息，不要泛泛而谈
3. 数据要合理，逻辑要清晰
4. 使用正式的公文语言
5. 参考模板示例的格式和风格，不要发挥
6. 不要包含 Markdown 格式（如 #、** 等），使用纯文本
7. 【重要】不要输出任何标题，只输出正文内容！标题已经由系统自动生成
8. 【重要】不要以"{section_title}"或类似标题开头，直接开始正文内容

请为"{section_title}"这一小节撰写正文内容（只输出正文，不要输出标题）："""
    else:
        # 没有模板示例时的降级 Prompt
        prompt = f"""你是一位专业的可行性研究报告撰写专家。请根据以下材料，为"{section_title}"这一小节撰写专业、详细的正文内容。

【需求文档内容】
{requirement_content[:5000] if requirement_content else '暂无需求文档'}

【模板格式参考】
{template_content[:3000] if template_content else '暂无模板'}

【用户额外要求】
{user_prompt if user_prompt else '无特殊要求'}

【重要撰写要求】
1. 内容要专业、严谨，符合可行性研究报告的规范
2. 结合需求文档中的具体信息，不要泛泛而谈
3. 数据要合理，逻辑要清晰
4. 使用正式的公文语言
5. 字数控制在 800-1500 字之间
6. 不要包含 Markdown 格式（如 #、** 等），使用纯文本
7. 【重要】不要输出任何标题，只输出正文内容！标题已经由系统自动生成
8. 【重要】不要以"{section_title}"或类似标题开头，直接开始正文内容

请为"{section_title}"这一小节撰写正文内容（只输出正文，不要输出标题）："""

    return call_bailian_api(prompt, api_key, model)


def load_style_config():
    """加载样式配置"""
    styles_path = os.path.join(app.config['UPLOAD_FOLDER'], 'style_config.json')
    default_styles = {
        'heading1': {
            'font_name': '黑体',
            'font_size': 22,
            'bold': True,
            'alignment': 'center'
        },
        'heading2': {
            'font_name': '楷体',
            'font_size': 16,
            'bold': True,
            'alignment': 'left'
        },
        'heading3': {
            'font_name': '楷体',
            'font_size': 14,
            'bold': True,
            'alignment': 'left'
        },
        'normal': {
            'font_name': '仿宋',
            'font_size': 10.5,
            'line_spacing': 1.5,
            'first_line_indent': 0.74
        }
    }
    
    if os.path.exists(styles_path):
        try:
            with open(styles_path, 'r', encoding='utf-8') as f:
                styles = json.load(f)
            # 合并用户配置和默认配置
            for key in default_styles:
                if key in styles:
                    default_styles[key].update(styles[key])
        except:
            pass
    
    return default_styles


def add_heading(doc, text, level=1, styles=None):
    """添加标题 - 支持自定义样式配置"""
    if styles is None:
        styles = load_style_config()
    
    if level == 1:
        style = styles.get('heading1', {})
        heading = doc.add_heading(text, level=1)
        
        # 对齐方式
        align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT, 
                     'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
        heading.alignment = align_map.get(style.get('alignment', 'center'), WD_ALIGN_PARAGRAPH.CENTER)
        
        # 设置字体样式
        for run in heading.runs:
            run.font.name = style.get('font_name', '黑体')
            run.font.size = Pt(style.get('font_size', 22))
            run.font.bold = style.get('bold', True)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style.get('font_name', '黑体'))
        
        heading.paragraph_format.space_before = Cm(0)
        heading.paragraph_format.space_after = Cm(0.5)
        
    elif level == 2:
        style = styles.get('heading2', {})
        heading = doc.add_heading(text, level=2)
        
        align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                     'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
        heading.alignment = align_map.get(style.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
        
        for run in heading.runs:
            run.font.name = style.get('font_name', '楷体')
            run.font.size = Pt(style.get('font_size', 16))
            run.font.bold = style.get('bold', True)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style.get('font_name', '楷体'))
        
        heading.paragraph_format.space_before = Cm(0.5)
        heading.paragraph_format.space_after = Cm(0.5)
        
    elif level == 3:
        style = styles.get('heading3', {})
        heading = doc.add_heading(text, level=3)
        
        align_map = {'center': WD_ALIGN_PARAGRAPH.CENTER, 'left': WD_ALIGN_PARAGRAPH.LEFT,
                     'right': WD_ALIGN_PARAGRAPH.RIGHT, 'justify': WD_ALIGN_PARAGRAPH.JUSTIFY}
        heading.alignment = align_map.get(style.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
        
        for run in heading.runs:
            run.font.name = style.get('font_name', '楷体')
            run.font.size = Pt(style.get('font_size', 14))
            run.font.bold = style.get('bold', True)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style.get('font_name', '楷体'))
        
        heading.paragraph_format.space_before = Cm(0.3)
        heading.paragraph_format.space_after = Cm(0.3)
        
    elif level == 4:
        style = styles.get('heading3', {})
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.font.name = style.get('font_name', '楷体')
        run.font.size = Pt(style.get('font_size', 14))
        run.font.bold = style.get('bold', True)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), style.get('font_name', '楷体'))
        heading.paragraph_format.first_line_indent = Cm(0.74)
        heading.paragraph_format.space_before = Cm(0)
        heading.paragraph_format.space_after = Cm(0)
        
    return heading


def add_normal_paragraph(doc, text, indent=True, styles=None):
    """添加正文段落 - 支持自定义样式配置"""
    if styles is None:
        styles = load_style_config()
    
    style = styles.get('normal', {})
    
    para = doc.add_paragraph()
    if not text.strip():
        return para
    
    run = para.add_run(text)
    run.font.name = style.get('font_name', '仿宋')
    run.font.size = Pt(style.get('font_size', 10.5))
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style.get('font_name', '仿宋'))
    
    if indent and style.get('first_line_indent'):
        para.paragraph_format.first_line_indent = Cm(style.get('first_line_indent', 0.74))
    
    if style.get('line_spacing'):
        para.paragraph_format.line_spacing = style.get('line_spacing', 1.5)
    
    para.paragraph_format.space_before = Cm(0)
    para.paragraph_format.space_after = Cm(0)
    return para


def add_toc(doc, styles=None):
    """添加自动目录（TOC）- 使用 Word 域代码"""
    if styles is None:
        styles = load_style_config()
    
    # 添加目录标题
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    run.font.name = '黑体'
    run.font.size = Pt(22)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 添加空行
    doc.add_paragraph()
    
    # 添加 TOC 域
    # 使用 Word 的 TOC 域代码：\o "1-3" 表示包含 1-3 级标题
    paragraph = doc.add_paragraph()
    
    # 创建 TOC 域的开始标记
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar)
    
    # 添加域代码指令
    run = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'  # 1-3级标题，带超链接，隐藏页码（临时）
    run._r.append(instrText)
    
    # 创建 TOC 域的结束标记
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar)
    
    # 添加分页符
    doc.add_page_break()


def generate_word_document(template_type, requirement_content, template_content, user_prompt, api_key, output_path, progress_callback=None, model='qwen-max'):
    """生成 Word 文档（AI 增强版）- 使用用户配置的目录结构"""

    doc = Document()
    styles = load_style_config()
    
    style = doc.styles['Normal']
    style.font.name = styles['normal'].get('font_name', '仿宋')
    style.font.size = Pt(styles['normal'].get('font_size', 10.5))
    style._element.rPr.rFonts.set(qn('w:eastAsia'), styles['normal'].get('font_name', '仿宋'))

    # 尝试加载用户保存的章节目录
    chapters_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chapter_config.json')
    user_chapters = None
    if os.path.exists(chapters_path):
        try:
            with open(chapters_path, 'r', encoding='utf-8') as f:
                user_chapters = json.load(f)
        except:
            user_chapters = None
    
    # 如果没有用户目录，使用硬编码的 TEMPLATE_TYPES
    if not user_chapters:
        template_config = TEMPLATE_TYPES.get(template_type, TEMPLATE_TYPES['future_community'])
        # 转换为统一格式
        user_chapters = []
        for chapter_title, sections in template_config['chapters']:
            chapter_node = {
                'number': chapter_title.split()[0] if ' ' in chapter_title else '1',
                'title': chapter_title,
                'level': 1,
                'children': []
            }
            for idx, section in enumerate(sections):
                section_num = section.split()[0] if ' ' in section else f'{chapter_node["number"]}.{idx+1}'
                chapter_node['children'].append({
                    'number': section_num,
                    'title': section,
                    'level': 2,
                    'children': []
                })
    
    project_name = "良熟新苑未来社区建设项目" if template_type == 'future_community' else "建设项目"

    # ========== 封面 ==========
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Cm(5)
    title.paragraph_format.space_after = Cm(1)
    run = title.add_run(project_name)
    run.font.name = '黑体'
    run.font.size = Pt(36)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_before = Cm(1)
    subtitle.paragraph_format.space_after = Cm(3)
    run = subtitle.add_run(template_config['name'] if 'template_config' in dir() else '可行性研究报告')
    run.font.name = '黑体'
    run.font.size = Pt(26)
    run.font.bold = True
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

    for _ in range(5):
        doc.add_paragraph()

    builder = doc.add_paragraph()
    builder.alignment = WD_ALIGN_PARAGRAPH.CENTER
    builder.paragraph_format.space_before = Cm(0.5)
    builder.paragraph_format.space_after = Cm(0.5)
    run = builder.add_run('建设单位：良熟社区居委会')
    run.font.name = '楷体'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    compiler = doc.add_paragraph()
    compiler.alignment = WD_ALIGN_PARAGRAPH.CENTER
    compiler.paragraph_format.space_before = Cm(0.5)
    compiler.paragraph_format.space_after = Cm(0.5)
    run = compiler.add_run('编制单位：XX 数字科技有限公司')
    run.font.name = '楷体'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para.paragraph_format.space_before = Cm(0.5)
    date_para.paragraph_format.space_after = Cm(0.5)
    run = date_para.add_run(f'编制日期：{datetime.now().strftime("%Y 年 %m 月")}')
    run.font.name = '楷体'
    run.font.size = Pt(16)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')

    doc.add_page_break()
    if progress_callback:
        progress_callback(5, "封面完成")

    # ========== 项目编审人员名单 ==========
    add_heading(doc, '项目编审人员名单', level=1, styles=styles)
    doc.add_paragraph()
    add_normal_paragraph(doc, '项目负责：张建国', styles=styles)
    add_normal_paragraph(doc, '编制小组：李明、王芳、陈伟、刘洋', styles=styles)
    add_normal_paragraph(doc, '勘误核稿：赵强', styles=styles)
    add_normal_paragraph(doc, '项目审定：周建华', styles=styles)
    doc.add_page_break()
    if progress_callback:
        progress_callback(10, "编审名单完成")

    # ========== 目录 ==========
    add_heading(doc, '目录', level=1, styles=styles)
    doc.add_paragraph()
    
    # 递归渲染目录
    render_doc_toc(doc, user_chapters, styles)

    doc.add_page_break()
    if progress_callback:
        progress_callback(15, "目录完成")

    # ========== 正文章节（使用用户配置的目录） ==========
    total_chapters = len(user_chapters)

    def render_chapter_content(node, level=1):
        """递归渲染章节内容"""
        node_number = node.get('number', '')
        node_title = node.get('title', '')
        node_level = node.get('level', level)
        children = node.get('children')
        
        # 添加章节标题
        add_heading(doc, f"{node_number} {node_title}", level=node_level, styles=styles)
        
        # 判断是否有子节点
        if children is not None and len(children) > 0:
            # 有子节点，递归处理
            for child in children:
                render_chapter_content(child, level=node_level + 1)
        else:
            # 叶子节点，生成内容
            content = get_chapter_content_template(node_title)
            for para_text in content.split('\n'):
                if para_text.strip():
                    add_normal_paragraph(doc, para_text.strip(), styles=styles)

    for idx, chapter in enumerate(user_chapters):
        # 渲染章节内容（递归处理所有层级）
        render_chapter_content(chapter, level=1)
        
        doc.add_page_break()

        # 进度回调
        progress = 15 + int((idx + 1) / total_chapters * 80)
        if progress_callback:
            progress_callback(progress, f"正在生成第{idx + 1}章：{chapter['title']}")

    # 保存文档
    doc.save(output_path)
    if progress_callback:
        progress_callback(100, "文档保存成功")

    return True


def render_doc_toc(doc, chapters, styles, level=0):
    """递归渲染目录 - 正确处理空列表和 null"""
    for chapter in chapters:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(0.74 * level)
        run = para.add_run(f"{chapter['number']} {chapter['title']}")
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 修复：检查 children 是否存在且非空
        children = chapter.get('children')
        if children is not None and len(children) > 0:
            render_doc_toc(doc, children, styles, level + 1)

def generate_chapter_content(doc, chapter_node, requirement_content, template_content,
                              user_prompt, model_config, styles, depth=0):
    """
    递归生成章节内容
    严格按照目录树结构生成标题和内容

    优化：支持 AI 生成
    - 有模型配置：调用 AI 生成内容
    - 无模型配置：使用模板内容
    """
    from ai_engine import generate_section_content_with_ai
    
    children = chapter_node.get('children')

    # 判断是否有子节点
    if children is not None and len(children) > 0:
        # 有子节点，递归处理
        for child in children:
            child_level = child.get('level', 2)
            # 根据层级添加标题
            add_heading(doc, f"{child['number']} {child['title']}", level=child_level, styles=styles)

            # 递归处理子节点
            generate_chapter_content(doc, child, requirement_content, template_content,
                                      user_prompt, model_config, styles, depth + 1)
    else:
        # 叶子节点，生成内容
        section_title = chapter_node['title']

        # 判断是否使用 AI 生成
        # Ollama 本地模型不需要 API Key
        has_api_key = model_config and (model_config.api_key or model_config.provider_id == 'ollama')
        if has_api_key:
            # 调用 AI 生成
            content = generate_section_content_with_ai(
                section_title=section_title,
                requirement_text=requirement_content,
                template_text=template_content,
                user_instruction=user_prompt,
                model_config=model_config,
                fallback_to_template=False  # AI 失败时返回错误，不降级
            )
            # AI 返回后检查任务是否被取消
            task_info = task_manager.load_task_info(task_id)
            if task_info and task_info.status == 'cancelled':
                print(f'[CANCEL] 任务已取消，停止生成')
                return
            # 检查是否生成失败
            if content.startswith('[AI 生成失败]'):
                print(f'[ERROR] 章节生成失败：{section_title}')
                # 添加错误提示到文档
                add_normal_paragraph(doc, f"【生成错误】{content}", styles=styles)
                return
        else:
            # 使用模板内容
            content = get_chapter_content_template(section_title)

        # 添加正文内容
        for para_text in content.split('\n'):
            if para_text.strip():
                add_normal_paragraph(doc, para_text.strip(), styles=styles)


def generate_chapter(doc, chapter_title, sections, requirement_content, template_content,
                     user_prompt, model_config, section_title_prefix=''):
    """生成单个章节的内容（使用 AI 生成）"""
    from ai_engine import generate_section_content_with_ai

    add_heading(doc, chapter_title, level=1)

    for section_title in sections:
        add_heading(doc, section_title, level=2)

        # 使用 AI 生成内容
        # Ollama 本地模型不需要 API Key
        has_api_key = model_config and (model_config.api_key or model_config.provider_id == 'ollama')
        if has_api_key:
            # AI 调用前检查取消状态
            task_info = task_manager.load_task_info(task_id)
            if task_info and task_info.status == 'cancelled':
                print(f'[CANCEL] 任务已取消，停止生成')
                return
                
            content = generate_section_content_with_ai(
                section_title=section_title,
                requirement_text=requirement_content,
                template_text=template_content,
                user_instruction=user_prompt,
                model_config=model_config,
                fallback_to_template=False  # AI 失败时返回错误，不降级
            )
            # AI 返回后检查取消状态
            task_info = task_manager.load_task_info(task_id)
            if task_info and task_info.status == 'cancelled':
                print(f'[CANCEL] 任务已取消，停止生成')
                return
            # 检查是否生成失败
            if content.startswith('[AI 生成失败]'):
                print(f'[ERROR] 章节生成失败：{section_title}')
                # 添加错误提示到文档
                add_normal_paragraph(doc, f"【生成错误】{content}")
                continue
        else:
            # 没有模型配置，使用模板
            content = get_chapter_content_template(section_title)

        for para_text in content.split('\n'):
            if para_text.strip():
                add_normal_paragraph(doc, para_text.strip())

    doc.add_page_break()
    return doc


def process_document_async(task_id, template_type, requirement_content, template_content,
                           user_prompt, model_config, output_path):
    """异步处理文档生成任务 - 分章节生成，每章等待确认（集成 AI 内容优化）
    
    注意：此函数已弃用，请使用 process_document_async_v2
    """

    # 全局变量用于持续追踪文档
    doc_container = {'doc': None}

    # 初始化优化服务
    from ai_engine import reset_optimization_services, get_data_point_manager, get_requirement_analyzer
    reset_optimization_services()
    dp_manager = get_data_point_manager()
    req_analyzer = get_requirement_analyzer()
    quality_reviewer = QualityReviewer()

    # 记录所有章节内容（用于质量审校）
    chapter_contents = {}

    def progress_callback(progress, message):
        """进度回调函数"""
        task_manager.update_task_progress(task_id, progress=progress, message=message)

    def save_partial_document():
        """保存部分生成的文档"""
        if doc_container['doc']:
            # 保存到任务目录下，使用完整 task_id 作为文件名
            task_dir = task_manager._get_task_directory(task_id)
            partial_path = os.path.join(task_dir, f'partial_{task_id}.docx')
            doc_container['doc'].save(partial_path)
            print(f'[INFO] 临时文档已保存：{partial_path}')

    try:
        # 标记任务开始
        task_manager.set_task_started(task_id)

        # 更新状态：解析文件
        task_manager.update_task_progress(task_id, progress=5,
            status=TaskStatus.PARSING_FILE.value, message='解析上传文件中...')
        
        # 从需求文档提取初始数据点和需求点
        if requirement_content:
            task_manager.update_task_progress(task_id, progress=7,
                message='分析需求文档，提取关键信息...')
            # 提取数据点
            initial_data = dp_manager.extract_from_text(requirement_content, "需求文档")
            if initial_data:
                dp_manager.update(initial_data, "需求文档")
                print(f'[INFO] 从需求文档提取数据点: {list(initial_data.keys())}')
            # 提取需求点
            req_analyzer.extract(requirement_content)
            print(f'[INFO] 从需求文档提取需求点完成')

        # 更新状态：AI 生成
        task_manager.update_task_progress(task_id, progress=10,
            status=TaskStatus.GENERATING_AI.value, message='准备生成文档...')
        
        # 获取模板配置
        template_config = TEMPLATE_TYPES.get(template_type, TEMPLATE_TYPES['future_community'])
        total_chapters = len(template_config['chapters'])
        
        # 初始化 Word 文档（封面、编审名单、目录）
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = '仿宋'
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        
        project_name = "良熟新苑未来社区建设项目" if template_type == 'future_community' else "建设项目"
        
        # 封面
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Cm(5)
        title.paragraph_format.space_after = Cm(1)
        run = title.add_run(project_name)
        run.font.name = '黑体'
        run.font.size = Pt(36)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Cm(1)
        subtitle.paragraph_format.space_after = Cm(3)
        run = subtitle.add_run(template_config['name'])
        run.font.name = '黑体'
        run.font.size = Pt(26)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        for _ in range(5):
            doc.add_paragraph()
        
        builder = doc.add_paragraph()
        builder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        builder.paragraph_format.space_before = Cm(0.5)
        builder.paragraph_format.space_after = Cm(0.5)
        run = builder.add_run('建设单位：良熟社区居委会')
        run.font.name = '楷体'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        compiler = doc.add_paragraph()
        compiler.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compiler.paragraph_format.space_before = Cm(0.5)
        compiler.paragraph_format.space_after = Cm(0.5)
        run = compiler.add_run('编制单位：XX 数字科技有限公司')
        run.font.name = '楷体'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Cm(0.5)
        date_para.paragraph_format.space_after = Cm(0.5)
        run = date_para.add_run(f'编制日期：{datetime.now().strftime("%Y 年 %m 月")}')
        run.font.name = '楷体'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        doc.add_page_break()
        
        # 编审名单
        add_heading(doc, '项目编审人员名单', level=1)
        doc.add_paragraph()
        add_normal_paragraph(doc, '项目负责：张建国')
        add_normal_paragraph(doc, '编制小组：李明、王芳、陈伟、刘洋')
        add_normal_paragraph(doc, '勘误核稿：赵强')
        add_normal_paragraph(doc, '项目审定：周建华')
        doc.add_page_break()
        
        # 目录
        add_heading(doc, '目录', level=1)
        doc.add_paragraph()
        for chapter_title, sections in template_config['chapters']:
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0)
            run = para.add_run(chapter_title)
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            for section in sections:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(0.74)
                run = para.add_run(section)
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        doc.add_page_break()
        
        # 保存初始文档（含封面、名单、目录）
        doc_container['doc'] = doc
        save_partial_document()
        
        task_manager.update_task_progress(task_id, progress=15,
            status=TaskStatus.CREATING_DOC.value, message='目录完成，准备生成正文')
        
        # 逐章生��
        for idx, (chapter_title, sections) in enumerate(template_config['chapters']):
            # 检查任务是否被取消
            task_info = task_manager.get_task_status(task_id)
            if task_info and task_info.get('status') == TaskStatus.CANCELLED.value:
                return
            
            # 记录章节步骤
            task_manager.add_chapter_step(task_id, idx, chapter_title, status='pending')
            
            # 更新当前章节
            task_manager.set_current_chapter(task_id, chapter_title)
            task_manager.update_chapter_step(task_id, idx, status='generating', 
                message=f'正在生成第{idx + 1}章：{chapter_title}')
            
            task_manager.update_task_progress(task_id, 
                progress=15 + int((idx + 1) / total_chapters * 80),
                status=TaskStatus.GENERATING_AI.value,
                message=f'正在生成第{idx + 1}章：{chapter_title}')
            
            # 生成当前章节
            doc = generate_chapter(
                doc, chapter_title, sections,
                requirement_content, template_content, user_prompt,
                model_config
            )
            
            # 标记章节完成
            task_manager.add_completed_chapter(task_id, chapter_title)
            task_manager.update_chapter_step(task_id, idx, status='completed',
                message=f'第{idx + 1}章：{chapter_title} 生成完成')
            
            # 保存部分文档
            save_partial_document()
            
            # 如果不是最后一章，等待用户确认
            if idx < total_chapters - 1:
                task_manager.set_pending_confirmation(task_id, True)
                task_manager.update_task_progress(task_id,
                    progress=15 + int((idx + 1) / total_chapters * 80),
                    status=TaskStatus.CREATING_DOC.value,
                    message=f'第{idx + 1}章已完成，请预览并确认后继续生成下一章')
                
                # 等待用户确认（轮询检查）
                wait_start = time.time()
                while task_manager.get_task_status(task_id).get('pending_confirmation'):
                    time.sleep(1)
                    # 超时检查（5 分钟）
                    if time.time() - wait_start > 300:
                        task_manager.update_task_progress(task_id,
                            message='等待确认超时，继续生成下一章')
                        break
                    
                    # 检查是否被取消
                    task_info = task_manager.get_task_status(task_id)
                    if task_info and task_info.get('status') == TaskStatus.CANCELLED.value:
                        return
        
        # 所有章节生成完成
        doc.save(output_path)
        task_manager.set_partial_filename(task_id, None)
        
        # 质量审校
        task_manager.update_task_progress(task_id, progress=98,
            message='正在进行质量审校...')
        try:
            # 获取所有数据点和需求点
            data_points = dp_manager.get_all()
            requirements = req_analyzer.get_all_requirements()
            
            # 生成审校报告
            review_report = quality_reviewer.generate_report(
                chapter_contents=chapter_contents,
                data_points=data_points,
                requirements=requirements
            )
            
            # 保存审校报告
            report_filename = f'review_report_{task_id[:8]}.md'
            report_path = os.path.join(app.config['OUTPUT_FOLDER'], report_filename)
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(review_report.to_markdown())
            
            print(f'[INFO] 质量审校完成：{report_path}')
            print(f'[INFO] 综合评分：{review_report.overall_score:.1f}/100')
            print(f'[INFO] 数据一致性：{review_report.data_consistency_rate:.1%}')
            print(f'[INFO] 需求覆盖率：{review_report.requirement_coverage_rate:.1%}')
            
            # 如果有冲突或未覆盖需求，输出警告
            if review_report.consistency_issues:
                print(f'[WARN] 发现 {len(review_report.consistency_issues)} 个数据一致性问题')
            if review_report.coverage_issues:
                print(f'[WARN] 发现 {len(review_report.coverage_issues)} 个未覆盖需求')
                
        except Exception as e:
            print(f'[WARN] 质量审校失败：{e}')
        
        task_manager.mark_task_completed(task_id, output_filename=os.path.basename(output_path))
        print(f'[INFO] 文档生成完成：{output_path}')

    except Exception as e:
        import traceback
        print(f'[ERROR] 文档生成失败：{e}')
        task_manager.mark_task_failed(task_id, f'{str(e)}\n{traceback.format_exc()}')


# ==================== 修复后的 process_document_async 函数 ==========
# 使用保存的目录配置生成文档，而不是硬编码的 TEMPLATE_TYPES

def process_document_async_v2(task_id, template_type, requirement_content, template_content,
                           user_prompt, model_config, output_path):
    """异步处理文档生成任务 - 使用保存的目录配置（集成 AI 内容优化）
    
    Args:
        task_id: 任务ID
        template_type: 模板类型
        requirement_content: 需求文档内容
        template_content: 模板文档内容
        user_prompt: 用户补充要求
        model_config: 模型配置对象
        output_path: 输出文件路径
    """
    from task_manager import get_task_manager

    task_manager = get_task_manager()
    doc_container = {'doc': None}

    # 初始化优化服务
    from ai_engine import reset_optimization_services, get_data_point_manager, get_requirement_analyzer
    reset_optimization_services()
    dp_manager = get_data_point_manager()
    req_analyzer = get_requirement_analyzer()
    quality_reviewer = QualityReviewer()

    def save_partial_document():
        if doc_container['doc']:
            # 保存到任务目录下，使用完整 task_id 作为文件名
            task_dir = task_manager._get_task_directory(task_id)
            partial_path = os.path.join(task_dir, f'partial_{task_id}.docx')
            doc_container['doc'].save(partial_path)
            print(f'[INFO] 临时文档已保存：{partial_path}')

    try:
        task_manager.set_task_started(task_id)
        task_manager.update_task_progress(task_id, progress=5,
            status=TaskStatus.PARSING_FILE.value, message='解析上传文件中...')
        
        # 从需求文档提取初始数据点和需求点
        if requirement_content:
            task_manager.update_task_progress(task_id, progress=7,
                message='分析需求文档，提取关键信息...')
            # 提取数据点
            initial_data = dp_manager.extract_from_text(requirement_content, "需求文档")
            if initial_data:
                dp_manager.update(initial_data, "需求文档")
                print(f'[INFO] 从需求文档提取数据点: {list(initial_data.keys())}')
            # 提取需求点
            req_analyzer.extract(requirement_content)
            print(f'[INFO] 从需求文档提取需求点完成')
        
        task_manager.update_task_progress(task_id, progress=10,
            status=TaskStatus.GENERATING_AI.value, message='准备生成文档...')

        # 加载保存的目录配置
        chapters_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chapter_config.json')
        user_chapters = None
        if os.path.exists(chapters_path):
            try:
                with open(chapters_path, 'r', encoding='utf-8') as f:
                    user_chapters = json.load(f)
                print(f'[INFO] 已加载用户目录配置：{len(user_chapters)} 个一级章节')
            except Exception as e:
                print(f'[ERROR] 加载目录配置失败：{e}')
        
        if not user_chapters:
            template_config = TEMPLATE_TYPES.get(template_type, TEMPLATE_TYPES['future_community'])
            user_chapters = []
            for chapter_title, sections in template_config['chapters']:
                chapter_node = {
                    'number': chapter_title.split()[0] if ' ' in chapter_title else '1',
                    'title': chapter_title,
                    'level': 1,
                    'children': []
                }
                for idx, section in enumerate(sections):
                    section_num = section.split()[0] if ' ' in section else f'{chapter_node["number"]}.{idx+1}'
                    chapter_node['children'].append({
                        'number': section_num,
                        'title': section,
                        'level': 2,
                        'children': []
                    })
            print('[INFO] 使用默认模板配置')

        # 初始化章节列表（使用树形结构）
        task_manager.initialize_chapters_with_tree(task_id, user_chapters)
        print(f'[TASK] 初始化章节列表（树形）：{len(user_chapters)} 个一级章节')

        styles = load_style_config()
        doc = Document()
        normal_style = styles.get('normal', {})
        style = doc.styles['Normal']
        style.font.name = normal_style.get('font_name', '仿宋')
        style.font.size = Pt(normal_style.get('font_size', 10.5))
        style._element.rPr.rFonts.set(qn('w:eastAsia'), normal_style.get('font_name', '仿宋'))

        # 获取项目名称
        project_name = "建设项目"
        if dp_manager.get('项目名称'):
            project_name = dp_manager.get('项目名称')
        print(f'[INFO] 项目名称：{project_name}')

        # 封面
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Cm(5)
        run = title.add_run(project_name)
        run.font.name = '黑体'
        run.font.size = Pt(36)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        doc.add_page_break()

        # 编审名单
        add_heading(doc, '项目编审人员名单', level=1, styles=styles)
        doc.add_page_break()

        # 目录 - 使用自动目录（TOC 域）
        add_toc(doc, styles=styles)

        doc_container['doc'] = doc
        save_partial_document()
        task_manager.update_task_progress(task_id, progress=15, message='目录完成')

        # 正文 - 生成所有章节
        # 计算所有章节标题（包括子章节）
        all_chapter_titles = []
        def collect_chapter_titles(nodes):
            for node in nodes:
                all_chapter_titles.append(node.get('title', ''))
                children = node.get('children', [])
                if children:
                    collect_chapter_titles(children)
        collect_chapter_titles(user_chapters)

        total_chapters = len(all_chapter_titles)

        def render_chapter_with_status(node, level=1):
            """递归渲染章节，同时更新章节状态"""
            node_title = node.get('title', '')
            children = node.get('children', [])
            # 使用节点自身的 index，而不是递增的 chapter_index
            node_index = node.get('index', 0)

            # 每次递归前检查任务是否被取消或暂停
            task_info = task_manager.load_task_info(task_id)
            if task_info and task_info.status == 'cancelled':
                print(f'[CANCEL] 任务已取消，停止生成')
                return
            
            # 检查是否暂停
            if task_info and task_info.status == 'paused':
                print(f'[PAUSE] 任务已暂停，等待继续...')
                while True:
                    time.sleep(2)
                    task_info = task_manager.load_task_info(task_id)
                    if task_info.status == 'cancelled':
                        print(f'[CANCEL] 任务已取消，停止生成')
                        return
                    if task_info.status != 'paused':
                        print(f'[RESUME] 任务已恢复')
                        break

            if children:
                # 有子节点，递归处理
                for child in children:
                    # 每次递归前检查取消状态
                    task_info = task_manager.load_task_info(task_id)
                    if task_info and task_info.status == 'cancelled':
                        return
                    render_chapter_with_status(child, level + 1)
            else:
                # 叶子节点，检查是否已完成（避免覆盖用户编辑）
                chapters = task_manager.load_chapters(task_id)
                if node_index < len(chapters):
                    existing_chapter = chapters[node_index]
                    if existing_chapter.status == 'completed' and existing_chapter.content:
                        # 已经有内容且完成，跳过重新生成，直接使用已有内容
                        print(f'[INFO] 章节 {node_index} "{node_title}" 已完成，跳过重新生成')
                        add_heading(doc, f"{node.get('number', '')} {node_title}", level=level, styles=styles)
                        for para_text in existing_chapter.content.split('\n'):
                            if para_text.strip():
                                add_normal_paragraph(doc, para_text.strip(), styles=styles)
                        return
                
                # 叶子节点，更新状态并生成内容
                # 生成前检查取消状态
                task_info = task_manager.load_task_info(task_id)
                if task_info and task_info.status == 'cancelled':
                    print(f'[CANCEL] 任务已取消，停止生成')
                    return

                task_manager.update_chapter_status(task_id, node_index, status='generating')

                add_heading(doc, f"{node.get('number', '')} {node_title}", level=level, styles=styles)

                # 生成内容
                # Ollama 本地模型不需要 API Key
                has_api_key = model_config and (model_config.api_key or model_config.provider_id == 'ollama')
                if has_api_key:
                    # 判断是否为信息提取类章节
                    if is_info_chapter(node_title):
                        # 从数据点管理器中获取信息（保证与封面一致）
                        data_point_key = node_title
                        # 去除编号前缀，如"1.1 项目名称" -> "项目名称"
                        for key in ['项目名称', '项目建设单位', '负责人', '联系方式',
                                    '建设工期', '总投资', '资金来源', '编制单位']:
                            if key in node_title:
                                data_point_key = key
                                break
                        content = dp_manager.get(data_point_key)
                        if not content:
                            # 如果数据点管理器中没有，尝试从需求文档中提取
                            from ai_engine import extract_info_field
                            content = extract_info_field(
                                section_title=node_title,
                                requirement_text=requirement_content,
                                ai_call_func=None,
                                model_config=None
                            )
                        if not content:
                            content = f"【未找到相关信息】请在需求文档中补充{node_title}的具体信息"
                        print(f'[INFO] 信息提取：{node_title} = {content}')
                    else:
                        # 使用 AI 生成内容
                        # 调用 AI 前检查任务是否被取消
                        task_info = task_manager.load_task_info(task_id)
                        if task_info and task_info.status == 'cancelled':
                            print(f'[CANCEL] 任务已取消，跳过 AI 生成')
                            return
                            
                        from ai_engine import generate_section_content_with_ai
                        content = generate_section_content_with_ai(
                            section_title=node_title,
                            requirement_text=requirement_content,
                            template_text=template_content,
                            user_instruction=user_prompt,
                            model_config=model_config,
                            fallback_to_template=False
                        )
                        
                        # AI 返回后检查任务是否被取消
                        task_info = task_manager.load_task_info(task_id)
                        if task_info and task_info.status == 'cancelled':
                            print(f'[CANCEL] 任务已取消，停止生成')
                            return

                    if content.startswith('[AI 生成失败]') or content.startswith('[API'):
                        print(f'[ERROR] 章节生成失败：{node_title}')
                        add_normal_paragraph(doc, f"【生成错误】{content}", styles=styles)
                        task_manager.update_chapter_status(task_id, node_index, status='failed',
                            error_message=content)
                    else:
                        for para_text in content.split('\n'):
                            if para_text.strip():
                                add_normal_paragraph(doc, para_text.strip(), styles=styles)
                        task_manager.update_chapter_status(task_id, node_index, status='completed',
                            content=content, word_count=len(content))
                else:
                    content = get_chapter_content_template(node_title)
                    for para_text in content.split('\n'):
                        if para_text.strip():
                            add_normal_paragraph(doc, para_text.strip(), styles=styles)
                    task_manager.update_chapter_status(task_id, node_index, status='completed',
                        content=content, word_count=len(content))

                # 更新进度
                progress = 15 + int((node_index + 1) / total_chapters * 80)
                task_manager.update_task_progress(task_id, progress=progress,
                    message=f'正在生成第{node_index + 1}章：{node_title}')

                save_partial_document()

        # 遍历一级章节
        for idx, chapter in enumerate(user_chapters):
            render_chapter_with_status(chapter, level=1)

        doc.save(output_path)
        
        # 质量审校
        task_manager.update_task_progress(task_id, progress=98,
            message='正在进行质量审校...')
        try:
            # 获取所有数据点和需求点
            data_points = dp_manager.get_all()
            requirements = req_analyzer.get_all_requirements()
            
            # 生成审校报告
            review_report = quality_reviewer.generate_report(
                chapter_contents={},  # v2 版本没有逐章记录内容
                data_points=data_points,
                requirements=requirements
            )
            
            # 保存审校报告
            report_filename = f'review_report_{task_id[:8]}.md'
            report_path = os.path.join(app.config['OUTPUT_FOLDER'], report_filename)
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(review_report.to_markdown())
            
            print(f'[INFO] 质量审校完成：{report_path}')
            print(f'[INFO] 综合评分：{review_report.overall_score:.1f}/100')
            print(f'[INFO] 数据一致性：{review_report.data_consistency_rate:.1%}')
            print(f'[INFO] 需求覆盖率：{review_report.requirement_coverage_rate:.1%}')
            
            # 如果有冲突或未覆盖需求，输出警告
            if review_report.consistency_issues:
                print(f'[WARN] 发现 {len(review_report.consistency_issues)} 个数据一致性问题')
            if review_report.coverage_issues:
                print(f'[WARN] 发现 {len(review_report.coverage_issues)} 个未覆盖需求')
                
        except Exception as e:
            print(f'[WARN] 质量审校失败：{e}')
        
        task_manager.mark_task_completed(task_id, output_filename=os.path.basename(output_path))
        print(f'[INFO] 文档生成完成：{output_path}')

    except Exception as e:
        import traceback
        print(f'[ERROR] 文档生成失败：{e}')
        task_manager.mark_task_failed(task_id, f'{str(e)}\n{traceback.format_exc()}')


def process_document_async_v3(task_id, template_type, requirement_content, template_content,
                              user_prompt, model_config, output_path):
    """
    异步处理文档生成任务 - v3 支持持久化、实时预览和章节编辑
    
    主要改进:
    1. 使用 task_manager 持久化任务状态
    2. 每完成一章就保存状态和临时文档
    3. 支持暂停/继续
    4. 详细的日志输出
    
    Args:
        task_id: 任务 ID
        template_type: 模板类型
        requirement_content: 需求文档内容
        template_content: 模板文档内容
        user_prompt: 用户补充要求
        model_config: 模型配置对象
        output_path: 输出文件路径
    """
    from task_manager import get_task_manager
    from ai_engine import generate_section_content_with_ai, extract_template_section, get_document_structure
    
    task_manager = get_task_manager()
    doc_container = {'doc': None}
    
    # 初始化优化服务
    reset_optimization_services()
    dp_manager = get_data_point_manager()
    req_analyzer = get_requirement_analyzer()
    quality_reviewer = QualityReviewer()
    
    def save_partial_document():
        """保存临时文档到任务目录"""
        if doc_container['doc']:
            task_manager.save_partial_document(task_id, doc_container['doc'])
    
    try:
        # ========== 任务启动日志 ==========
        print(f'\n{"="*60}')
        print(f'[TASK] 任务启动：{task_id}')
        print(f'[TASK] 模板类型：{template_type}')
        print(f'[TASK] 模型：{model_config.name} ({model_config.id})')
        print(f'[TASK] 需求文档长度：{len(requirement_content) if requirement_content else 0} 字符')
        print(f'[TASK] 模板文档长度：{len(template_content) if template_content else 0} 字符')
        print(f'[TASK] 用户补充要求：{user_prompt[:50] if user_prompt else "无"}...')
        print(f'[TASK] 输出路径：{output_path}')
        print(f'{"="*60}\n')
        
        # 更新任务状态
        task_manager.update_task_status(task_id, status='generating', started_at=datetime.now().strftime('%Y-%m-%d %H:%M:%S'))

        # ========== 加载目录配置并初始化章节列表 ==========
        # 注意：先初始化章节列表，以便前端可以立即显示章节
        chapters_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chapter_config.json')
        user_chapters = None

        if os.path.exists(chapters_path):
            try:
                with open(chapters_path, 'r', encoding='utf-8') as f:
                    user_chapters = json.load(f)
                print(f'\n[CHAPTER] 已加载用户目录配置：{len(user_chapters)} 个一级章节')
            except Exception as e:
                print(f'[CHAPTER] 加载目录配置失败：{e}')
                user_chapters = None

        if not user_chapters:
            # 使用默认模板配置
            template_config = TEMPLATE_TYPES.get(template_type, TEMPLATE_TYPES['future_community'])
            user_chapters = []
            for chapter_title, sections in template_config['chapters']:
                chapter_node = {
                    'number': chapter_title.split()[0] if ' ' in chapter_title else '1',
                    'title': chapter_title,
                    'level': 1,
                    'children': []
                }
                for idx, section in enumerate(sections):
                    section_num = f'{chapter_node["number"]}.{idx+1}'
                    chapter_node['children'].append({
                        'number': section_num,
                        'title': section,
                        'level': 2,
                        'children': []
                    })
            print(f'\n[CHAPTER] 使用默认模板配置：{len(template_config["chapters"])} 个一级章节')

        # 初始化章节列表（使用树形结构）
        task_manager.initialize_chapters_with_tree(task_id, user_chapters)
        print(f'\n[CHAPTER] 初始化章节列表（树形）：{len(user_chapters)} 个一级章节')

        # ========== 从需求文档提取数据点 ==========
        if requirement_content:
            print(f'\n[DATA] 开始从需求文档提取数据点...')
            initial_data = dp_manager.extract_from_text(requirement_content, "需求文档")
            if initial_data:
                dp_manager.update(initial_data, "需求文档")
                print(f'[DATA] 提取到数据点：{list(initial_data.keys())}')
                for key, value in initial_data.items():
                    print(f'       - {key}: {value}')
            else:
                print(f'[DATA] 未提取到数据点')

            # 提取需求点
            print(f'\n[REQ] 开始从需求文档提取需求点...')
            req_analyzer.extract(requirement_content)
            req_summary = req_analyzer.get_summary()
            print(f'[REQ] 提取到需求点：')
            for req_type, items in req_summary.items():
                if items:
                    print(f'       - {req_type}: {len(items)} 个')

        # ========== 创建 Word 文档 ==========
        styles = load_style_config()
        doc = Document()
        normal_style = styles.get('normal', {})
        style = doc.styles['Normal']
        style.font.name = normal_style.get('font_name', '仿宋')
        style.font.size = Pt(normal_style.get('font_size', 10.5))
        style._element.rPr.rFonts.set(qn('w:eastAsia'), normal_style.get('font_name', '仿宋'))
        
        project_name = "建设项目"
        if dp_manager.get('项目名称'):
            project_name = dp_manager.get('项目名称')
        
        # 添加封面
        print(f'\n[DOC] 开始生成封面...')
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Cm(3)
        title.paragraph_format.space_after = Cm(1)
        run = title.add_run(project_name)
        run.font.name = '黑体'
        run.font.size = Pt(36)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Cm(1)
        subtitle.paragraph_format.space_after = Cm(3)
        run = subtitle.add_run(template_config['name'])
        run.font.name = '黑体'
        run.font.size = Pt(26)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        for _ in range(5):
            doc.add_paragraph()
        
        # 保存初始文档
        doc_container['doc'] = doc
        save_partial_document()
        
        # ========== 逐章生成 ==========
        total_chapters = len(chapter_titles)
        
        for idx, chapter_title in enumerate(chapter_titles):
            # 检查任务状态
            task_info = task_manager.load_task_info(task_id)
            
            if task_info and task_info.status == 'paused':
                print(f'\n[PAUSE] 任务已暂停，等待用户继续...')
                while True:
                    time.sleep(2)
                    task_info = task_manager.load_task_info(task_id)
                    if task_info.status != 'paused':
                        break
                print(f'[RESUME] 任务已恢复')
            
            if task_info and task_info.status == 'cancelled':
                print(f'\n[CANCEL] 任务已取消，停止生成')
                return
            
            # 更新当前章节索引
            task_manager.update_task_status(task_id, current_chapter_index=idx)
            
            # 更新章节状态
            task_manager.update_chapter_status(task_id, idx, status='generating')
            
            print(f'\n{"-"*60}')
            print(f'[CHAPTER {idx+1}/{total_chapters}] {chapter_title}')
            print(f'{"-"*60}')
            
            try:
                # 提取模板章节内容
                template_section = extract_template_section(chapter_title, template_content)
                doc_structure = get_document_structure(template_content)
                
                print(f'[TEMPLATE] 提取模板章节内容长度：{len(template_section)} 字符')
                print(f'[TEMPLATE] 文档结构：{len(doc_structure.split(chr(10)))} 章')

                # AI 调用前检查取消状态
                task_info = task_manager.load_task_info(task_id)
                if task_info and task_info.status == 'cancelled':
                    print(f'[CANCEL] 任务已取消，停止生成')
                    return

                # 生成章节内容
                content = generate_section_content_with_ai(
                    section_title=chapter_title,
                    requirement_text=requirement_content,
                    template_text=template_section,
                    user_instruction=user_prompt,
                    model_config=model_config,
                    fallback_to_template=False
                )
                
                # AI 返回后检查取消状态
                task_info = task_manager.load_task_info(task_id)
                if task_info and task_info.status == 'cancelled':
                    print(f'[CANCEL] 任务已取消，停止生成')
                    return

                # 检查生成结果
                if content.startswith('[AI 生成失败]'):
                    raise Exception(content)
                
                print(f'[CONTENT] 生成成功，字数：{len(content)}')
                
                # 清理内容
                from ai_engine import clean_ai_content
                content = clean_ai_content(content, chapter_title)
                print(f'[CONTENT] 清理后字数：{len(content)}')
                
                # 更新章节状态
                task_manager.update_chapter_status(
                    task_id, idx,
                    status='completed',
                    content=content,
                    word_count=len(content)
                )
                
                # 添加到文档
                add_heading(doc, chapter_title, level=1)
                for para_text in content.split('\n'):
                    if para_text.strip():
                        add_normal_paragraph(doc, para_text.strip())
                
                # 保存临时文档
                save_partial_document()
                
            except Exception as e:
                print(f'[ERROR] 章节生成失败：{str(e)}')
                task_manager.update_chapter_status(
                    task_id, idx,
                    status='failed',
                    error_message=str(e)
                )
        
        # ========== 完成文档 ==========
        print(f'\n{"="*60}')
        print(f'[COMPLETE] 所有章节生成完成')
        print(f'[COMPLETE] 保存最终文档：{output_path}')
        print(f'{"="*60}\n')
        
        doc.save(output_path)
        
        # 质量审校
        print(f'\n[REVIEW] 开始质量审校...')
        try:
            data_points = dp_manager.get_all()
            requirements = req_analyzer.get_all_requirements()
            
            review_report = quality_reviewer.generate_report(
                chapter_contents={},
                data_points=data_points,
                requirements=requirements
            )
            
            report_filename = f'review_report_{task_id[:8]}.md'
            report_path = os.path.join(app.config['OUTPUT_FOLDER'], report_filename)
            with open(report_path, 'w', encoding='utf-8') as f:
                f.write(review_report.to_markdown())
            
            print(f'[REVIEW] 审校报告已保存：{report_path}')
            print(f'[REVIEW] 综合评分：{review_report.overall_score:.1f}/100')
            print(f'[REVIEW] 数据一致性：{review_report.data_consistency_rate:.1%}')
            print(f'[REVIEW] 需求覆盖率：{review_report.requirement_coverage_rate:.1%}')
            
        except Exception as e:
            print(f'[REVIEW] 质量审校失败：{e}')
        
        # 保存最终文档到任务目录
        task_manager.save_final_document(task_id, output_path)
        
        # 更新任务状态
        task_manager.update_task_status(task_id, status='completed')
        
        print(f'\n[TASK] 任务完成：{task_id}')
        
    except Exception as e:
        import traceback
        print(f'\n[ERROR] 任务失败：{task_id}')
        print(f'[ERROR] 错误信息：{str(e)}')
        print(f'[ERROR] 堆栈跟踪:\n{traceback.format_exc()}')
        
        task_manager.update_task_status(task_id, status='failed')


# ==================== Flask 路由 ====================

@app.route('/task-monitor')
def task_monitor():
    """任务监控页面"""
    response = make_response(render_template('task_monitor.html'))
    response.headers['Cache-Control'] = 'no-cache, no-store, must-revalidate'
    response.headers['Pragma'] = 'no-cache'
    response.headers['Expires'] = '0'
    return response


@app.route('/')
def index():
    return render_template('index.html', template_types=TEMPLATE_TYPES)


@app.route('/task-center')
def task_center():
    """文件生成管理中心"""
    return render_template('task_center.html')


@app.route('/model-config')
def model_config_page():
    """模型配置管理页面"""
    return render_template('model_config.html')


# ==================== 代码生成接口（A 方案） ====================

@app.route('/api/generate_code', methods=['POST'])
def generate_code():
    """
    根据目录结构和需求文档生成完整的 Python 代码
    """
    try:
        # 获取参数
        template_type = request.form.get('template_type', 'future_community')
        
        # 读取需求文档
        requirement_file = request.files.get('requirement_file')
        requirement_content = ''

        if requirement_file and allowed_file(requirement_file.filename):
            filename = secure_filename(requirement_file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], f'{uuid.uuid4()}_{filename}')
            requirement_file.save(filepath)

            if filename.endswith('.docx'):
                requirement_content = read_docx_text(filepath)
                print(f'[DEBUG] 读取需求文档：{filepath}, 长度：{len(requirement_content)}', flush=True)
            elif filename.endswith('.txt'):
                requirement_content = read_txt_text(filepath)
                print(f'[DEBUG] 读取需求文档：{filepath}, 长度：{len(requirement_content)}', flush=True)
            elif filename.endswith('.doc'):
                print(f'[WARNING] 暂不支持读取 .doc 格式文件，请另存为 .docx 格式：{filepath}', flush=True)
            elif filename.endswith('.pdf'):
                print(f'[WARNING] 暂不支持读取 .pdf 格式文件，请转换为 .txt 或 .docx 格式：{filepath}', flush=True)
            else:
                print(f'[WARNING] 未知文件格式：{filepath}', flush=True)

            # 同时保存为 requirement.txt 方便后续使用
            req_path = os.path.join(app.config['UPLOAD_FOLDER'], 'requirement.txt')
            with open(req_path, 'w', encoding='utf-8') as f:
                f.write(requirement_content)
        else:
            print(f'[WARNING] 需求文件未上传或格式不允许：{requirement_file}', flush=True)
        
        # 加载目录配置
        chapters_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chapter_config.json')
        if not os.path.exists(chapters_path):
            return jsonify({
                'success': False,
                'message': '请先扫描模板并保存目录结构'
            }), 400
        
        with open(chapters_path, 'r', encoding='utf-8') as f:
            chapters = json.load(f)
        
        # 加载样式配置
        styles = load_style_config()
        
        # 生成代码
        from generators import generate_word_code
        python_code = generate_word_code(chapters, styles, requirement_content, template_type)
        
        # 保存生成的代码
        code_filename = f'generated_{datetime.now().strftime("%Y%m%d_%H%M%S")}.py'
        code_path = os.path.join(app.config['UPLOAD_FOLDER'], code_filename)
        with open(code_path, 'w', encoding='utf-8') as f:
            f.write(python_code)
        
        return jsonify({
            'success': True,
            'message': '代码生成成功',
            'code_filename': code_filename,
            'code_preview': python_code[:2000]  # 返回前 2000 字符预览
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'生成失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/execute_code', methods=['POST'])
def execute_code():
    """
    执行生成的 Python 代码，生成 Word 文档
    """
    try:
        data = request.get_json()
        code_filename = data.get('code_filename', '')
        
        if not code_filename:
            return jsonify({
                'success': False,
                'message': '请指定要执行的代码文件'
            }), 400
        
        code_path = os.path.join(app.config['UPLOAD_FOLDER'], code_filename)
        if not os.path.exists(code_path):
            return jsonify({
                'success': False,
                'message': '代码文件不存在'
            }), 404
        
        # 读取代码
        with open(code_path, 'r', encoding='utf-8') as f:
            python_code = f.read()
        
        # 创建临时目录执行代码
        import tempfile
        import subprocess
        
        # 在执行时重定向日志
        temp_dir = tempfile.mkdtemp()
        temp_script = os.path.join(temp_dir, 'run.py')
        
        # 修改代码，添加输出路径
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'generated_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
        
        # 执行代码
        exec_globals = {
            '__name__': '__main__',
            '__file__': code_path
        }

        # 为了安全，使用受限的执行环境
        # 这里直接执行生成的代码
        exec(compile(python_code, code_path, 'exec'), exec_globals)
        
        # 查找生成的文件
        output_files = [f for f in os.listdir('.') if f.endswith('.docx') and f.startswith('generated_')]
        output_files.sort(reverse=True)
        
        if output_files:
            output_file = output_files[0]
            return jsonify({
                'success': True,
                'message': '文档生成成功',
                'output_file': output_file,
                'download_url': f'/api/download/{output_file}'
            })
        else:
            return jsonify({
                'success': False,
                'message': '未找到生成的文件'
            }), 500
        
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'执行失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/download/<filename>', methods=['GET'])
def download_file(filename):
    """下载生成的文件"""
    from flask import send_file
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return jsonify({'success': False, 'message': '文件不存在'}), 404


# ==================== 原有接口 ====================

@app.route('/api/generate', methods=['POST'])
def generate():
    """异步提交文档生成任务"""
    try:
        # 获取表单数据
        template_type = request.form.get('template_type', 'future_community')
        user_prompt = request.form.get('requirement', '')
        model = request.form.get('model', 'qwen-max')

        # 获取模型配置（使用 v2 版本）
        from model_config_v2 import get_model_config_v2
        model_config = get_model_config_v2(model)
        if not model_config:
            return jsonify({'success': False, 'message': f'模型 {model} 不存在'}), 400

        if not model_config.enabled:
            return jsonify({'success': False, 'message': f'模型 {model_config.name} 已禁用，请前往模型配置管理页面启用'}), 400

        # Ollama 本地模型不需要 API Key
        if model_config.provider_id != 'ollama' and not model_config.api_key:
            return jsonify({'success': False, 'message': f'模型 {model_config.name} 的 API Key 未配置，请前往模型配置管理页面配置'}), 400

        # 使用新的 v2 任务管理器（支持持久化）
        from task_manager import get_task_manager
        task_manager_v2 = get_task_manager()

        # 处理上传的文件
        template_file = request.files.get('template')
        requirement_file = request.files.get('requirement_file')

        template_content = ''
        requirement_content = ''
        template_filename = ''
        requirement_filename = ''

        # 先创建任务，获取 task_id
        task_id = task_manager_v2.create_task(
            template_type=template_type,
            user_prompt=user_prompt,
            model=model
        )

        # 读取模板文件并保存到任务目录
        if template_file and allowed_file(template_file.filename):
            original_filename = secure_filename(template_file.filename)
            # 确保保留扩展名
            if '.' not in original_filename:
                ext = template_file.filename.rsplit('.', 1)[1].lower()
                original_filename = f'{original_filename}.{ext}'
            
            # 保存到任务目录
            file_ext = original_filename.rsplit('.', 1)[-1]
            task_manager_v2.save_file(task_id, f'template.{file_ext}', template_file.read())
            template_filename = original_filename

            # 读取内容
            tmpl_path = task_manager_v2.get_file_path(task_id, f'template.{file_ext}')
            if tmpl_path:
                if original_filename.endswith('.docx'):
                    template_content = read_docx_text(tmpl_path)
                    print(f'[DEBUG] 读取模板文件：{tmpl_path}, 长度：{len(template_content)}', flush=True)
                elif original_filename.endswith('.txt'):
                    template_content = read_txt_text(tmpl_path)
                    print(f'[DEBUG] 读取模板文件：{tmpl_path}, 长度：{len(template_content)}', flush=True)
        else:
            print(f'[WARNING] 模板文件未上传或格式不允许：{template_file}', flush=True)

        # 读取需求文件并保存到任务目录
        if requirement_file and allowed_file(requirement_file.filename):
            original_filename = secure_filename(requirement_file.filename)
            # 确保保留扩展名
            if '.' not in original_filename:
                ext = requirement_file.filename.rsplit('.', 1)[1].lower()
                original_filename = f'{original_filename}.{ext}'
            
            # 保存到任务目录
            file_ext = original_filename.rsplit('.', 1)[-1]
            task_manager_v2.save_file(task_id, f'requirement.{file_ext}', requirement_file.read())
            requirement_filename = original_filename

            # 读取内容
            req_path = task_manager_v2.get_file_path(task_id, f'requirement.{file_ext}')
            if req_path:
                if original_filename.endswith('.docx'):
                    requirement_content = read_docx_text(req_path)
                    print(f'[DEBUG] 读取需求文件：{req_path}, 长度：{len(requirement_content)}', flush=True)
                elif original_filename.endswith('.txt'):
                    requirement_content = read_txt_text(req_path)
                    print(f'[DEBUG] 读取需求文件：{req_path}, 长度：{len(requirement_content)}', flush=True)
                elif original_filename.endswith('.doc'):
                    print(f'[WARNING] 暂不支持读取 .doc 格式文件，请另存为 .docx 格式：{req_path}', flush=True)
                elif original_filename.endswith('.pdf'):
                    print(f'[WARNING] 暂不支持读取 .pdf 格式文件，请转换为 .txt 或 .docx 格式：{req_path}', flush=True)
                else:
                    print(f'[WARNING] 未知文件格式：{req_path}', flush=True)
        else:
            print(f'[WARNING] 需求文件未上传或格式不允许：{requirement_file}', flush=True)

        # 更新任务信息（添加文件名）
        task_manager_v2.update_task_status(
            task_id,
            requirement_filename=requirement_filename,
            template_filename=template_filename
        )

        # 生成输出文件名
        task_id_short = task_id[:8]
        output_filename = f'{template_type}_可行性研究报告_{task_id_short}.docx'
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], output_filename)

        # 启动异步线程处理任务 - 使用修复后的 v2 版本
        thread = threading.Thread(
            target=process_document_async_v2,
            args=(task_id, template_type, requirement_content, template_content,
                  user_prompt, model_config, output_path)
        )
        thread.daemon = True
        thread.start()

        return jsonify({
            'success': True,
            'message': '任务已提交，正在后台处理中',
            'task_id': task_id,
            'status': 'pending'
        })

    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'提交失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/task/status/<task_id>', methods=['GET'])
def task_status(task_id):
    """查询任务状态和进度 - 使用新版 TaskManager"""
    try:
        from task_manager import get_task_manager

        task_manager_v2 = get_task_manager()
        status = task_manager_v2.get_task_status(task_id)

        if not status:
            return jsonify({
                'success': False,
                'message': '任务不存在'
            }), 404

        # 状态映射：持久化版本 -> 前端版本
        STATUS_MAP = {
            'pending': 'pending',
            'generating': 'generating_ai',
            'paused': 'processing',
            'completed': 'completed',
            'failed': 'failed',
            'cancelled': 'cancelled',
            'partially_completed': 'completed'
        }

        # 获取已完成的章节标题列表（用于前端显示）
        completed_chapter_titles = []
        if status.get('chapters'):
            for ch in status.get('chapters', []):
                if ch.get('status') == 'completed':
                    completed_chapter_titles.append(ch.get('title'))

        # 转换为旧版格式兼容前端
        safe_info = {
            'task_id': status.get('task_id'),
            'template_type': status.get('template_type'),
            'status': STATUS_MAP.get(status.get('status'), status.get('status')),
            'progress': status.get('progress'),
            'message': status.get('message'),
            'created_at': status.get('created_at'),
            'updated_at': status.get('updated_at'),
            'started_at': status.get('started_at'),
            'completed_at': status.get('completed_at'),
            'output_filename': status.get('output_filename'),
            'partial_filename': status.get('partial_doc_url'),  # URL 路径
            'model': status.get('model'),
            'chapter_steps': [],
            'current_chapter': status.get('current_chapter'),
            'completed_chapters': completed_chapter_titles,  # 章节标题列表（用于显示）
            'completed_chapters_count': status.get('completed_chapters', 0),  # 已完成章节数（数字）
            'total_chapters': status.get('total_chapters', 0),
            'pending_confirmation': False
        }

        return jsonify({
            'success': True,
            'task': safe_info
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'查询失败：{str(e)}'
        }), 500


@app.route('/api/task/list', methods=['GET'])
def task_list():
    """获取任务列表（支持分页和搜索）- 使用新版 TaskManager"""
    try:
        from task_manager import get_task_manager, TASKS_ROOT
        import os

        task_manager_v2 = get_task_manager()
        page = request.args.get('page', 1, type=int)
        page_size = request.args.get('page_size', 10, type=int)
        keyword = request.args.get('keyword', '')
        status_filter = request.args.get('status', '')

        tasks = []

        # 遍历任务目录
        if os.path.exists(TASKS_ROOT):
            for task_id in os.listdir(TASKS_ROOT):
                task_dir = os.path.join(TASKS_ROOT, task_id)
                if not os.path.isdir(task_dir):
                    continue

                task_info = task_manager_v2.load_task_info(task_id)
                if task_info:
                    # 应用状态筛选
                    if status_filter and task_info.status != status_filter:
                        continue

                    # 应用关键词筛选
                    if keyword:
                        task_str = f"{task_info.task_id} {task_info.template_type} {task_info.model}"
                        if keyword.lower() not in task_str.lower():
                            continue

                    tasks.append({
                        'task_id': task_info.task_id,
                        'template_type': task_info.template_type,
                        'status': task_info.status,
                        'progress': task_info.progress,
                        'message': task_info.message or '',
                        'model': task_info.model,
                        'created_at': task_info.created_at,
                        'started_at': task_info.started_at,
                        'completed_at': task_info.completed_at,
                        'output_filename': task_info.output_filename,
                        'partial_filename': task_info.partial_filename,
                        'requirement_file': task_info.requirement_filename,
                        'template_file': task_info.template_filename
                    })

        # 按创建时间排序（最新的在前）
        tasks.sort(key=lambda x: x['created_at'], reverse=True)

        total = len(tasks)
        start = (page - 1) * page_size
        end = start + page_size
        tasks = tasks[start:end]

        return jsonify({
            'success': True,
            'tasks': tasks,
            'total': total,
            'page': page,
            'page_size': page_size
        })

    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'查询失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/task/cancel/<task_id>', methods=['POST'])
def cancel_task(task_id):
    """取消任务 - 使用新版 TaskManager"""
    try:
        from task_manager import get_task_manager

        task_manager_v2 = get_task_manager()
        success = task_manager_v2.cancel_task(task_id)

        if success:
            return jsonify({
                'success': True,
                'message': '任务已取消'
            })
        else:
            return jsonify({
                'success': False,
                'message': '任务已完成或失败，无法取消'
            }), 400

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'取消失败：{str(e)}'
        }), 500


@app.route('/api/task/delete/<task_id>', methods=['POST'])
def delete_task(task_id):
    """删除任务 - 使用新版 TaskManager"""
    try:
        from task_manager import get_task_manager

        task_manager_v2 = get_task_manager()
        success = task_manager_v2.delete_task(task_id)

        if success:
            return jsonify({
                'success': True,
                'message': '任务已删除'
            })
        else:
            return jsonify({
                'success': False,
                'message': '任务不存在或删除失败'
            }), 400

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'删除失败：{str(e)}'
        }), 500


@app.route('/api/task/continue/<task_id>', methods=['POST'])
def continue_task(task_id):
    """继续生成下一章（用户确认后的回调）- 使用新版 TaskManager"""
    try:
        from task_manager import get_task_manager

        task_manager_v2 = get_task_manager()
        task_info = task_manager_v2.get_task_status(task_id)

        if not task_info:
            return jsonify({
                'success': False,
                'message': '任务不存在'
            }), 404

        if task_info.get('status') not in ['processing', 'generating_ai', 'creating_doc']:
            return jsonify({
                'success': False,
                'message': '任务状态不允许继续生成'
            }), 400

        # 继续任务
        task_manager_v2.continue_task(task_id)

        return jsonify({
            'success': True,
            'message': '已确认，继续生成下一章'
        })

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'确认失败：{str(e)}'
        }), 500


@app.route('/api/task/preview/<task_id>', methods=['GET'])
def preview_task(task_id):
    """预览当前已生成的文档内容 - 使用新版 TaskManager"""
    try:
        from task_manager import get_task_manager

        task_manager_v2 = get_task_manager()
        task_info = task_manager_v2.load_task_info(task_id)

        if not task_info:
            return jsonify({
                'success': False,
                'message': '任务不存在'
            }), 404

        # 检查是否有部分生成的文档
        if task_info.partial_filename:
            filepath = os.path.join(app.config['OUTPUT_FOLDER'], task_info.partial_filename)
            if os.path.exists(filepath):
                return send_file(filepath)

        # 或者检查已完成的文档
        if task_info.output_filename:
            filepath = os.path.join(app.config['OUTPUT_FOLDER'], task_info.output_filename)
            if os.path.exists(filepath):
                return send_file(filepath)

        return jsonify({
            'success': False,
            'message': '暂无可预览的文档'
        }), 404

    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'预览失败：{str(e)}'
        }), 500


@app.route('/api/download/<filename>')
def download(filename):
    """下载生成的文档"""
    filepath = os.path.join(app.config['OUTPUT_FOLDER'], filename)
    if os.path.exists(filepath):
        return send_file(filepath, as_attachment=True)
    return jsonify({'error': '文件不存在'}), 404


@app.route('/api/template/scan', methods=['POST'])
def scan_template():
    """扫描上传的模板文件，提取章节目录和样式信息"""
    try:
        if 'template' not in request.files:
            return jsonify({'success': False, 'message': '请上传模板文件'}), 400
        
        file = request.files['template']
        if file.filename == '':
            return jsonify({'success': False, 'message': '未选择文件'}), 400
        
        if not file.filename.endswith('.docx'):
            return jsonify({'success': False, 'message': '请上传 .docx 格式的 Word 文件'}), 400
        
        # 保存临时文件
        temp_path = os.path.join(app.config['UPLOAD_FOLDER'], f'temp_scan_{uuid.uuid4().hex[:8]}.docx')
        file.save(temp_path)
        
        try:
            # 扫描模板样式
            result = scan_template_styles(temp_path)
            
            if result['success']:
                return jsonify({
                    'success': True,
                    'message': result['message'],
                    'chapters': result['chapters'],
                    'total_nodes': result.get('total_nodes', 0)
                })
            else:
                return jsonify({
                    'success': False,
                    'message': result['message']
                }), 400
        finally:
            # 删除临时文件
            if os.path.exists(temp_path):
                os.remove(temp_path)
                
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'扫描失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/chapters/save', methods=['POST'])
def save_chapters():
    """保存用户编辑后的章节目录配置"""
    try:
        data = request.json
        chapters = data.get('chapters', [])
        
        if not chapters:
            return jsonify({
                'success': False,
                'message': '章节目录不能为空'
            }), 400
        
        # 将章节目录保存到文件
        chapters_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chapter_config.json')
        with open(chapters_path, 'w', encoding='utf-8') as f:
            json.dump(chapters, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': '章节目录配置已保存'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'保存失败：{str(e)}'
        }), 500


@app.route('/api/chapters/load', methods=['GET'])
def load_chapters():
    """加载已保存的章节目录配置"""
    try:
        chapters_path = os.path.join(app.config['UPLOAD_FOLDER'], 'chapter_config.json')
        
        if os.path.exists(chapters_path):
            with open(chapters_path, 'r', encoding='utf-8') as f:
                chapters = json.load(f)
            return jsonify({
                'success': True,
                'chapters': chapters,
                'message': '已加载保存的章节目录配置'
            })
        else:
            return jsonify({
                'success': True,
                'chapters': [],
                'message': '未找到保存的章节目录配置'
            })
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'加载失败：{str(e)}'
        }), 500


@app.route('/api/styles/save', methods=['POST'])
def save_styles():
    """保存用户配置的样式"""
    try:
        data = request.json
        
        # 样式配置包括：
        # - heading1: 一级标题样式（章标题）
        # - heading2: 二级标题样式（节标题）
        # - heading3: 三级标题样式（小节标题）
        # - normal: 正文样式
        styles = data.get('styles', {})
        
        # 将样式配置保存到文件
        styles_path = os.path.join(app.config['UPLOAD_FOLDER'], 'style_config.json')
        with open(styles_path, 'w', encoding='utf-8') as f:
            json.dump(styles, f, ensure_ascii=False, indent=2)
        
        return jsonify({
            'success': True,
            'message': '样式配置已保存'
        })
        
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'保存失败：{str(e)}'
        }), 500


@app.route('/api/styles/load', methods=['GET'])
def load_styles():
    """加载已保存的样式配置"""
    try:
        styles_path = os.path.join(app.config['UPLOAD_FOLDER'], 'style_config.json')
        
        if os.path.exists(styles_path):
            with open(styles_path, 'r', encoding='utf-8') as f:
                styles = json.load(f)
            return jsonify({
                'success': True,
                'styles': styles
            })
        else:
            # 返回默认样式
            default_styles = {
                'heading1': {
                    'font_name': '黑体',
                    'font_size': 22,
                    'bold': True,
                    'alignment': 'center'
                },
                'heading2': {
                    'font_name': '楷体',
                    'font_size': 16,
                    'bold': True,
                    'alignment': 'left'
                },
                'heading3': {
                    'font_name': '楷体',
                    'font_size': 14,
                    'bold': True,
                    'alignment': 'left'
                },
                'normal': {
                    'font_name': '仿宋',
                    'font_size': 10.5,
                    'line_spacing': 1.5,
                    'first_line_indent': 0.74
                }
            }
            return jsonify({
                'success': True,
                'styles': default_styles,
                'message': '未找到自定义配置，返回默认样式'
            })
            
    except Exception as e:
        return jsonify({
            'success': False,
            'message': f'加载失败：{str(e)}'
        }), 500


@app.route('/api/validate-key', methods=['POST'])
def validate_api_key():
    """验证 API Key 是否有效"""
    try:
        data = request.json
        api_key = data.get('api_key', '')
        model = data.get('model', 'qwen-max')

        if not api_key:
            return jsonify({'valid': False, 'message': 'API Key 不能为空'})

        # 检测 API Key 类型
        is_app_api = api_key.startswith('sk-sp-') or api_key.startswith('sk-app-')
        
        if is_app_api:
            # 百炼应用 API 验证
            app_id = None
            test_key = api_key
            if '#' in api_key:
                test_key, app_id = api_key.split('#', 1)
            else:
                app_id = APP_CONFIGS.get('default', {}).get('app_id')
            
            if not app_id:
                return jsonify({
                    'valid': False, 
                    'message': '百炼应用 API Key (sk-sp-) 需要配置 APP_ID。请在 app.py 的 APP_CONFIGS 中设置，或在 API Key 后添加 #APP_ID（例如：sk-sp-xxx#your_app_id）'
                })
            
            url = f"https://dashscope.aliyuncs.com/api/v1/apps/{app_id}/completion"
            payload = {
                "input": {"prompt": "你好"},
                "parameters": {"max_tokens": 10}
            }
        else:
            # DashScope 模型 API 验证
            url = "https://dashscope.aliyuncs.com/api/v1/services/aigc/text-generation/generation"
            payload = {
                "model": model,
                "input": {"messages": [{"role": "user", "content": "你好"}]},
                "parameters": {"max_tokens": 10}
            }

        headers = {"Authorization": f"Bearer {test_key if is_app_api else api_key}", "Content-Type": "application/json"}
        response = requests.post(url, headers=headers, json=payload, timeout=15)

        if response.status_code == 200:
            if is_app_api:
                return jsonify({'valid': True, 'message': f'百炼应用 API Key 有效 (APP_ID: {app_id})'})
            else:
                return jsonify({'valid': True, 'message': f'{MODEL_CONFIGS.get(model, {}).get("name", model)} API Key 有效'})
        elif response.status_code == 401:
            return jsonify({'valid': False, 'message': 'API Key 无效或已过期，请检查 Key 是否正确'})
        elif response.status_code == 403:
            return jsonify({'valid': False, 'message': '权限不足，请检查是否开通了对应���务'})
        elif response.status_code == 404 and is_app_api:
            return jsonify({'valid': False, 'message': f'APP_ID ({app_id}) 不存在或无权访问'})
        elif response.status_code == 429:
            return jsonify({'valid': False, 'message': '请求频率超限，请稍后重试'})
        elif response.status_code == 500:
            return jsonify({'valid': False, 'message': '服务器内部错误，请稍后重试'})
        else:
            error_msg = response.text[:200] if response.text else '无响应内容'
            return jsonify({'valid': False, 'message': f'验证失败 (状态码 {response.status_code}): {error_msg}'})

    except requests.exceptions.Timeout:
        return jsonify({'valid': False, 'message': '验证超时，请检查网络连接'})
    except requests.exceptions.ConnectionError:
        return jsonify({'valid': False, 'message': '无法连接到 API 服务器，请检查网络'})
    except Exception as e:
        return jsonify({'valid': False, 'message': f'验证异常：{str(e)}'})


@app.route('/api/cleanup', methods=['POST'])
def cleanup():
    """清理过期文件"""
    try:
        import time
        current_time = time.time()
        cleaned = 0

        for folder in [app.config['UPLOAD_FOLDER'], app.config['OUTPUT_FOLDER']]:
            for filename in os.listdir(folder):
                filepath = os.path.join(folder, filename)
                if os.path.isfile(filepath) and current_time - os.path.getctime(filepath) > 3600:
                    os.remove(filepath)
                    cleaned += 1

        return jsonify({'success': True, 'cleaned': cleaned})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 模型配置管理 API ====================

@app.route('/api/models', methods=['GET'])
def get_models():
    """获取所有模型配置列表"""
    try:
        manager = get_model_config_manager()
        models = manager.get_config_list()
        return jsonify({'success': True, 'models': models})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/models/enabled', methods=['GET'])
def get_enabled_models_list():
    """获取启用的模型列表"""
    try:
        models = get_enabled_models()
        return jsonify({'success': True, 'models': models})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/models/<model_id>', methods=['GET'])
def get_model_detail(model_id):
    """获取模型配置详情"""
    try:
        manager = get_model_config_manager()
        config = manager.get_config(model_id)
        if config:
            # 返回配置，但隐藏 API Key
            data = config.to_dict()
            data['api_key'] = '******' if data['api_key'] else ''
            return jsonify({'success': True, 'model': data})
        else:
            return jsonify({'success': False, 'error': '模型不存在'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/models/<model_id>/apikey', methods=['POST'])
def update_model_api_key(model_id):
    """更新模型 API Key"""
    try:
        data = request.get_json()
        api_key = data.get('api_key', '')
        
        manager = get_model_config_manager()
        config = manager.get_config(model_id)
        if not config:
            return jsonify({'success': False, 'error': '模型不存在'}), 404
        
        # 更新 API Key
        success = manager.update_config(model_id, api_key=api_key)
        if success:
            return jsonify({'success': True, 'message': 'API Key 更新成功'})
        else:
            return jsonify({'success': False, 'error': '更新失败'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/models/<model_id>/toggle', methods=['POST'])
def toggle_model(model_id):
    """启用/禁用模型"""
    try:
        data = request.get_json()
        enabled = data.get('enabled', True)
        
        manager = get_model_config_manager()
        config = manager.get_config(model_id)
        if not config:
            return jsonify({'success': False, 'error': '模型不存在'}), 404
        
        success = manager.update_config(model_id, enabled=enabled)
        if success:
            return jsonify({'success': True, 'message': f'模型已{"启用" if enabled else "禁用"}'})
        else:
            return jsonify({'success': False, 'error': '更新失败'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/models/<model_id>/test', methods=['POST'])
def test_model(model_id):
    """测试模型配置"""
    try:
        manager = get_model_config_manager()
        config = manager.get_config(model_id)
        if not config:
            return jsonify({'success': False, 'error': '模型不存在'}), 404

        # Ollama 本地模型不需要 API Key
        if config.provider_id != 'ollama' and not config.api_key:
            return jsonify({'success': False, 'error': 'API Key 未配置'}), 400
        
        # 测试调用
        from ai_engine import call_ai_api
        test_prompt = "你好，请用一句话介绍你自己。"
        result = call_ai_api(test_prompt, config)
        
        if result.startswith('[') and ('API' in result or 'Error' in result or 'error' in result):
            return jsonify({'success': False, 'error': result}), 400
        else:
            return jsonify({'success': True, 'message': '连接成功', 'response': result[:100]})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/models/custom', methods=['POST'])
def add_custom_model():
    """添加自定义模型配置"""
    try:
        data = request.get_json()
        
        # 创建新配置
        config = ModelConfig(
            id=data.get('id'),
            name=data.get('name'),
            provider=data.get('provider', 'custom'),
            model=data.get('model'),
            api_key=data.get('api_key', ''),
            base_url=data.get('base_url'),
            max_tokens=data.get('max_tokens', 2000),
            temperature=data.get('temperature', 0.7),
            timeout=data.get('timeout', 120),
            enabled=True,
            description=data.get('description', ''),
            request_format=data.get('request_format', 'openai'),
            response_path=data.get('response_path', 'choices.0.message.content')
        )
        
        manager = get_model_config_manager()
        success = manager.add_config(config)
        
        if success:
            return jsonify({'success': True, 'message': '模型配置添加成功'})
        else:
            return jsonify({'success': False, 'error': '添加失败'}), 500
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 模型配置管理 API v2 ====================

@app.route('/model-config-v2')
def model_config_v2_page():
    """模型配置管理 v2 页面"""
    return render_template('model_config_v2.html')


@app.route('/api/v2/models/providers', methods=['GET'])
def get_providers_v2():
    """获取所有厂商列表"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        providers = manager.get_all_providers()
        
        result = [
            {
                'id': p.id,
                'name': p.name,
                'icon': p.icon,
                'base_url': p.base_url,
                'is_custom': p.is_custom,
                'model_count': len(p.models)
            }
            for p in providers.values()
        ]
        
        return jsonify({'success': True, 'providers': result})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/providers/<provider_id>', methods=['GET'])
def get_provider_models_v2(provider_id):
    """获取指定厂商的所有模型"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        models = manager.get_models_by_provider(provider_id)
        return jsonify({'success': True, 'models': models})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/providers', methods=['POST'])
def add_provider_v2():
    """添加自定义厂商"""
    try:
        from model_config_v2 import get_model_config_manager_v2, ProviderConfig
        manager = get_model_config_manager_v2()
        
        data = request.get_json()
        provider = ProviderConfig(
            id=data['id'],
            name=data['name'],
            icon=data.get('icon', '🔧'),
            base_url=data['base_url'],
            is_custom=True
        )
        
        success = manager.add_custom_provider(provider)
        if success:
            return jsonify({'success': True, 'message': '厂商添加成功'})
        else:
            return jsonify({'success': False, 'message': '厂商 ID 已存在'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models', methods=['POST'])
def add_model_v2():
    """添加自定义模型"""
    try:
        from model_config_v2 import get_model_config_manager_v2, ModelConfig
        manager = get_model_config_manager_v2()
        
        data = request.get_json()
        model = ModelConfig(
            id=data['id'],
            provider_id=data['provider_id'],
            name=data['name'],
            type=data['type'],
            model=data['model'],
            api_key=data.get('api_key', ''),
            base_url=data.get('base_url', ''),
            max_tokens=data.get('max_tokens', 2000),
            temperature=data.get('temperature', 0.7),
            description=data.get('description', ''),
            is_custom=True
        )
        
        success = manager.add_custom_model(data['provider_id'], model)
        if success:
            return jsonify({'success': True, 'message': '模型添加成功'})
        else:
            return jsonify({'success': False, 'message': '厂商不存在'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/<model_id>', methods=['GET'])
def get_model_v2(model_id):
    """获取模型配置详情"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        model = manager.get_model(model_id)
        
        if not model:
            return jsonify({'success': False, 'message': '模型不存在'}), 404
        
        return jsonify({'success': True, 'model': model.to_dict()})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/<model_id>', methods=['PUT'])
def update_model_v2(model_id):
    """更新模型配置"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        
        data = request.get_json()
        success = manager.update_model_config(model_id, **data)
        
        if success:
            return jsonify({'success': True, 'message': '模型更新成功'})
        else:
            return jsonify({'success': False, 'message': '模型不存在'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/<model_id>', methods=['DELETE'])
def delete_model_v2(model_id):
    """删除自定义模型"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        success = manager.delete_custom_model(model_id)
        
        if success:
            return jsonify({'success': True, 'message': '模型删除成功'})
        else:
            return jsonify({'success': False, 'message': '模型不存在或不是自定义模型'}), 400
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/<model_id>/toggle', methods=['POST'])
def toggle_model_v2(model_id):
    """切换模型启用状态"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        
        data = request.get_json()
        enabled = data.get('enabled', True)
        success = manager.update_model_config(model_id, enabled=enabled)
        
        if success:
            return jsonify({'success': True, 'message': f'模型已{"启用" if enabled else "禁用"}'})
        else:
            return jsonify({'success': False, 'message': '模型不存在'}), 404
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/models/<model_id>/test', methods=['POST'])
def test_model_v2(model_id):
    """测试模型配置"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        from ai_engine import call_ai_api
        from model_config_v2 import ModelConfig
        
        manager = get_model_config_manager_v2()
        model = manager.get_model(model_id)
        
        if not model:
            return jsonify({'success': False, 'message': '模型不存在'}), 404

        # Ollama 本地模型不需要 API Key
        if model.provider_id != 'ollama' and not model.api_key:
            return jsonify({
                'success': False,
                'message': 'API Key 未配置，请先编辑模型配置 API Key'
            }), 400
        
        # 检查启用状态
        if not model.enabled:
            return jsonify({
                'success': False, 
                'message': '模型未启用，请先启用模型'
            }), 400
        
        # 构建测试 Prompt
        test_prompt = "请用一句话介绍你自己。"

        # 如果 base_url 为空，使用厂商的 base_url
        base_url = model.base_url
        if not base_url:
            provider = manager.get_provider(model.provider_id)
            if provider:
                base_url = provider.base_url
                print(f'[DEBUG] 模型 {model.id} 的 base_url 为空，使用厂商默认值：{base_url}')

        # 创建临时模型配置用于测试
        # GLM-4.7-Flash 等模型会输出 reasoning_content，需要更多 token
        test_max_tokens = 300 if "glm-4.7" in model.model.lower() or "glm-4-7" in model.model.lower() else 100
        temp_config = ModelConfig(
            id=model.id,
            provider_id=model.provider_id,
            name=model.name,
            type=model.type,
            model=model.model,
            api_key=model.api_key,
            base_url=base_url,
            max_tokens=test_max_tokens,  # 测试时限制输出长度
            temperature=model.temperature,
            timeout=model.timeout,
            enabled=model.enabled,
            request_format=model.request_format,
            response_path=model.response_path
        )
        
        # 调用 AI API
        import time
        start_time = time.time()
        result = call_ai_api(test_prompt, temp_config)
        response_time = round((time.time() - start_time) * 1000)  # 毫秒
        
        # 检查结果
        if result.startswith('[') and ('API' in result or 'Error' in result or 'error' in result):
            return jsonify({
                'success': False,
                'message': f'测试失败：{result}',
                'response_time': response_time
            }), 400
        
        # 获取厂商信息
        provider = manager.get_provider(model.provider_id)
        
        # 返回测试结果和模型信息
        return jsonify({
            'success': True,
            'message': '模型测试成功',
            'model_info': {
                'id': model.id,
                'name': model.name,
                'provider_name': provider.name if provider else model.provider_id,
                'provider_icon': provider.icon if provider else '🔧',
                'model': model.model,
                'type': model.type,
                'max_tokens': model.max_tokens,
                'temperature': model.temperature,
                'base_url': model.base_url,
                'enabled': model.enabled,
                'has_api_key': bool(model.api_key)
            },
            'test_result': {
                'response': result[:200],  # 返回前 200 字符
                'response_length': len(result),
                'response_time': response_time,
                'tokens_used': '约 50-100'  # 估算
            }
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'测试失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/v2/models/enabled', methods=['GET'])
def get_enabled_models_v2():
    """获取启用的模型列表（用于首页选择）"""
    try:
        from model_config_v2 import get_model_config_manager_v2
        manager = get_model_config_manager_v2()
        models = manager.get_enabled_models()
        return jsonify({'success': True, 'models': models})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


# ==================== 新版任务管理 API（支持持久化和实时预览） ====================

@app.route('/api/v2/task/submit', methods=['POST'])
def submit_task_v2():
    """
    提交文档生成任务（新版，支持持久化和实时预览）
    
    请求参数:
    - template_type: 模板类型
    - requirement_file: 需求文档（File）
    - template_file: 模板文档（File，可选）
    - model: 模型名称
    - user_prompt: 用户补充要求
    """
    try:
        from task_manager import get_task_manager
        
        # 获取参数
        template_type = request.form.get('template_type', 'future_community')
        user_prompt = request.form.get('user_prompt', '')
        model = request.form.get('model', 'qwen-max')
        
        # 获取模型配置
        model_config = get_model_config(model)
        if not model_config:
            return jsonify({'success': False, 'message': f'模型 {model} 不存在'}), 400
        
        if not model_config.enabled:
            return jsonify({'success': False, 'message': f'模型 {model_config.name} 已禁用'}), 400

        # Ollama 本地模型不需要 API Key
        if model_config.provider_id != 'ollama' and not model_config.api_key:
            return jsonify({'success': False, 'message': f'模型 {model_config.name} 的 API Key 未配置'}), 400
        
        # 创建任务
        task_manager = get_task_manager()
        task_id = task_manager.create_task(
            template_type=template_type,
            user_prompt=user_prompt,
            model=model
        )
        
        # 处理上传的文件
        template_file = request.files.get('template_file')
        requirement_file = request.files.get('requirement_file')
        
        requirement_filename = None
        template_filename = None
        
        # 保存需求文件
        if requirement_file and allowed_file(requirement_file.filename):
            original_filename = secure_filename(requirement_file.filename)
            if '.' not in original_filename:
                ext = requirement_file.filename.rsplit('.', 1)[1].lower()
                original_filename = f'{original_filename}.{ext}'
            
            # 保存到任务目录
            task_manager.save_file(task_id, f'requirement.{original_filename.rsplit(".", 1)[-1]}', 
                                   requirement_file.read())
            requirement_filename = original_filename
        
        # 保存模板文件
        if template_file and allowed_file(template_file.filename):
            original_filename = secure_filename(template_file.filename)
            if '.' not in original_filename:
                ext = template_file.filename.rsplit('.', 1)[1].lower()
                original_filename = f'{original_filename}.{ext}'
            
            task_manager.save_file(task_id, f'template.{original_filename.rsplit(".", 1)[-1]}', 
                                   template_file.read())
            template_filename = original_filename
        
        # 更新任务信息
        task_manager.update_task_status(
            task_id, 
            requirement_filename=requirement_filename,
            template_filename=template_filename
        )
        
        # 读取文件内容
        requirement_content = ''
        template_content = ''
        
        if requirement_filename:
            req_path = task_manager.get_file_path(task_id, f'requirement.{requirement_filename.rsplit(".", 1)[-1]}')
            if req_path:
                if requirement_filename.endswith('.docx'):
                    requirement_content = read_docx_text(req_path)
                elif requirement_filename.endswith('.txt'):
                    requirement_content = read_txt_text(req_path)
        
        if template_filename:
            tmpl_path = task_manager.get_file_path(task_id, f'template.{template_filename.rsplit(".", 1)[-1]}')
            if tmpl_path:
                if template_filename.endswith('.docx'):
                    template_content = read_docx_text(tmpl_path)
                elif template_filename.endswith('.txt'):
                    template_content = read_txt_text(tmpl_path)
        
        # 生成输出路径
        output_path = os.path.join(app.config['OUTPUT_FOLDER'], f'{template_type}_{task_id}.docx')
        
        # 启动异步线程
        thread = threading.Thread(
            target=process_document_async_v3,
            args=(task_id, template_type, requirement_content, template_content,
                  user_prompt, model_config, output_path)
        )
        thread.daemon = True
        thread.start()
        
        return jsonify({
            'success': True,
            'message': '任务已提交',
            'task_id': task_id
        })
        
    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'提交失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/v2/task/<task_id>/status', methods=['GET'])
def get_task_status_v2(task_id):
    """
    获取任务状态（轮询接口）

    返回:
    {
        "success": true,
        "task": {
            "task_id": "...",
            "status": "generating",
            "progress": 45,
            "current_chapter": "第三章",
            "total_chapters": 15,
            "completed_chapters": 6,
            "chapters": [...],
            "partial_doc_url": "/api/v2/task/abc123/download-partial"
        }
    }
    """
    try:
        from task_manager import get_task_manager

        task_manager = get_task_manager()
        status = task_manager.get_task_status(task_id)

        if not status:
            return jsonify({'success': False, 'message': '任务不存在'}), 404

        return jsonify({'success': True, 'task': status})

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/tasks', methods=['GET'])
def list_tasks_v2():
    """
    获取所有任务列表

    查询参数:
    - status: 筛选状态 (pending/processing/completed/failed)
    - limit: 返回数量限制 (默认 50)
    - offset: 分页偏移 (默认 0)

    返回:
    {
        "success": true,
        "tasks": [
            {
                "task_id": "...",
                "status": "completed",
                "progress": 100,
                "template_type": "future_community",
                "model": "qwen-max",
                "created_at": "2026-03-10 10:00:00",
                "updated_at": "2026-03-10 10:05:00"
            }
        ],
        "total": 10
    }
    """
    try:
        from task_manager import get_task_manager, TASKS_ROOT
        import os

        task_manager = get_task_manager()
        status_filter = request.args.get('status')
        limit = int(request.args.get('limit', 50))
        offset = int(request.args.get('offset', 0))

        tasks = []

        # 遍历任务目录
        if os.path.exists(TASKS_ROOT):
            for task_id in os.listdir(TASKS_ROOT):
                task_dir = os.path.join(TASKS_ROOT, task_id)
                if not os.path.isdir(task_dir):
                    continue

                task_info = task_manager.load_task_info(task_id)
                if task_info:
                    # 应用状态筛选
                    if status_filter and task_info.status != status_filter:
                        continue

                    tasks.append({
                        'task_id': task_info.task_id,
                        'status': task_info.status,
                        'progress': task_info.progress,
                        'template_type': task_info.template_type,
                        'model': task_info.model,
                        'created_at': task_info.created_at,
                        'updated_at': task_info.updated_at
                    })

        # 按创建时间排序（最新的在前）
        tasks.sort(key=lambda x: x['created_at'], reverse=True)

        total = len(tasks)
        tasks = tasks[offset:offset + limit]

        return jsonify({
            'success': True,
            'tasks': tasks,
            'total': total
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/pause', methods=['POST'])
def pause_task_v2(task_id):
    """暂停任务"""
    try:
        from task_manager import get_task_manager
        task_manager = get_task_manager()
        task_manager.pause_task(task_id)
        return jsonify({'success': True, 'message': '任务已暂停'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/continue', methods=['POST'])
def continue_task_v2(task_id):
    """继续任务"""
    try:
        from task_manager import get_task_manager
        task_manager = get_task_manager()
        task_manager.continue_task(task_id)
        return jsonify({'success': True, 'message': '任务已继续'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/cancel', methods=['POST'])
def cancel_task_v2(task_id):
    """取消任务"""
    try:
        from task_manager import get_task_manager
        task_manager = get_task_manager()
        task_manager.cancel_task(task_id)
        return jsonify({'success': True, 'message': '任务已取消'})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/download-partial', methods=['GET'])
def download_partial_v2(task_id):
    """下载临时文档"""
    try:
        from task_manager import get_task_manager
        task_manager = get_task_manager()
        
        filepath = task_manager.get_partial_document_path(task_id)
        if not filepath:
            return jsonify({'success': False, 'message': '临时文档不存在'}), 404
        
        return send_file(
            filepath,
            as_attachment=True,
            download_name=f'partial_{task_id[:8]}.docx'
        )
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/download', methods=['GET'])
def download_final_v2(task_id):
    """下载最终文档"""
    try:
        from task_manager import get_task_manager
        task_manager = get_task_manager()
        
        # 查找最终文档
        task_dir = task_manager._get_task_directory(task_id)
        final_file = os.path.join(task_dir, f'final_{task_id}.docx')
        
        if not os.path.exists(final_file):
            # 如果最终文档不存在，返回临时文档
            partial_file = os.path.join(task_dir, f'partial_{task_id}.docx')
            if os.path.exists(partial_file):
                return send_file(partial_file, as_attachment=True, download_name=f'partial_{task_id[:8]}.docx')
            return jsonify({'success': False, 'message': '文档尚未生成完成'}), 404
        
        return send_file(final_file, as_attachment=True, download_name=f'final_{task_id[:8]}.docx')
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/chapters/<int:chapter_index>/regenerate', methods=['POST'])
def regenerate_chapter_v2(task_id, chapter_index):
    """
    重新生成章节（带用户补充需求）

    请求体:
    {
        "user_instruction": "请结合回迁小区的特点，重点说明...",
        "based_on_content": "用户编辑的内容（可选）"
    }
    """
    try:
        from task_manager import get_task_manager, MAX_REGENERATE_COUNT

        task_manager = get_task_manager()

        # 检查是否可以重新生成
        if not task_manager.can_regenerate(task_id, chapter_index):
            return jsonify({
                'success': False,
                'message': f'该章节已重新生成{MAX_REGENERATE_COUNT}次，达到上限'
            }), 400

        data = request.get_json()
        user_instruction = data.get('user_instruction', '')
        based_on_content = data.get('based_on_content', '')  # 用户编辑的内容

        # 获取任务信息
        task_info = task_manager.load_task_info(task_id)
        chapters = task_manager.load_chapters(task_id)

        if not task_info or not chapters:
            return jsonify({'success': False, 'message': '任务不存在'}), 404

        if chapter_index >= len(chapters):
            return jsonify({'success': False, 'message': '章节不存在'}), 404

        chapter = chapters[chapter_index]

        # 读取文件内容
        requirement_content = ''
        template_content = ''

        if task_info.requirement_filename:
            req_path = task_manager.get_file_path(task_id, f'requirement.{task_info.requirement_filename.rsplit(".", 1)[-1]}')
            if req_path and task_info.requirement_filename.endswith('.docx'):
                requirement_content = read_docx_text(req_path)
            elif req_path and task_info.requirement_filename.endswith('.txt'):
                requirement_content = read_txt_text(req_path)

        if task_info.template_filename:
            tmpl_path = task_manager.get_file_path(task_id, f'template.{task_info.template_filename.rsplit(".", 1)[-1]}')
            if tmpl_path and task_info.template_filename.endswith('.docx'):
                template_content = read_docx_text(tmpl_path)
            elif tmpl_path and task_info.template_filename.endswith('.txt'):
                template_content = read_txt_text(tmpl_path)

        # 获取模型配置
        model_config = get_model_config(task_info.model)
        if not model_config:
            return jsonify({'success': False, 'message': '模型配置不存在'}), 404

        # 异步重新生成
        def regenerate_async():
            from ai_engine import generate_section_content_with_ai

            # 更新章节状态
            task_manager.update_chapter_status(task_id, chapter_index, status='generating')

            try:
                # 提取模板章节内容
                from ai_engine import extract_template_section
                template_section = extract_template_section(chapter.title, template_content)

                # 构建提示词：如果有 based_on_content，则基于它进行优化
                final_instruction = user_instruction
                if based_on_content:
                    final_instruction = f"""请基于以下用户编辑的内容进行优化和扩展，保持核心信息不变，但提升表达质量：

用户编辑的内容：
{based_on_content}

用户额外要求：
{user_instruction}

请生成优化后的内容："""

                # AI 调用前检查取消状态
                task_info_check = task_manager.load_task_info(task_id)
                if task_info_check and task_info_check.status == 'cancelled':
                    print(f'[CANCEL] 任务已取消，停止生成')
                    return

                # 生成内容
                content = generate_section_content_with_ai(
                    section_title=chapter.title,
                    requirement_text=requirement_content,
                    template_text=template_section,
                    user_instruction=final_instruction,
                    model_config=model_config,
                    fallback_to_template=False
                )

                # AI 返回后检查取消状态
                task_info_check = task_manager.load_task_info(task_id)
                if task_info_check and task_info_check.status == 'cancelled':
                    print(f'[CANCEL] 任务已取消，停止生成')
                    return

                if content.startswith('[AI 生成失败]'):
                    raise Exception(content)

                # 更新章节状态
                task_manager.update_chapter_status(
                    task_id, chapter_index,
                    status='completed',
                    content=content,
                    word_count=len(content),
                    regenerated_count=chapter.regenerated_count + 1,
                    last_user_instruction=user_instruction
                )

                # 更新临时文档
                try:
                    chapters_updated = task_manager.load_chapters(task_id)
                    if chapters_updated:
                        from ai_engine import create_partial_document
                        chapter_dicts = [ch.to_dict() for ch in chapters_updated]
                        create_partial_document(task_id, chapter_dicts, task_info.template_type if task_info else 'future_community')
                except Exception as e:
                    print(f'[WARN] 更新临时文档失败：{e}')

            except Exception as e:
                task_manager.update_chapter_status(
                    task_id, chapter_index,
                    status='failed',
                    error_message=str(e)
                )

        thread = threading.Thread(target=regenerate_async)
        thread.daemon = True
        thread.start()

        return jsonify({
            'success': True,
            'message': '正在重新生成章节'
        })

    except Exception as e:
        import traceback
        return jsonify({
            'success': False,
            'message': f'请求失败：{str(e)}',
            'detail': traceback.format_exc()
        }), 500


@app.route('/api/v2/task/<task_id>/chapters/<int:chapter_index>/content', methods=['PUT'])
def update_chapter_content_v2(task_id, chapter_index):
    """
    更新章节内容（手动编辑）

    请求体:
    {
        "content": "编辑后的章节内容..."
    }
    """
    try:
        from task_manager import get_task_manager

        task_manager = get_task_manager()
        data = request.get_json()
        new_content = data.get('content', '')

        if not new_content.strip():
            return jsonify({'success': False, 'message': '内容不能为空'}), 400

        # 更新章节内容
        task_manager.update_chapter_status(
            task_id,
            chapter_index,
            content=new_content,
            status='completed',  # 确保状态为已完成
            error_message=None   # 清除错误信息
        )

        # 更新部分文档
        try:
            chapters = task_manager.load_chapters(task_id)
            if chapters:
                from ai_engine import create_partial_document
                task_info = task_manager.load_task_info(task_id)
                # 将章节对象转换为字典列表
                chapter_dicts = [ch.to_dict() for ch in chapters]
                create_partial_document(task_id, chapter_dicts, task_info.template_type if task_info else 'future_community')
        except Exception as e:
            print(f'[WARN] 更新部分文档失败：{e}')
            import traceback
            print(f'[WARN] 详细错误：{traceback.format_exc()}')

        return jsonify({
            'success': True,
            'message': '章节内容已保存'
        })

    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


@app.route('/api/v2/task/<task_id>/chapters/retry-failed', methods=['POST'])
def retry_failed_chapters_v2(task_id):
    """
    批量重试失败章节
    
    请求体:
    {
        "chapter_indices": [7, 9, 11],  // 空则重试所有失败
        "user_instruction": "请使用更保守的估算方法..."  // 可选
    }
    """
    try:
        from task_manager import get_task_manager
        
        task_manager = get_task_manager()
        data = request.get_json()
        chapter_indices = data.get('chapter_indices', [])
        user_instruction = data.get('user_instruction', '')
        
        # 获取失败章节
        failed_chapters = task_manager.get_failed_chapters(task_id)
        
        if not failed_chapters:
            return jsonify({'success': True, 'message': '没有失败的章节'})
        
        # 如果没有指定章节，重试所有失败章节
        if not chapter_indices:
            chapter_indices = [ch.index for ch in failed_chapters]
        
        # 逐个重试（简化实现，实际应该用队列管理）
        for idx in chapter_indices:
            # 类似 regenerate_chapter_v2 的逻辑
            # TODO: 实现批量重试逻辑
            pass
        
        return jsonify({
            'success': True,
            'message': f'开始重试 {len(chapter_indices)} 个章节'
        })
        
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500


if __name__ == '__main__':
    print("=" * 50)
    print("  未来社区建设方案生成系统 (AI 增强版)")
    print("=" * 50)
    print()
    print("  访问地址：http://localhost:5000")
    print("  访问地址：http://127.0.0.1:5000")
    print()
    print("  按 Ctrl+C 停止服务")
    print()
    app.run(debug=True, host='0.0.0.0', port=5000)
