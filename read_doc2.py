# -*- coding: utf-8 -*-
from docx import Document
import sys

doc = Document('partial_5c456593.docx')

print(f'段落数量: {len(doc.paragraphs)}')
print()

# 查找有实际内容的段落（不是标题，长度大于30）
content_paras = []
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if text and len(text) > 30:
        content_paras.append((i+1, text))

print(f'=== 有实际内容的段落（共{len(content_paras)}段）===')
for idx, (i, text) in enumerate(content_paras[:20]):
    print(f'{idx+1}. [第{i}段] {text[:200]}')
    print()
