# -*- coding: utf-8 -*-
"""测试修复后的文档生成 - 直接复制模板"""
import sys
sys.path.insert(0, 'backend')

from app.services.template_analyzer import TemplateAnalyzer
from app.utils.doc_builder import DocBuilder
import os

# 模板文件路径
template_path = r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx'
output_path = r'e:\Qwen\xmjl\outputs\test_fixed_template_v2.docx'

print('=' * 80)
print('步骤 1: 分析模板文档')
print('=' * 80)

analyzer = TemplateAnalyzer(template_path)
print(f'\n章节树节点数：{len(analyzer.chapter_tree)}')
print(f'样式定义数：{len(analyzer.styles)}')

# 获取模板结构
template_structure = {
    'chapter_tree': analyzer.get_chapter_tree()
}

print(f'\n前 3 个一级章节:')
for i, chapter in enumerate(template_structure['chapter_tree'][:3]):
    print(f'  {i+1}. {chapter["title"]} (样式：{chapter["style"]})')
    for section in chapter.get('children', [])[:2]:
        print(f'      - {section["title"]} (样式：{section["style"]})')

print('\n' + '=' * 80)
print('步骤 2: 使用 DocBuilder 生成文档（直接复制模板）')
print('=' * 80)

# 创建 DocBuilder（使用模板路径，会直接复制模板文档）
builder = DocBuilder(template_path=template_path)

# 渲染文档（不传入内容，只保留模板结构）
builder.render_from_template_structure(template_structure, {})

# 保存文档
print(f'\n保存到：{output_path}')
os.makedirs(os.path.dirname(output_path), exist_ok=True)
builder.save(output_path)

print('\n文档生成成功！')

# 验证生成的文档
print('\n' + '=' * 80)
print('步骤 3: 验证生成的文档')
print('=' * 80)

from docx import Document
doc = Document(output_path)

styles_count = {}
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    styles_count[style_name] = styles_count.get(style_name, 0) + 1

print(f'\n生成文档总段落数：{len(doc.paragraphs)}')
print(f'\n样式分布（前 10 个）:')
for style, count in sorted(styles_count.items(), key=lambda x: -x[1])[:10]:
    print(f'  {style}: {count}')

# 对比原始模板
print('\n' + '=' * 80)
print('对比原始模板')
print('=' * 80)

orig_doc = Document(template_path)
orig_styles = {}
for para in orig_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    orig_styles[style_name] = orig_styles.get(style_name, 0) + 1

print(f'\n原始模板样式：{len(orig_styles)} 种')
print(f'生成文档样式：{len(styles_count)} 种')

# 检查关键样式是否保留
key_styles = ['Body Text First Indent 2', 'Heading 1', 'Heading 2', 'BZ_正文']
all_passed = True
for style in key_styles:
    orig_count = orig_styles.get(style, 0)
    gen_count = styles_count.get(style, 0)
    status = '✓' if gen_count > 0 else '✗'
    if gen_count == 0:
        all_passed = False
    print(f'  {status} {style}: 原始={orig_count}, 生成={gen_count}')

if all_passed:
    print('\n✓ 所有关键样式都已保留！')
else:
    print('\n✗ 部分样式丢失，需要进一步修复。')

print(f'\n文档已保存到：{output_path}')
