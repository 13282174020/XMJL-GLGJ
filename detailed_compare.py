# -*- coding: utf-8 -*-
"""详细对比模板文档和生成文档的章节结构"""
import sys
sys.path.insert(0, 'backend')
sys.path.insert(0, 'web_system')

from docx import Document
import re

def analyze_doc_structure(doc_path):
    """分析文档结构"""
    print(f'\n{"="*80}')
    print(f'分析文档：{doc_path}')
    print(f'{"="*80}')
    
    doc = Document(doc_path)
    
    # 统计信息
    total_paras = len(doc.paragraphs)
    styles_count = {}
    headings = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else 'Normal'
        styles_count[style_name] = styles_count.get(style_name, 0) + 1
        
        # 检测是否是标题或编号
        level = detect_heading_level(text, style_name)
        if level > 0:
            headings.append({
                'index': i,
                'level': level,
                'style': style_name,
                'text': text[:60]
            })
    
    print(f'\n总段落数：{total_paras}')
    print(f'\n样式分布（前 10 个）:')
    for style, count in sorted(styles_count.items(), key=lambda x: -x[1])[:10]:
        print(f'  {style}: {count}')
    
    print(f'\n识别到的标题数：{len(headings)}')
    print(f'\n标题列表（前 30 个）:')
    for h in headings[:30]:
        indent = '  ' * (h['level'] - 1)
        print(f"  {h['index']:4d}. L{h['level']} [{h['style']}] {indent}{h['text']}")
    
    if len(headings) > 30:
        print(f'  ... 还有 {len(headings) - 30} 个标题')
    
    return headings, styles_count

def detect_heading_level(text, style_name):
    """检测标题层级"""
    style_lower = style_name.lower()
    
    # 检查样式名
    if 'heading 1' in style_lower or '标题 1' in style_lower:
        return 1
    elif 'heading 2' in style_lower or '标题 2' in style_lower:
        return 2
    elif 'heading 3' in style_lower or '标题 3' in style_lower:
        return 3
    elif 'heading 4' in style_lower or '标题 4' in style_lower:
        return 4
    
    # 检查编号格式
    if re.match(r'^第 [一二三四五六七八九十\d]+ 章', text):
        return 1
    if re.match(r'^\d+\.\d+$', text) or re.match(r'^\d+\.\d+\s', text):
        return 2
    if re.match(r'^\d+\.\d+\.\d+$', text) or re.match(r'^\d+\.\d+\.\d+\s', text):
        return 3
    if re.match(r'^\d+\.\d+\.\d+\.\d+$', text) or re.match(r'^\d+\.\d+\.\d+\.\d+\s', text):
        return 4
    
    return 0

# 分析模板文档
template_headings, template_styles = analyze_doc_structure('03-临平区数字慈善系统建设方案模板.docx')

# 尝试查找生成的文档
import glob
generated_files = glob.glob('**/*新模板*.docx', recursive=True)
if not generated_files:
    generated_files = glob.glob('**/template_*.docx', recursive=True)
if not generated_files:
    generated_files = glob.glob('outputs/*.docx', recursive=True)

if generated_files:
    print(f'\n\n找到生成的文档：{generated_files[0]}')
    generated_headings, generated_styles = analyze_doc_structure(generated_files[0])
    
    # 对比
    print(f'\n{"="*80}')
    print('对比分析')
    print(f'{"="*80}')
    print(f'模板文档标题数：{len(template_headings)}')
    print(f'生成文档标题数：{len(generated_headings)}')
    print(f'差异：{len(template_headings) - len(generated_headings)}')
else:
    print('\n\n未找到生成的文档！')
    print('请确认生成的文档路径，或手动指定路径进行分析。')
