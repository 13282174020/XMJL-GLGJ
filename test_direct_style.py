# -*- coding: utf-8 -*-
"""测试直接设置样式"""
from docx import Document
import shutil
import tempfile
import os

template_path = r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx'
output_path = r'e:\Qwen\xmjl\outputs\test_direct_style.docx'

# 复制模板并清除内容
temp_dir = tempfile.mkdtemp()
temp_template = os.path.join(temp_dir, 'template.docx')
shutil.copy2(template_path, temp_template)

template_doc = Document(temp_template)

# 清除段落
for para in list(template_doc.paragraphs):
    para._element.getparent().remove(para._element)

# 添加测试段落
print('添加段落使用 add_paragraph(style="Heading 1")...')
para1 = template_doc.add_paragraph(style='Heading 1')
para1.add_run('测试标题 1')

print('添加段落使用 add_paragraph() 后设置样式...')
para2 = template_doc.add_paragraph()
para2.style = 'Heading 1'
para2.add_run('测试标题 2')

print('添加段落使用 add_paragraph()...')
para3 = template_doc.add_paragraph()
para3.add_run('测试标题 3 (Normal)')

# 保存
template_doc.save(output_path)

# 清理
os.remove(temp_template)
os.rmdir(temp_dir)

# 验证
print(f'\n保存到：{output_path}')
print('\n验证生成的文档:')
doc = Document(output_path)
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    print(f'  段落 {i}: 样式={para.style.name}, 文本={text}')
