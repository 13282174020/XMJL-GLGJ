from docx import Document
import os

os.makedirs('tests/fixtures', exist_ok=True)

# 标准样本
doc = Document()
doc.add_heading('1 项目概况', level=1)
doc.add_heading('1.1 项目背景', level=2)
doc.add_heading('1.1.1 政策背景', level=3)
doc.add_heading('1.1.1.1 详细政策', level=4)
doc.add_heading('2 需求分析', level=1)
doc.add_heading('2.1 业务需求', level=2)
doc.add_heading('2.1.1 用户需求', level=3)
doc.save('tests/fixtures/sample_standard.docx')
print('OK: sample_standard.docx')

# 非标准样本
doc2 = Document()
doc2.add_heading('1 项目概述', level=1)
doc2.add_heading('2 现状分析', level=1)
p = doc2.add_paragraph('2.1 建设背景', style='Normal')
p.runs[0].font.bold = True
p = doc2.add_paragraph('2.2 现有系统不足', style='Normal')
p.runs[0].font.bold = True
p = doc2.add_paragraph('2.2.1 操作复杂', style='Normal')
p.runs[0].font.bold = True
p = doc2.add_paragraph('2.2.2 功能单一', style='Normal')
p.runs[0].font.bold = True
p = doc2.add_paragraph('2.2.2.1 缺乏移动端', style='Normal')
p.runs[0].font.bold = True
p = doc2.add_paragraph('2.2.2.2 数据分析弱', style='Normal')
p.runs[0].font.bold = True
p = doc2.add_paragraph('2.2.2.2.1 无报表功能', style='Normal')
p.runs[0].font.bold = True
doc2.add_heading('3 建设目标', level=1)
doc2.save('tests/fixtures/sample_nonstandard.docx')
print('OK: sample_nonstandard.docx')
print('Done!')

# 创建测试文件
test_content = '''# -*- coding: utf-8 -*-
import pytest
from web_system.app import get_heading_level, scan_template_styles

def test_heading_1():
    assert get_heading_level('Heading 1') == 1

def test_heading_6():
    assert get_heading_level('Heading 6') == 6

def test_not_heading():
    assert get_heading_level('Normal') == 0

def test_scan_standard(sample_standard_doc):
    result = scan_template_styles(str(sample_standard_doc))
    assert result['success'] is True
    print('OK:', result['message'])

def test_scan_nonstandard(sample_nonstandard_doc):
    result = scan_template_styles(str(sample_nonstandard_doc))
    assert result['success'] is True
    print('OK:', result['message'])
'''

with open('tests/test_document_analyzer.py', 'w', encoding='utf-8') as f:
    f.write(test_content)
print('OK: tests/test_document_analyzer.py')

# 创建 template_analyzer.py 模块
analyzer_content = '''# -*- coding: utf-8 -*-
"""模板文档分析器模块"""
import re
from dataclasses import dataclass, field
from typing import Dict, List, Set
from collections import defaultdict


@dataclass
class NumberingPattern:
    pattern: str
    suggested_level: int
    count: int = 0
    example_texts: List[str] = field(default_factory=list)


@dataclass
class StyleAnalysis:
    style_name: str
    total_count: int = 0
    is_heading: bool = False
    heading_level: int = 0
    numbering_patterns: Dict[str, NumberingPattern] = field(default_factory=dict)


@dataclass
class AnalysisReport:
    file_path: str
    total_paragraphs: int = 0
    styles_usage: Dict[str, int] = field(default_factory=dict)
    heading_styles: Dict[str, StyleAnalysis] = field(default_factory=dict)
    potential_heading_styles: List[str] = field(default_factory=list)
    
    def generate_mapping_rules(self) -> List[dict]:
        rules = []
        for style_name, analysis in self.heading_styles.items():
            if analysis.is_heading and analysis.heading_level > 0:
                rules.append({
                    'source_style': style_name,
                    'pattern': None,
                    'target_style': f'Heading {analysis.heading_level}',
                    'rule_type': 'heading'
                })
        for style_name in self.potential_heading_styles:
            if style_name in self.heading_styles:
                continue
            analysis = self.heading_styles.get(style_name)
            if analysis and analysis.numbering_patterns:
                for pattern_str, pattern_info in analysis.numbering_patterns.items():
                    rules.append({
                        'source_style': style_name,
                        'pattern': pattern_str,
                        'target_style': f'Heading {pattern_info.suggested_level}',
                        'rule_type': 'numbering'
                    })
        return rules


class DocumentAnalyzer:
    def __init__(self):
        self.numbering_regex_patterns = [
            (r'^(\\d+)\\.(\\d+)\\.(\\d+)\\.(\\d+)', 4),
            (r'^(\\d+)\\.(\\d+)\\.(\\d+)', 3),
            (r'^(\\d+)\\.(\\d+)', 2),
            (r'^(\\d+)', 1),
        ]
    
    def analyze(self, file_path: str) -> AnalysisReport:
        from docx import Document
        doc = Document(file_path)
        styles_usage = defaultdict(int)
        heading_styles = {}
        potential_heading_styles = set()
        style_numbering = defaultdict(dict)
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else 'Normal'
            styles_usage[style_name] += 1
            level = self._get_heading_level(style_name, text)
            
            if level > 0:
                if style_name not in heading_styles:
                    heading_styles[style_name] = StyleAnalysis(
                        style_name=style_name, is_heading=True, heading_level=level)
                heading_styles[style_name].total_count += 1
            else:
                for pattern, suggested_level in self.numbering_regex_patterns:
                    if re.match(pattern, text):
                        potential_heading_styles.add(style_name)
                        if pattern not in style_numbering[style_name]:
                            style_numbering[style_name][pattern] = NumberingPattern(
                                pattern=pattern, suggested_level=suggested_level)
                        pattern_info = style_numbering[style_name][pattern]
                        pattern_info.count += 1
                        if len(pattern_info.example_texts) < 3:
                            pattern_info.example_texts.append(text[:50])
                        break
        
        for style_name in potential_heading_styles:
            if style_name not in heading_styles:
                heading_styles[style_name] = StyleAnalysis(
                    style_name=style_name, is_heading=False, heading_level=0,
                    total_count=styles_usage.get(style_name, 0))
            heading_styles[style_name].numbering_patterns = style_numbering.get(style_name, {})
        
        return AnalysisReport(
            file_path=file_path, total_paragraphs=len(doc.paragraphs),
            styles_usage=dict(styles_usage), heading_styles=heading_styles,
            potential_heading_styles=list(potential_heading_styles))
    
    def _get_heading_level(self, style_name: str, text: str = '') -> int:
        if not style_name:
            return 0
        style_lower = style_name.lower()
        if 'heading' in style_lower:
            try:
                num_str = ''.join(filter(str.isdigit, style_lower.replace('heading', '')))
                if num_str:
                    level = int(num_str)
                    if 1 <= level <= 6:
                        return level
            except:
                pass
        if '标题' in style_lower:
            try:
                num_str = ''.join(filter(str.isdigit, style_lower.replace('标题', '')))
                if num_str:
                    level = int(num_str)
                    if 1 <= level <= 6:
                        return level
            except:
                pass
        return 0
'''

with open('web_system/template_analyzer.py', 'w', encoding='utf-8') as f:
    f.write(analyzer_content)
print('OK: web_system/template_analyzer.py')

# 创建模板分析器测试文件
analyzer_test_content = '''# -*- coding: utf-8 -*-
"""模板分析器测试"""
import pytest
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))
sys.path.insert(0, str(Path(__file__).parent.parent / 'web_system'))

from template_analyzer import DocumentAnalyzer

class TestDocumentAnalyzer:
    def test_analyze_standard(self, sample_standard_doc):
        analyzer = DocumentAnalyzer()
        report = analyzer.analyze(str(sample_standard_doc))
        assert report.total_paragraphs > 0
        assert len(report.heading_styles) > 0
        print(f"标准文档：{len(report.heading_styles)} 种样式")
    
    def test_analyze_nonstandard(self, sample_nonstandard_doc):
        analyzer = DocumentAnalyzer()
        report = analyzer.analyze(str(sample_nonstandard_doc))
        assert len(report.potential_heading_styles) > 0
        rules = report.generate_mapping_rules()
        assert len(rules) > 0
        print(f"非标准文档：{len(rules)} 条规则")
    
    def test_get_heading_level(self):
        analyzer = DocumentAnalyzer()
        assert analyzer._get_heading_level('Heading 1') == 1
        assert analyzer._get_heading_level('Normal') == 0

if __name__ == '__main__':
    pytest.main([__file__, '-v'])
'''

with open('tests/test_template_analyzer.py', 'w', encoding='utf-8') as f:
    f.write(analyzer_test_content)
print('OK: tests/test_template_analyzer.py')

# 创建预处理引擎模块
preprocessor_content = '''# -*- coding: utf-8 -*-
"""模板预处理引擎模块"""
import re
from pathlib import Path
from dataclasses import dataclass
from typing import List, Optional, Dict
from docx import Document


@dataclass
class PreprocessResult:
    success: bool
    output_path: Optional[str]
    message: str
    stats: Dict[str, int]
    
    def to_dict(self):
        return {'success': self.success, 'output_path': self.output_path, 
                'message': self.message, 'stats': self.stats}


class TemplatePreprocessor:
    def __init__(self):
        self.standard_styles = ['Heading 1', 'Heading 2', 'Heading 3', 
                                'Heading 4', 'Heading 5', 'Heading 6', 'Normal']
    
    def preprocess(self, input_path: str, output_path: str, 
                   rules: List[dict]) -> PreprocessResult:
        try:
            doc = Document(input_path)
            stats = {}
            
            compiled_rules = []
            for rule in rules:
                compiled_rule = rule.copy()
                if rule.get('pattern'):
                    try:
                        compiled_rule['compiled_pattern'] = re.compile(rule['pattern'])
                    except re.error as e:
                        return PreprocessResult(False, None, f'正则错误：{e}', {})
                compiled_rules.append(compiled_rule)
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                current_style = para.style.name if para.style else 'Normal'
                new_style = self._apply_rules(current_style, text, compiled_rules)
                
                if new_style and new_style != current_style:
                    para.style = new_style
                    key = f'{current_style}->{new_style}'
                    stats[key] = stats.get(key, 0) + 1
                else:
                    stats['unchanged'] = stats.get('unchanged', 0) + 1
            
            output_path = Path(output_path)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(output_path))
            
            total = sum(v for k, v in stats.items() if k != 'unchanged')
            return PreprocessResult(True, str(output_path), 
                                   f'处理完成，{total} 个段落已转换', stats)
        except Exception as e:
            return PreprocessResult(False, None, f'处理失败：{e}', {})
    
    def _apply_rules(self, style_name: str, text: str, rules: List[dict]) -> Optional[str]:
        for rule in rules:
            if rule['source_style'] != style_name:
                continue
            if rule.get('compiled_pattern'):
                if rule['compiled_pattern'].match(text):
                    return rule['target_style']
            elif rule.get('pattern') is None:
                return rule['target_style']
        return None
    
    def preprocess_with_analysis(self, input_path: str, output_path: str) -> PreprocessResult:
        from template_analyzer import DocumentAnalyzer
        analyzer = DocumentAnalyzer()
        report = analyzer.analyze(input_path)
        rules = report.generate_mapping_rules()
        return self.preprocess(input_path, output_path, rules)


def preprocess_template(input_path: str, output_path: str, 
                       rules: List[dict]) -> PreprocessResult:
    preprocessor = TemplatePreprocessor()
    return preprocessor.preprocess(input_path, output_path, rules)
'''

with open('web_system/template_preprocessor.py', 'w', encoding='utf-8') as f:
    f.write(preprocessor_content)
print('OK: web_system/template_preprocessor.py')

# 创建预处理引擎测试文件
preprocessor_test_content = '''# -*- coding: utf-8 -*-
"""预处理引擎测试"""
import pytest
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))
sys.path.insert(0, str(Path(__file__).parent.parent / 'web_system'))

from template_preprocessor import TemplatePreprocessor, preprocess_template


class TestTemplatePreprocessor:
    def test_preprocess_with_rules(self, sample_nonstandard_doc, output_dir):
        analyzer = __import__('template_analyzer', fromlist=['DocumentAnalyzer'])
        analyzer_instance = analyzer.DocumentAnalyzer()
        report = analyzer_instance.analyze(str(sample_nonstandard_doc))
        rules = report.generate_mapping_rules()
        
        preprocessor = TemplatePreprocessor()
        output_path = output_dir / 'processed.docx'
        result = preprocessor.preprocess(str(sample_nonstandard_doc), str(output_path), rules)
        
        assert result.success is True
        assert output_path.exists()
        print(f"预处理成功：{result.message}")
        print(f"统计：{result.stats}")
    
    def test_preprocess_with_analysis(self, sample_nonstandard_doc, output_dir):
        preprocessor = TemplatePreprocessor()
        output_path = output_dir / 'processed_auto.docx'
        result = preprocessor.preprocess_with_analysis(str(sample_nonstandard_doc), str(output_path))
        
        assert result.success is True
        assert output_path.exists()
        print(f"自动预处理成功：{result.message}")
    
    def test_preprocess_standard_doc(self, sample_standard_doc, output_dir):
        preprocessor = TemplatePreprocessor()
        output_path = output_dir / 'processed_standard.docx'
        
        analyzer = __import__('template_analyzer', fromlist=['DocumentAnalyzer'])
        report = analyzer.DocumentAnalyzer().analyze(str(sample_standard_doc))
        rules = report.generate_mapping_rules()
        
        result = preprocessor.preprocess(str(sample_standard_doc), str(output_path), rules)
        
        assert result.success is True
        print(f"标准文档预处理：{result.message}")

if __name__ == '__main__':
    pytest.main([__file__, '-v'])
'''

with open('tests/test_template_preprocessor.py', 'w', encoding='utf-8') as f:
    f.write(preprocessor_test_content)
print('OK: tests/test_template_preprocessor.py')

# 创建模板库管理模块
library_content = '''# -*- coding: utf-8 -*-
"""模板库管理模块"""
import json
import shutil
from pathlib import Path
from datetime import datetime
from typing import List, Optional, Dict
from dataclasses import dataclass, asdict


@dataclass
class TemplateInfo:
    """模板信息"""
    id: str
    name: str
    type: str
    file_path: str
    chapter_structure: Dict
    style_config: Dict
    created_at: str
    is_default: bool = False
    
    def to_dict(self):
        return asdict(self)
    
    @classmethod
    def from_dict(cls, data):
        return cls(**data)


class TemplateLibrary:
    """模板库管理器"""
    
    def __init__(self, base_dir: str = 'templates'):
        self.base_dir = Path(base_dir)
        self.base_dir.mkdir(exist_ok=True)
        self.metadata_file = self.base_dir / 'metadata.json'
        self.templates = self._load_metadata()
    
    def _load_metadata(self) -> Dict[str, TemplateInfo]:
        if self.metadata_file.exists():
            with open(self.metadata_file, 'r', encoding='utf-8') as f:
                data = json.load(f)
                return {k: TemplateInfo.from_dict(v) for k, v in data.items()}
        return {}
    
    def _save_metadata(self):
        with open(self.metadata_file, 'w', encoding='utf-8') as f:
            json.dump({k: v.to_dict() for k, v in self.templates.items()}, 
                     f, ensure_ascii=False, indent=2)
    
    def add_template(self, 
                     name: str,
                     file_path: str,
                     template_type: str,
                     chapter_structure: Dict,
                     style_config: Dict,
                     is_default: bool = False) -> TemplateInfo:
        import uuid
        template_id = f'template_{uuid.uuid4().hex[:8]}'
        
        # 复制文件到模板库
        dest_path = self.base_dir / f'{template_id}.docx'
        shutil.copy2(file_path, dest_path)
        
        info = TemplateInfo(
            id=template_id,
            name=name,
            type=template_type,
            file_path=str(dest_path),
            chapter_structure=chapter_structure,
            style_config=style_config,
            created_at=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
            is_default=is_default
        )
        
        self.templates[template_id] = info
        self._save_metadata()
        return info
    
    def get_all_templates(self) -> List[TemplateInfo]:
        return list(self.templates.values())
    
    def get_template(self, template_id: str) -> Optional[TemplateInfo]:
        return self.templates.get(template_id)
    
    def delete_template(self, template_id: str) -> bool:
        if template_id in self.templates:
            template = self.templates[template_id]
            file_path = Path(template.file_path)
            if file_path.exists():
                file_path.unlink()
            del self.templates[template_id]
            self._save_metadata()
            return True
        return False
    
    def download_template(self, template_id: str, output_path: str) -> bool:
        template = self.get_template(template_id)
        if not template:
            return False
        file_path = Path(template.file_path)
        if not file_path.exists():
            return False
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        shutil.copy2(file_path, output_path)
        return True
'''

with open('web_system/template_library.py', 'w', encoding='utf-8') as f:
    f.write(library_content)
print('OK: web_system/template_library.py')

# 创建模板库测试文件
library_test_content = '''# -*- coding: utf-8 -*-
"""模板库管理测试"""
import pytest
from pathlib import Path
import sys

sys.path.insert(0, str(Path(__file__).parent.parent))
sys.path.insert(0, str(Path(__file__).parent.parent / 'web_system'))

from template_library import TemplateLibrary, TemplateInfo


class TestTemplateLibrary:
    @pytest.fixture
    def library(self, tmp_path):
        lib = TemplateLibrary(str(tmp_path / 'templates'))
        return lib
    
    def test_add_template(self, library, sample_standard_doc):
        info = library.add_template(
            name='测试模板',
            file_path=str(sample_standard_doc),
            template_type='test',
            chapter_structure={'chapters': []},
            style_config={}
        )
        
        assert info.id.startswith('template_')
        assert info.name == '测试模板'
        assert Path(info.file_path).exists()
        print(f"添加模板：{info.name}, ID: {info.id}")
    
    def test_get_all_templates(self, library, sample_standard_doc):
        library.add_template('模板 1', str(sample_standard_doc), 'test', {}, {})
        library.add_template('模板 2', str(sample_standard_doc), 'test', {}, {})
        
        templates = library.get_all_templates()
        assert len(templates) == 2
        print(f"模板库中有 {len(templates)} 个模板")
    
    def test_get_template(self, library, sample_standard_doc):
        info = library.add_template('测试', str(sample_standard_doc), 'test', {}, {})
        
        retrieved = library.get_template(info.id)
        assert retrieved is not None
        assert retrieved.name == '测试'
    
    def test_delete_template(self, library, sample_standard_doc):
        info = library.add_template('测试', str(sample_standard_doc), 'test', {}, {})
        
        assert library.delete_template(info.id) is True
        assert library.get_template(info.id) is None
    
    def test_download_template(self, library, sample_standard_doc, tmp_path):
        info = library.add_template('测试', str(sample_standard_doc), 'test', {}, {})
        
        output_path = tmp_path / 'downloaded.docx'
        assert library.download_template(info.id, str(output_path)) is True
        assert output_path.exists()
        print(f"下载模板到：{output_path}")

if __name__ == '__main__':
    pytest.main([__file__, '-v'])
'''

with open('tests/test_template_library.py', 'w', encoding='utf-8') as f:
    f.write(library_test_content)
print('OK: tests/test_template_library.py')
