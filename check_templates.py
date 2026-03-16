# -*- coding: utf-8 -*-
"""检查模板目录中的文档结构"""
from docx import Document
import glob

def check_doc_structure(doc_path):
    """检查文档结构"""
    print(f'\n{"="*80}')
    print(f'文档：{doc_path}')
    print(f'{"="*80}')
    
    doc = Document(doc_path)
    
    styles_count = {}
    headings = []
    
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else 'Normal'
        styles_count[style_name] = styles_count.get(style_name, 0) + 1
        
        # 检测标题
        import re
        if 'Heading' in style_name or '标题' in style_name:
            level = 1 if '1' in style_name else (2 if '2' in style_name else (3 if '3' in style_name else 4))
            headings.append({'level': level, 'style': style_name, 'text': text[:50]})
        elif re.match(r'^\d+\.\d+', text):
            headings.append({'level': 0, 'style': style_name, 'text': text[:50]})
    
    print(f'\n总段落数：{len(doc.paragraphs)}')
    print(f'\n样式分布:')
    for style, count in sorted(styles_count.items(), key=lambda x: -x[1])[:10]:
        print(f'  {style}: {count}')
    
    print(f'\n标题数：{len(headings)}')
    print(f'\n前 20 个标题:')
    for h in headings[:20]:
        level_str = f'L{h["level"]}' if h['level'] > 0 else 'N '
        print(f'  [{h["style"]}] {level_str} {h["text"]}')
    
    return styles_count, headings

# 检查模板目录中的文件
template_files = glob.glob(r'e:\Qwen\xmjl\templates\*.docx')

if template_files:
    print(f'找到 {len(template_files)} 个模板文件')
    
    # 检查最后一个模板
    latest_template = template_files[-1]
    styles, headings = check_doc_structure(latest_template)
    
    # 与原始模板对比
    print(f'\n\n{"="*80}')
    print('与原始模板对比')
    print(f'{"="*80}')
    original_styles, original_headings = check_doc_structure(r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx')
    
    print(f'\n原始模板样式：{len(original_styles)} 种，标题：{len(original_headings)} 个')
    print(f'生成模板样式：{len(styles)} 种，标题：{len(headings)} 个')
    
    # 对比 Heading 2 的数量
    orig_h2 = original_styles.get('Heading 2', 0)
    gen_h2 = styles.get('Heading 2', 0)
    print(f'\nHeading 2 样式：原始={orig_h2}, 生成={gen_h2}')
    
    orig_body = original_styles.get('Body Text First Indent 2', 0)
    gen_body = styles.get('Body Text First Indent 2', 0)
    print(f'Body Text First Indent 2 样式：原始={orig_body}, 生成={gen_body}')
else:
    print('未找到模板文件')
