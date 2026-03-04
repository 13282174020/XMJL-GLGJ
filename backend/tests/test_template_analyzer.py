# -*- coding: utf-8 -*-
"""
模板分析服务测试用例
"""

import unittest
import os
import sys

# 添加父目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.services.template_analyzer import TemplateAnalyzer, analyze_template


class TestTemplateAnalyzer(unittest.TestCase):
    """模板分析器测试"""

    def setUp(self):
        """测试前准备"""
        self.template_path = None
    
    def tearDown(self):
        """测试后清理"""
        pass

    def test_analyze_template_structure(self):
        """测试分析模板结构"""
        # 创建一个简单的测试文档
        from docx import Document
        doc = Document()
        doc.add_heading('第 1 章 项目概况', level=1)
        doc.add_heading('1.1 项目名称', level=2)
        doc.add_heading('1.2 项目建设单位', level=2)
        doc.add_heading('1.2.1 建设单位信息', level=3)
        doc.add_paragraph('这是测试内容。')
        
        filename = os.path.join(os.path.dirname(__file__), 'test_template.docx')
        doc.save(filename)
        self.template_path = filename
        
        try:
            # 测试分析
            analyzer = TemplateAnalyzer(filename)
            tree = analyzer.get_chapter_tree()
            
            # 验证章节树
            self.assertGreater(len(tree), 0)
            
            # 验证第一章
            chapter1 = tree[0]
            self.assertEqual(chapter1['title'], '第 1 章 项目概况')
            self.assertEqual(chapter1['level'], 1)
            
            # 验证小节
            self.assertGreater(len(chapter1['subsections']), 0)
            section1 = chapter1['subsections'][0]
            self.assertEqual(section1['title'], '1.1 项目名称')
            self.assertEqual(section1['level'], 2)
            
        finally:
            if os.path.exists(filename):
                os.remove(filename)

    def test_get_style_for_level(self):
        """测试根据层级获取样式"""
        from docx import Document
        doc = Document()
        doc.add_heading('第 1 章 项目概况', level=1)
        doc.add_paragraph('测试内容。')
        
        filename = os.path.join(os.path.dirname(__file__), 'test_style.docx')
        doc.save(filename)
        self.template_path = filename
        
        try:
            analyzer = TemplateAnalyzer(filename)
            
            # 测试获取各级样式
            style1 = analyzer.get_style_for_level(1)
            self.assertIsNotNone(style1)
            
            style2 = analyzer.get_style_for_level(2)
            self.assertIsNotNone(style2)
            
        finally:
            if os.path.exists(filename):
                os.remove(filename)

    def test_get_full_structure(self):
        """测试获取完整结构"""
        from docx import Document
        doc = Document()
        doc.add_heading('第 1 章 项目概况', level=1)
        doc.add_heading('1.1 项目名称', level=2)
        doc.add_paragraph('测试内容。')
        
        filename = os.path.join(os.path.dirname(__file__), 'test_full.docx')
        doc.save(filename)
        self.template_path = filename
        
        try:
            analyzer = TemplateAnalyzer(filename)
            structure = analyzer.get_full_structure()
            
            self.assertIn('chapter_tree', structure)
            self.assertIn('styles', structure)
            self.assertIn('total_chapters', structure)
            
        finally:
            if os.path.exists(filename):
                os.remove(filename)


if __name__ == '__main__':
    unittest.main()
