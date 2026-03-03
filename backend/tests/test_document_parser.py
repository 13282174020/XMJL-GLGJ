# -*- coding: utf-8 -*-
"""
文档解析服务测试用例
"""

import unittest
import os
import sys

# 添加父目录到路径
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from app.services.document_parser import DocumentParser, parse_document


class TestDocumentParser(unittest.TestCase):
    """文档解析器测试"""

    def setUp(self):
        """测试前准备"""
        self.test_file = None
    
    def tearDown(self):
        """测试后清理"""
        if self.test_file and os.path.exists(self.test_file):
            os.remove(self.test_file)
    
    def _create_test_docx(self, content: str) -> str:
        """创建测试文档"""
        from docx import Document
        doc = Document()
        doc.add_heading('第一章 项目概况', level=1)
        doc.add_paragraph(content)
        
        filename = os.path.join(os.path.dirname(__file__), 'test_doc.docx')
        doc.save(filename)
        self.test_file = filename
        return filename

    # ========== 正常业务场景测试 ==========

    def test_extract_text_success(self):
        """测试成功提取文本"""
        filepath = self._create_test_docx('这是测试内容。')
        parser = DocumentParser(filepath)
        text = parser.extract_text()
        
        self.assertIsNotNone(text)
        self.assertIn('这是测试内容。', text)

    def test_extract_paragraphs_success(self):
        """测试成功提取段落"""
        filepath = self._create_test_docx('测试段落内容。')
        parser = DocumentParser(filepath)
        paragraphs = parser.extract_paragraphs()
        
        self.assertGreater(len(paragraphs), 0)
        self.assertIn('text', paragraphs[0])
        self.assertIn('level', paragraphs[0])

    def test_extract_full_structure(self):
        """测试提取完整结构"""
        filepath = self._create_test_docx('完整结构测试。')
        parser = DocumentParser(filepath)
        result = parser.extract_full_structure()
        
        self.assertIn('full_text', result)
        self.assertIn('paragraphs', result)
        self.assertIn('tables', result)
        self.assertIn('character_count', result)

    # ========== 边界条件测试 ==========

    def test_empty_document(self):
        """测试空文档"""
        from docx import Document
        doc = Document()
        filename = os.path.join(os.path.dirname(__file__), 'test_empty.docx')
        doc.save(filename)
        self.test_file = filename
        
        parser = DocumentParser(filename)
        text = parser.extract_text()
        
        self.assertEqual(text, '')

    def test_file_not_exists(self):
        """测试文件不存在"""
        with self.assertRaises(FileNotFoundError):
            DocumentParser('/non/existent/file.docx')

    # ========== 异常处理测试 ==========

    def test_invalid_file(self):
        """测试无效文件"""
        # 创建一个非 docx 文件
        filename = os.path.join(os.path.dirname(__file__), 'test_invalid.docx')
        with open(filename, 'w') as f:
            f.write('这不是一个真正的 docx 文件')
        self.test_file = filename
        
        with self.assertRaises(Exception):
            DocumentParser(filename)


class TestParseDocument(unittest.TestCase):
    """parse_document 便捷函数测试"""

    def test_parse_document(self):
        """测试 parse_document 函数"""
        from docx import Document
        
        # 创建测试文档
        doc = Document()
        doc.add_heading('第一章 测试', level=1)
        doc.add_paragraph('测试内容。')
        
        filename = os.path.join(os.path.dirname(__file__), 'test_parse.docx')
        doc.save(filename)
        
        try:
            result = parse_document(filename, 'requirement')
            
            self.assertEqual(result['file_type'], 'requirement')
            self.assertIn('full_text', result)
            self.assertIn('character_count', result)
        finally:
            if os.path.exists(filename):
                os.remove(filename)


if __name__ == '__main__':
    unittest.main()
