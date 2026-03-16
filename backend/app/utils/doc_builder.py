# -*- coding: utf-8 -*-
"""
文档渲染工具 - SKILL-007
DocBuilder 工具类，封装 Word 文档生成操作
支持从模板加载样式，确保生成文档与模板格式一致
"""

import os
import json
from typing import Dict, List, Any, Optional
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class DocBuilder:
    """文档构建器类"""
    
    # 默认样式配置
    DEFAULT_STYLES = {
        'page': {
            'width_cm': 21,
            'height_cm': 29.7,
            'top_margin_cm': 2.54,
            'bottom_margin_cm': 2.54,
            'left_margin_cm': 3.17,
            'right_margin_cm': 3.17
        },
        'styles': {
            'chapter_title': {
                'font': '黑体',
                'size_pt': 22,
                'bold': True,
                'alignment': 'center',
                'space_before_pt': 18,
                'space_after_pt': 12
            },
            'section_title': {
                'font': '黑体',
                'size_pt': 16,
                'bold': True,
                'alignment': 'left',
                'space_before_pt': 12,
                'space_after_pt': 6
            },
            'subsection_title': {
                'font': '黑体',
                'size_pt': 14,
                'bold': True,
                'alignment': 'left',
                'space_before_pt': 8,
                'space_after_pt': 4
            },
            'body_text': {
                'font': '仿宋',
                'size_pt': 10.5,
                'bold': False,
                'alignment': 'justify',
                'first_indent_chars': 2,
                'line_spacing': 1.5
            },
            'table_header': {
                'font': '黑体',
                'size_pt': 10.5,
                'bold': True,
                'alignment': 'center',
                'bg_color': 'D9E2F3'
            },
            'table_cell': {
                'font': '宋体',
                'size_pt': 10.5,
                'alignment': 'center'
            },
            'cover_title': {
                'font': '黑体',
                'size_pt': 36,
                'bold': True,
                'alignment': 'center'
            },
            'cover_subtitle': {
                'font': '黑体',
                'size_pt': 26,
                'bold': True,
                'alignment': 'center'
            }
        }
    }
    
    def __init__(self, template_path: Optional[str] = None, style_config: Optional[Dict] = None):
        """初始化文档构建器
        
        Args:
            template_path: Word 模板文件路径（可选，用于加载样式）
            style_config: 样式配置字典（可选）
        """
        self.template_path = template_path
        self.style_config = style_config or self.DEFAULT_STYLES['styles']
        self.template_styles = {}  # 从模板加载的样式
        
        # 如果提供了模板文件，加载模板样式
        if template_path and os.path.exists(template_path):
            self._load_template_styles(template_path)
            # 使用模板中的文档作为基础
            self.doc = Document(template_path)
            # 清除模板中的内容，保留样式
            self._clear_document_content()
        else:
            self.doc = Document()
        
        self._setup_styles()
    
    def _load_template_styles(self, template_path: str) -> None:
        """从模板文档加载样式定义

        Args:
            template_path: 模板文件路径
        """
        try:
            template_doc = Document(template_path)

            # 提取所有样式的定义
            # 注意：template_doc.styles 返回的是样式对象列表
            for style in template_doc.styles:
                style_name = style.name
                style_info = {
                    'name': style_name,
                    'font_name': None,
                    'font_size': None,
                    'bold': False,
                    'italic': False,
                    'alignment': 'left',
                    'space_before': 0,
                    'space_after': 0,
                    'line_spacing': 1.0,
                    'first_indent': 0
                }

                # 提取字体信息
                if hasattr(style, 'font') and style.font:
                    font = style.font
                    if font.name:
                        style_info['font_name'] = font.name
                    if font.size:
                        style_info['font_size'] = font.size.pt
                    style_info['bold'] = font.bold if font.bold is not None else False
                    style_info['italic'] = font.italic if font.italic is not None else False

                # 提取段落格式
                if hasattr(style, 'paragraph_format') and style.paragraph_format:
                    pf = style.paragraph_format
                    if hasattr(pf, 'alignment') and pf.alignment:
                        style_info['alignment'] = str(pf.alignment)
                    if hasattr(pf, 'space_before') and pf.space_before:
                        style_info['space_before'] = pf.space_before.pt if pf.space_before else 0
                    if hasattr(pf, 'space_after') and pf.space_after:
                        style_info['space_after'] = pf.space_after.pt if pf.space_after else 0
                    if hasattr(pf, 'line_spacing') and pf.line_spacing:
                        style_info['line_spacing'] = pf.line_spacing
                    if hasattr(pf, 'first_line_indent') and pf.first_line_indent:
                        style_info['first_indent'] = pf.first_line_indent.cm if pf.first_line_indent else 0

                self.template_styles[style_name] = style_info

            # 映射模板样式到内部样式
            self._map_template_styles()

        except Exception as e:
            print(f"警告：无法加载模板样式：{str(e)}")
            import traceback
            traceback.print_exc()
    
    def _map_template_styles(self) -> None:
        """将模板样式映射到内部样式配置"""
        # 标题样式映射
        style_mapping = {
            'chapter_title': ['Heading 1', '标题 1', '章标题', '1'],
            'section_title': ['Heading 2', '标题 2', '节标题', '2'],
            'subsection_title': ['Heading 3', '标题 3', '小节标题', '3'],
            'body_text': ['Normal', '正文', 'GP 正文 (首行缩进)']
        }
        
        for internal_name, template_names in style_mapping.items():
            for template_name in template_names:
                if template_name in self.template_styles:
                    ts = self.template_styles[template_name]
                    self.style_config[internal_name] = {
                        'font': ts.get('font_name') or self.style_config[internal_name].get('font', '仿宋'),
                        'size_pt': ts.get('font_size') or self.style_config[internal_name].get('size_pt', 10.5),
                        'bold': ts.get('bold', False),
                        'alignment': self._convert_alignment(ts.get('alignment', 'left')),
                        'space_before_pt': ts.get('space_before', 0),
                        'space_after_pt': ts.get('space_after', 0),
                        'first_indent_chars': 2 if ts.get('first_indent', 0) > 0 else 0,
                        'line_spacing': ts.get('line_spacing', 1.5)
                    }
                    break
    
    def _convert_alignment(self, alignment_str: str) -> str:
        """转换对齐方式字符串"""
        alignment_map = {
            'center': 'center',
            'left': 'left',
            'right': 'right',
            'justify': 'justify',
            'WD_ALIGN_PARAGRAPH.CENTER': 'center',
            'WD_ALIGN_PARAGRAPH.LEFT': 'left',
            'WD_ALIGN_PARAGRAPH.RIGHT': 'right',
            'WD_ALIGN_PARAGRAPH.JUSTIFY': 'justify'
        }
        return alignment_map.get(str(alignment_str), 'left')
    
    def _clear_document_content(self) -> None:
        """清除文档内容，保留样式"""
        # 保留第一个段落（通常包含样式定义），清除其他内容
        while len(self.doc.paragraphs) > 1:
            p = self.doc.paragraphs[-1]
            p._element.getparent().remove(p._element)
        
        # 清除表格
        while len(self.doc.tables) > 0:
            t = self.doc.tables[-1]
            t._element.getparent().remove(t._element)
    
    def _setup_styles(self) -> None:
        """设置文档样式"""
        style = self.doc.styles['Normal']
        body_style = self.style_config.get('body_text', {})
        
        style.font.name = body_style.get('font', '仿宋')
        style.font.size = Pt(body_style.get('size_pt', 10.5))
        style._element.rPr.rFonts.set(qn('w:eastAsia'), body_style.get('font', '仿宋'))
    
    def add_cover(self, project_info: Dict[str, Any]) -> None:
        """添加封面页
        
        Args:
            project_info: 项目信息字典
        """
        # 主标题
        project_name = project_info.get('name', '建设项目')
        title = self.doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Cm(5)
        title.paragraph_format.space_after = Cm(1)
        
        run = title.add_run(project_name)
        self._apply_style(run, self.style_config.get('cover_title', {}))
        
        # 副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle.paragraph_format.space_before = Cm(1)
        subtitle.paragraph_format.space_after = Cm(3)
        
        run = subtitle.add_run('可行性研究报告')
        self._apply_style(run, self.style_config.get('cover_subtitle', {}))
        
        # 空白行
        for _ in range(5):
            self.doc.add_paragraph()
        
        # 建设单位
        builder = self.doc.add_paragraph()
        builder.alignment = WD_ALIGN_PARAGRAPH.CENTER
        builder.paragraph_format.space_before = Cm(0.5)
        builder.paragraph_format.space_after = Cm(0.5)
        run = builder.add_run(f"建设单位：{project_info.get('org_name', 'XX 单位')}")
        run.font.name = '楷体'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        # 编制单位
        compiler = self.doc.add_paragraph()
        compiler.alignment = WD_ALIGN_PARAGRAPH.CENTER
        compiler.paragraph_format.space_before = Cm(0.5)
        compiler.paragraph_format.space_after = Cm(0.5)
        run = compiler.add_run("编制单位：XX 数字科技有限公司")
        run.font.name = '楷体'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        # 日期
        date_para = self.doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_para.paragraph_format.space_before = Cm(0.5)
        date_para.paragraph_format.space_after = Cm(0.5)
        run = date_para.add_run(f'编制日期：{datetime.now().strftime("%Y 年 %m 月")}')
        run.font.name = '楷体'
        run.font.size = Pt(16)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
        
        self.add_page_break()
    
    def add_toc(self, chapters: List[Dict[str, Any]]) -> None:
        """添加目录页
        
        Args:
            chapters: 章节列表
        """
        self._add_heading_style('目录', level=1)
        self.doc.add_paragraph()
        
        for chapter in chapters:
            # 章标题
            para = self.doc.add_paragraph()
            para.paragraph_format.left_indent = Cm(0)
            run = para.add_run(chapter.get('number', '') + ' ' + chapter.get('title', ''))
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            
            # 小节
            for subsection in chapter.get('subsections', []):
                para = self.doc.add_paragraph()
                para.paragraph_format.left_indent = Cm(0.74)
                run = para.add_run(subsection.get('number', '') + ' ' + subsection.get('title', ''))
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        self.add_page_break()
    
    def _add_heading_style(self, text: str, level: int = 1, style_name: str = None) -> None:
        """添加标题，使用模板样式

        Args:
            text: 标题文本
            level: 标题层级 (1-4)
            style_name: 原始样式名（可选，用于保留模板中的样式）
        """
        import logging
        
        # 如果有原始样式名，尝试使用模板样式
        if style_name and style_name in self.template_styles:
            logging.info(f'[DOC_BUILD]           模板样式 "{style_name}" 存在，使用模板样式')
            self._add_heading_with_template_style(text, style_name)
            return
        
        # 样式名不在 template_styles 中，检查是否是标准样式
        if style_name:
            logging.info(f'[DOC_BUILD]           模板样式 "{style_name}" 不存在（template_styles 有 {len(self.template_styles)} 个样式）')
            logging.info(f'[DOC_BUILD]           尝试直接使用样式名创建段落')
            # 尝试直接使用样式名创建段落
            try:
                para = self.doc.add_paragraph(style=style_name)
                run = para.add_run(text)
                logging.info(f'[DOC_BUILD]           成功使用样式 "{style_name}" 创建标题')
                return
            except KeyError:
                logging.warning(f'[DOC_BUILD]           样式 "{style_name}" 在文档中不存在，使用默认样式')
        
        # 否则使用默认样式配置
        logging.info(f'[DOC_BUILD]           使用默认 L{level} 样式配置')
        # 获取对应层级的样式配置
        style_config = self._get_style_for_level(level)

        # 添加段落
        para = self.doc.add_paragraph()

        # 设置对齐方式
        alignment = style_config.get('alignment', 'left')
        if alignment == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 设置间距
        space_before = style_config.get('space_before_pt', 0)
        space_after = style_config.get('space_after_pt', 0)
        if space_before:
            para.paragraph_format.space_before = Cm(space_before / 72 * 2.54)
        if space_after:
            para.paragraph_format.space_after = Cm(space_after / 72 * 2.54)

        # 添加文本
        run = para.add_run(text)
        self._apply_style(run, style_config)

    def _add_heading_with_template_style(self, text: str, style_name: str) -> None:
        """使用模板中的原始样式添加标题

        Args:
            text: 标题文本
            style_name: 模板中的样式名
        """
        if style_name not in self.template_styles:
            # 如果模板样式不存在，回退到默认方式
            self._add_heading_style(text, level=1)
            return

        ts = self.template_styles[style_name]
        
        import logging
        logging.info(f'[DOC_BUILD]           使用样式 "{style_name}" 创建标题')
        
        # 尝试直接使用样式名创建段落
        # 注意：样式必须已经在文档中定义
        try:
            para = self.doc.add_paragraph(style=style_name)
        except KeyError:
            # 如果样式不存在，使用 Normal 样式
            logging.warning(f'[DOC_BUILD]           样式 "{style_name}" 不存在，使用 Normal 样式')
            para = self.doc.add_paragraph()
        
        # 设置对齐方式
        alignment = self._convert_alignment(ts.get('alignment', 'left'))
        if alignment == 'center':
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif alignment == 'right':
            para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif alignment == 'justify':
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # 设置间距
        space_before = ts.get('space_before', 0)
        space_after = ts.get('space_after', 0)
        if space_before:
            para.paragraph_format.space_before = Cm(space_before / 72 * 2.54)
        if space_after:
            para.paragraph_format.space_after = Cm(space_after / 72 * 2.54)
        
        # 首行缩进
        first_indent = ts.get('first_indent', 0)
        if first_indent:
            para.paragraph_format.first_line_indent = Cm(first_indent)
        
        # 行距
        line_spacing = ts.get('line_spacing', 1.0)
        para.paragraph_format.line_spacing = line_spacing
        
        # 添加文本
        run = para.add_run(text)
        run.font.name = ts.get('font_name') or '宋体'
        run.font.size = Pt(ts.get('font_size') or 10.5)
        run.font.bold = ts.get('bold', False)
        run._element.rPr.rFonts.set(qn('w:eastAsia'), ts.get('font_name') or '宋体')
    
    def _get_style_for_level(self, level: int) -> Dict[str, Any]:
        """根据层级获取样式配置
        
        Args:
            level: 标题层级 (1-4)
            
        Returns:
            样式配置字典
        """
        styles = {
            1: self.style_config.get('chapter_title', {}),
            2: self.style_config.get('section_title', {}),
            3: self.style_config.get('subsection_title', {}),
            4: self.style_config.get('subsection_title', {})
        }
        return styles.get(level, styles[4])
    
    def add_chapter_title(self, title: str, style_name: str = None) -> None:
        """添加章标题

        Args:
            title: 章标题
            style_name: 原始样式名（可选，用于保留模板中的样式）
        """
        self._add_heading_style(title, level=1, style_name=style_name)
    
    def add_section_title(self, title: str, style_name: str = None) -> None:
        """添加节标题

        Args:
            title: 节标题
            style_name: 原始样式名（可选，用于保留模板中的样式）
        """
        self._add_heading_style(title, level=2, style_name=style_name)
    
    def add_subsection_title(self, title: str) -> None:
        """添加小节标题
        
        Args:
            title: 小节标题
        """
        self._add_heading_style(title, level=3)
    
    def add_body(self, text: str) -> None:
        """添加正文段落
        
        Args:
            text: 段落文本
        """
        if not text.strip():
            return
        
        para = self.doc.add_paragraph()
        run = para.add_run(text)
        
        body_style = self.style_config.get('body_text', {})
        run.font.name = body_style.get('font', '仿宋')
        run.font.size = Pt(body_style.get('size_pt', 10.5))
        run._element.rPr.rFonts.set(qn('w:eastAsia'), body_style.get('font', '仿宋'))
        
        # 首行缩进
        first_indent = body_style.get('first_indent_chars', 2)
        if first_indent > 0:
            para.paragraph_format.first_line_indent = Cm(first_indent * 0.37)
        
        # 行距
        line_spacing = body_style.get('line_spacing', 1.5)
        para.paragraph_format.line_spacing = line_spacing
        para.paragraph_format.space_before = Cm(0)
        para.paragraph_format.space_after = Cm(0)
    
    def add_table(self, headers: List[str], rows: List[List[str]]) -> None:
        """添加表格
        
        Args:
            headers: 表头
            rows: 数据行
        """
        if not headers or not rows:
            return
        
        table = self.doc.add_table(rows=len(rows) + 1, cols=len(headers))
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        header_row = table.rows[0]
        for i, header in enumerate(headers):
            cell = header_row.cells[i]
            cell.text = header
            
            # 设置表头样式
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = self.style_config.get('table_header', {}).get('font', '黑体')
                    run.font.size = Pt(self.style_config.get('table_header', {}).get('size_pt', 10.5))
                    run.font.bold = self.style_config.get('table_header', {}).get('bold', True)
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        
        # 设置数据行
        for i, row_data in enumerate(rows):
            row = table.rows[i + 1]
            for j, cell_text in enumerate(row_data):
                cell = row.cells[j]
                cell.text = cell_text
                
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = self.style_config.get('table_cell', {}).get('font', '宋体')
                        run.font.size = Pt(self.style_config.get('table_cell', {}).get('size_pt', 10.5))
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    def add_page_break(self) -> None:
        """添加分页符"""
        self.doc.add_page_break()
    
    def _apply_style(self, run, style: Dict[str, Any]) -> None:
        """应用样式到 run
        
        Args:
            run: docx run 对象
            style: 样式配置
        """
        if style.get('font'):
            run.font.name = style['font']
            run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font'])
        if style.get('size_pt'):
            run.font.size = Pt(style['size_pt'])
        if style.get('bold'):
            run.font.bold = True
    
    def save(self, output_path: str) -> str:
        """保存文档
        
        Args:
            output_path: 输出文件路径
            
        Returns:
            保存的文件路径
        """
        # 确保目录存在
        os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
        self.doc.save(output_path)
        return output_path
    
    def render_chapter(self, chapter_data: Dict[str, Any]) -> None:
        """渲染单个章节
        
        Args:
            chapter_data: 章节数据 JSON
        """
        chapter_number = chapter_data.get('chapter_number', '')
        chapter_title = chapter_data.get('chapter_title', '')
        
        # 添加章标题
        self.add_chapter_title(f"{chapter_number} {chapter_title}")
        
        # 渲染小节
        for subsection in chapter_data.get('subsections', []):
            self._render_subsection(subsection)
        
        self.add_page_break()
    
    def _render_subsection(self, subsection: Dict[str, Any]) -> None:
        """渲染小节
        
        Args:
            subsection: 小节数据
        """
        number = subsection.get('number', '')
        title = subsection.get('title', '')
        level = subsection.get('level', 2)
        
        # 添加标题
        if level == 2:
            self.add_section_title(f"{number} {title}")
        elif level == 3:
            self.add_subsection_title(f"{number} {title}")
        elif level >= 4:
            self._add_heading_style(f"{number} {title}", level=4)
        
        # 渲染内容
        content_type = subsection.get('type', 'text')
        
        if content_type in ('text', 'list', 'mixed'):
            for para in subsection.get('paragraphs', []):
                self.add_body(para)
        
        if content_type in ('table', 'mixed'):
            table_data = subsection.get('table_data')
            if table_data:
                headers = table_data.get('headers', [])
                rows = table_data.get('rows', [])
                self.add_table(headers, rows)
        
        # 递归渲染子节点
        for child in subsection.get('children', []):
            self._render_subsection(child)
    
    def render_document(self, json_data: Dict[str, Any]) -> None:
        """渲染完整文档
        
        Args:
            json_data: 完整文档 JSON 数据
        """
        # 封面
        project_info = json_data.get('project_info', {})
        self.add_cover(project_info)
        
        # 目录
        chapters = json_data.get('chapters', [])
        self.add_toc(chapters)
        
        # 正文
        for chapter in json_data.get('generated_chapters', []):
            self.render_chapter(chapter)
    
    def render_from_template_structure(self, template_structure: Dict[str, Any],
                                       chapter_contents: Dict[str, str]) -> None:
        """根据模板结构渲染文档

        Args:
            template_structure: 模板结构树（来自 TemplateAnalyzer）
            chapter_contents: 各章节内容字典 {章节编号：内容文本}
        """
        import logging
        import shutil
        import tempfile
        
        logging.info('=' * 80)
        logging.info('[DOC_BUILD] 开始根据模板结构渲染文档')
        logging.info(f'[DOC_BUILD] 模板路径：{self.template_path}')
        logging.info(f'[DOC_BUILD] 章节数：{len(template_structure.get("chapter_tree", []))}')
        logging.info(f'[DOC_BUILD] 内容章节数：{len(chapter_contents)}')
        
        # 如果有模板文件，复制模板文档并清除内容（保留样式）
        if self.template_path and os.path.exists(self.template_path):
            logging.info('[DOC_BUILD] 检测到模板文件，复制模板文档并清除内容...')
            
            # 创建临时文件复制模板
            temp_dir = tempfile.mkdtemp()
            temp_template = os.path.join(temp_dir, 'template.docx')
            shutil.copy2(self.template_path, temp_template)
            
            # 加载模板文档
            from docx import Document
            template_doc = Document(temp_template)
            
            # 清除所有段落和表格（保留样式定义）
            logging.info(f'[DOC_BUILD] 清除模板内容：{len(template_doc.paragraphs)} 个段落，{len(template_doc.tables)} 个表格')
            
            # 清除段落
            for para in list(template_doc.paragraphs):
                para._element.getparent().remove(para._element)
            
            # 清除表格
            for table in list(template_doc.tables):
                table._element.getparent().remove(table._element)
            
            # 使用清除后的文档
            self.doc = template_doc
            
            # 清理临时文件
            os.remove(temp_template)
            os.rmdir(temp_dir)
            
            logging.info('[DOC_BUILD] 模板内容已清除，样式已保留')
        else:
            logging.info('[DOC_BUILD] 无模板文件，创建新文档')
            from docx import Document
            self.doc = Document()
        
        # 添加封面
        logging.info('[DOC_BUILD] 正在添加封面...')
        self.add_cover({'name': '建设项目', 'org_name': 'XX 单位'})
        
        # 添加目录
        logging.info('[DOC_BUILD] 正在添加目录...')
        chapters = template_structure.get('chapter_tree', [])
        self.add_toc(chapters)
        logging.info(f'[DOC_BUILD] 目录条目数：{len(chapters)}')
        
        # 按模板结构逐章渲染
        logging.info('[DOC_BUILD] 开始渲染章节内容...')
        for i, chapter in enumerate(chapters):
            logging.info(f'[DOC_BUILD] 渲染第 {i+1}/{len(chapters)} 章：{chapter.get("title", "未知")}')
            logging.info(f'[DOC_BUILD]   章节样式：{chapter.get("style", "未知")}')
            logging.info(f'[DOC_BUILD]   子节点数：{len(chapter.get("children", []))}')
            self._render_template_chapter(chapter, chapter_contents)
        
        logging.info('[DOC_BUILD] 章节渲染完成')
        logging.info('=' * 80)

    def _update_template_headings(self, template_structure: Dict[str, Any],
                                   chapter_contents: Dict[str, str]) -> None:
        """更新模板文档中的标题文本
        
        Args:
            template_structure: 模板结构树
            chapter_contents: 章节内容字典
        """
        # 这个方法用于在模板文档基础上修改标题文本
        # 目前保持简单实现，直接使用 render_from_template_structure 的逻辑
        # 因为模板文档已经包含了所有样式和结构
        
        # 添加封面（如果需要）
        # 注意：模板文档可能已经有封面，这里选择添加新的封面
        # 可以在后续优化中判断是否已有封面
        
        # 实际上，如果模板文档已经包含完整的结构，我们只需要：
        # 1. 保留模板文档的样式
        # 2. 根据章节内容替换或添加内容段落
        
        # 简化处理：直接使用模板文档，不修改标题
        # 因为模板文档的标题已经是正确的格式
        pass
    
    def _render_template_chapter(self, chapter: Dict[str, Any],
                                  chapter_contents: Dict[str, str]) -> None:
        """渲染模板中的章节

        Args:
            chapter: 章节数据（来自模板分析）
            chapter_contents: 章节内容字典
        """
        import logging
        logging.info(f'[DOC_BUILD]   添加章标题：{chapter.get("title", "未知")}')
        logging.info(f'[DOC_BUILD]     章节样式：{chapter.get("style", "未知")}')
        
        # 添加章标题（传递样式名）
        style_name = chapter.get('style')
        self.add_chapter_title(chapter['title'], style_name=style_name)

        # 渲染子节点（修复：使用 'children' 而不是 'subsections'）
        for i, section in enumerate(chapter.get('children', [])):
            logging.info(f'[DOC_BUILD]     渲染小节 {i+1}/{len(chapter.get("children", []))}: {section.get("title", "未知")}')
            logging.info(f'[DOC_BUILD]       小节样式：{section.get("style", "未知")}')
            self._render_template_section(section, chapter_contents)

        self.add_page_break()
        logging.info(f'[DOC_BUILD]   章节 {chapter.get("title", "未知")} 渲染完成，添加分页符')
    
    def _render_template_section(self, section: Dict[str, Any],
                                   chapter_contents: Dict[str, str]) -> None:
        """渲染模板中的小节

        Args:
            section: 小节数据
            chapter_contents: 章节内容字典
        """
        import logging
        logging.info(f'[DOC_BUILD]       添加节标题：{section.get("title", "未知")}')
        logging.info(f'[DOC_BUILD]         原始样式：{section.get("style", "未知")}')
        
        # 添加节标题（使用原始样式名）
        style_name = section.get('style')
        self.add_section_title(section['title'], style_name=style_name)
        logging.info(f'[DOC_BUILD]         标题已添加')

        # 获取内容（如果有）
        content = chapter_contents.get(section['title'], '')
        if content:
            logging.info(f'[DOC_BUILD]         内容长度：{len(content)} 字符')
            # 将内容分割成段落
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    self.add_body(para.strip())
            logging.info(f'[DOC_BUILD]         已添加 {len(paragraphs)} 个段落')
        else:
            logging.info(f'[DOC_BUILD]         无内容，跳过')

        # 渲染子节点（三级标题等）
        for i, child in enumerate(section.get('children', [])):
            logging.info(f'[DOC_BUILD]         渲染子节点 {i+1}/{len(section.get("children", []))}: {child.get("title", "未知")}')
            self._render_template_subsection(child, chapter_contents)

    def _render_template_subsection(self, subsection: Dict[str, Any],
                                     chapter_contents: Dict[str, str]) -> None:
        """渲染模板中的小小节

        Args:
            subsection: 小小节数据
            chapter_contents: 章节内容字典
        """
        import logging
        logging.info(f'[DOC_BUILD]           添加小小节标题：{subsection.get("title", "未知")}')
        logging.info(f'[DOC_BUILD]             原始样式：{subsection.get("style", "未知")}')
        
        # 添加小小节标题（使用原始样式名）
        style_name = subsection.get('style')
        self._add_heading_style(subsection['title'], level=3, style_name=style_name)
        logging.info(f'[DOC_BUILD]             标题已添加')

        # 获取内容
        content = chapter_contents.get(subsection['title'], '')
        if content:
            logging.info(f'[DOC_BUILD]             内容长度：{len(content)} 字符')
            paragraphs = content.split('\n')
            for para in paragraphs:
                if para.strip():
                    self.add_body(para.strip())
            logging.info(f'[DOC_BUILD]             已添加 {len(paragraphs)} 个段落')
        else:
            logging.info(f'[DOC_BUILD]             无内容，跳过')


def create_doc_builder(template_path: Optional[str] = None, style_config: Optional[Dict] = None) -> DocBuilder:
    """创建 DocBuilder 实例的便捷函数"""
    return DocBuilder(template_path, style_config)


if __name__ == '__main__':
    # 测试代码
    builder = DocBuilder()
    
    # 测试封面
    builder.add_cover({
        'name': '测试项目',
        'org_name': '测试单位'
    })
    
    # 测试保存
    builder.save('test_output.docx')
    print("测试文档生成成功")
