# -*- coding: utf-8 -*-
"""
文档解析服务 - SKILL-001
解析 Word 文档，提取文本内容、表格数据、标题层级等信息
"""

import os
from typing import Dict, List, Any, Optional, Tuple
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.paragraph import Paragraph


class DocumentParser:
    """文档解析器类"""
    
    def __init__(self, file_path: str):
        """初始化文档解析器
        
        Args:
            file_path: Word 文档路径
        """
        self.file_path = file_path
        self.doc = None
        self._load_document()
    
    def _load_document(self) -> None:
        """加载 Word 文档"""
        if not os.path.exists(self.file_path):
            raise FileNotFoundError(f"文件不存在：{self.file_path}")
        
        try:
            self.doc = Document(self.file_path)
        except Exception as e:
            raise Exception(f"无法读取 Word 文档：{str(e)}")
    
    def extract_text(self) -> str:
        """提取文档全部文本内容
        
        Returns:
            文档全文文本
        """
        texts = []
        for para in self.doc.paragraphs:
            if para.text.strip():
                texts.append(para.text)
        return '\n'.join(texts)
    
    def extract_paragraphs(self) -> List[Dict[str, Any]]:
        """提取段落信息（含层级）
        
        Returns:
            段落信息列表，每项包含：text, level, style_name
        """
        paragraphs = []
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 判断段落层级
            level, style_name = self._detect_heading_level(para)
            
            paragraphs.append({
                'text': text,
                'level': level,
                'style_name': style_name,
                'is_heading': level > 0
            })
        
        return paragraphs
    
    def _detect_heading_level(self, para: Paragraph) -> Tuple[int, str]:
        """检测段落是否为标题及其层级
        
        Args:
            para: 段落对象
            
        Returns:
            (层级，样式名), 层级 0 表示正文，1-9 表示标题层级
        """
        style_name = para.style.name if para.style else ''
        
        # 检查样式名
        heading_styles = ['Heading', '标题', '章标题', '节标题', '小节标题']
        for i, h_style in enumerate(heading_styles):
            if h_style.lower() in style_name.lower():
                return (i + 1, style_name)
        
        # 检查字体大小（标题通常字体较大）
        if para.runs:
            font_size = para.runs[0].font.size
            if font_size and font_size.pt >= 16:
                return (1, style_name)  # 大字体视为一级标题
            elif font_size and font_size.pt >= 14:
                return (2, style_name)
            elif font_size and font_size.pt >= 12:
                return (3, style_name)
        
        # 检查是否以章节编号开头
        text = para.text.strip()
        if self._is_chapter_heading(text):
            return (1, style_name)
        elif self._is_section_heading(text):
            return (2, style_name)
        elif self._is_subsection_heading(text):
            return (3, style_name)
        
        return (0, style_name or 'Normal')
    
    def _is_chapter_heading(self, text: str) -> bool:
        """判断是否为章标题（如：第一章、第 1 章）"""
        import re
        patterns = [
            r'^第 [一二三四五六七八九十\d]+章',
            r'^Chapter\s*\d+',
        ]
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def _is_section_heading(self, text: str) -> bool:
        """判断是否为节标题（如：1.1、第一节）"""
        import re
        patterns = [
            r'^\d+\.\d+\s+',
            r'^第 [一二三四五六七八九十\d]+[节条]',
            r'^Section\s*\d+',
        ]
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def _is_subsection_heading(self, text: str) -> bool:
        """判断是否为小节标题（如：1.1.1）"""
        import re
        patterns = [
            r'^\d+\.\d+\.\d+\s+',
            r'^\([一二三四五六七八九十\d]+\)',
        ]
        for pattern in patterns:
            if re.match(pattern, text):
                return True
        return False
    
    def extract_tables(self) -> List[Dict[str, Any]]:
        """提取文档中的表格
        
        Returns:
            表格数据列表，每项包含：headers, rows, caption
        """
        tables = []
        
        for table in self.doc.tables:
            table_data = self._parse_table(table)
            if table_data:
                tables.append(table_data)
        
        return tables
    
    def _parse_table(self, table: Table) -> Optional[Dict[str, Any]]:
        """解析单个表格
        
        Args:
            table: 表格对象
            
        Returns:
            表格数据字典
        """
        try:
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):  # 跳过空行
                    rows.append(cells)
            
            if not rows:
                return None
            
            # 假设第一行是表头
            headers = rows[0] if rows else []
            data_rows = rows[1:] if len(rows) > 1 else []
            
            return {
                'headers': headers,
                'rows': data_rows,
                'row_count': len(rows),
                'col_count': len(headers) if headers else 0
            }
        except Exception as e:
            return None
    
    def extract_full_structure(self) -> Dict[str, Any]:
        """提取文档完整结构
        
        Returns:
            包含文本、段落、表格的完整结构
        """
        return {
            'full_text': self.extract_text(),
            'paragraphs': self.extract_paragraphs(),
            'tables': self.extract_tables(),
            'paragraph_count': len(self.doc.paragraphs),
            'table_count': len(self.doc.tables),
            'character_count': len(self.extract_text())
        }
    
    def is_too_long(self, max_chars: int = 120000) -> bool:
        """检查文档是否超长
        
        Args:
            max_chars: 最大字符数限制
            
        Returns:
            是否超长
        """
        return len(self.extract_text()) > max_chars
    
    def split_by_chapters(self) -> List[Dict[str, Any]]:
        """按章节拆分文档
        
        Returns:
            章节列表，每项包含：title, level, content
        """
        chapters = []
        current_chapter = None
        current_content = []
        
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            level, _ = self._detect_heading_level(para)
            
            if level == 1:  # 新章开始
                # 保存前一章
                if current_chapter:
                    current_chapter['content'] = '\n'.join(current_content)
                    chapters.append(current_chapter)
                
                # 开始新章
                current_chapter = {
                    'title': text,
                    'level': 1,
                    'content': ''
                }
                current_content = []
            else:
                current_content.append(text)
        
        # 保存最后一章
        if current_chapter:
            current_chapter['content'] = '\n'.join(current_content)
            chapters.append(current_chapter)
        
        return chapters


def parse_document(file_path: str, file_type: str = 'requirement') -> Dict[str, Any]:
    """解析文档的便捷函数
    
    Args:
        file_path: 文件路径
        file_type: 文件类型（requirement/template）
        
    Returns:
        解析结果字典
    """
    parser = DocumentParser(file_path)
    
    result = {
        'file_path': file_path,
        'file_type': file_type,
        'full_text': parser.extract_text(),
        'paragraphs': parser.extract_paragraphs(),
        'tables': parser.extract_tables(),
        'chapters': parser.split_by_chapters(),
        'is_too_long': parser.is_too_long(),
        'character_count': len(parser.extract_text()),
        'paragraph_count': len(parser.doc.paragraphs),
        'table_count': len(parser.doc.tables)
    }
    
    return result


if __name__ == '__main__':
    # 测试代码
    import sys
    if len(sys.argv) > 1:
        result = parse_document(sys.argv[1])
        print(f"文档：{sys.argv[1]}")
        print(f"字符数：{result['character_count']}")
        print(f"段落数：{result['paragraph_count']}")
        print(f"表格数：{result['table_count']}")
        print(f"章节数：{len(result['chapters'])}")
        if result['chapters']:
            print("章节列表:")
            for ch in result['chapters'][:5]:
                print(f"  - {ch['title']}")
