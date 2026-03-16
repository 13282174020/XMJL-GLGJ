# -*- coding: utf-8 -*-
"""调试文档分析"""
from docx import Document
import re

file_path = r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx'
doc = Document(file_path)

# 编号正则
numbering_patterns = [
    (r'^(\d+)\.(\d+)\.(\d+)\.(\d+)', 4),
    (r'^(\d+)\.(\d+)\.(\d+)', 3),
    (r'^(\d+)\.(\d+)', 2),
]

print('=' * 80)
print('检查所有非 Normal/Heading 样式段落')
print('=' * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    style_lower = style_name.lower()
    
    # 跳过标准标题样式
    if 'heading' in style_lower or '标题' in style_lower:
        continue
    
    # 检查是否匹配编号模式
    matched = False
    for pattern, level in numbering_patterns:
        if re.match(pattern, text):
            matched = True
            print(f'段落 {i}: [{style_name}] 匹配编号 L{level}: {text[:50]}')
            break
    
    # 检查是否以数字开头（可能是 1 级标题）
    if not matched and re.match(r'^\d+', text):
        print(f'段落 {i}: [{style_name}] 数字开头但未匹配编号：{text[:50]}')

print('\n' + '=' * 80)
print('样式统计')
print('=' * 80)

styles_count = {}
for para in doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    styles_count[style_name] = styles_count.get(style_name, 0) + 1

for style, count in sorted(styles_count.items(), key=lambda x: -x[1]):
    print(f'  {style}: {count} 次')
