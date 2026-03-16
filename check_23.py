# -*- coding: utf-8 -*-
from docx import Document

doc = Document('新模板_4.docx')

print('=' * 80)
print('查找 2.3 相关章节')
print('=' * 80)

found = False
for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name if para.style else 'Normal'
    
    # 查找 2.3 相关内容
    if '2.3' in text or '项目建设' in text or '政策文件' in text:
        found = True
        is_heading = 'Heading' in style
        status = '✓ Heading' if is_heading else '✗ 非 Heading'
        print(f'{status}: [{style:30s}] {text[:70]}')

if not found:
    print('未找到包含 2.3 的段落')
    
print()
print('=' * 80)
print('所有 Heading 2 样式的段落')
print('=' * 80)

for i, para in enumerate(doc.paragraphs):
    text = para.text.strip()
    if not text:
        continue
    style = para.style.name if para.style else 'Normal'
    
    if style == 'Heading 2':
        print(f'段落 {i:4d}: {text[:70]}')
