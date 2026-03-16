# -*- coding: utf-8 -*-
from docx import Document

try:
    doc = Document('新模板_4.docx')
    print(f'段落数：{len(doc.paragraphs)}')
    
    styles = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else 'Normal'
        styles[style] = styles.get(style, 0) + 1
    
    print('文档中的样式分布:')
    for style, count in sorted(styles.items(), key=lambda x: -x[1])[:15]:
        print(f'  {style}: {count}')
except Exception as e:
    print(f'错误：{e}')
    import traceback
    traceback.print_exc()
