#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
未来社区建设方案生成器
从"目录与样式配置"中加载章节结构和样式配置，生成 Word 文档
支持：
1. 从 uploads/chapter_config.json 加载目录结构
2. 从 uploads/style_config.json 加载样式配置
3. 根据配置的样式动态渲染文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os
import json
from pathlib import Path


# ==================== 默认样式配置 ====================

DEFAULT_STYLES = {
    'heading1': {
        'font': '黑体',
        'size_pt': 22,
        'bold': True,
        'alignment': 'center',
        'space_before_pt': 24,
        'space_after_pt': 12
    },
    'heading2': {
        'font': '黑体',
        'size_pt': 16,
        'bold': True,
        'alignment': 'left',
        'space_before_pt': 18,
        'space_after_pt': 6
    },
    'heading3': {
        'font': '黑体',
        'size_pt': 14,
        'bold': True,
        'alignment': 'left',
        'space_before_pt': 12,
        'space_after_pt': 4
    },
    'heading4': {
        'font': '黑体',
        'size_pt': 12,
        'bold': True,
        'alignment': 'left',
        'space_before_pt': 6,
        'space_after_pt': 3
    },
    'heading5': {
        'font': '黑体',
        'size_pt': 11,
        'bold': True,
        'alignment': 'left',
        'space_before_pt': 3,
        'space_after_pt': 2
    },
    'heading6': {
        'font': '黑体',
        'size_pt': 10,
        'bold': True,
        'alignment': 'left',
        'space_before_pt': 2,
        'space_after_pt': 1
    },
    'normal': {
        'font': '仿宋',
        'size_pt': 12,
        'bold': False,
        'alignment': 'justify',
        'first_indent_chars': 2,
        'line_spacing': 1.5
    }
}


# ==================== 配置加载函数 ====================

def load_style_config(style_config_path=None):
    """加载样式配置
    
    Args:
        style_config_path: 样式配置文件路径，默认使用 uploads/style_config.json
        
    Returns:
        样式配置字典
    """
    if style_config_path is None:
        # 默认从 uploads 目录加载
        base_dir = Path(__file__).parent
        style_config_path = base_dir / 'web_system' / 'uploads' / 'style_config.json'
    
    if os.path.exists(style_config_path):
        with open(style_config_path, 'r', encoding='utf-8') as f:
            styles = json.load(f)
        print(f'已加载样式配置：{style_config_path}')
        return styles
    else:
        print(f'未找到样式配置文件，使用默认样式：{style_config_path}')
        return DEFAULT_STYLES


def load_chapter_config(chapters_path=None):
    """加载章节目录配置
    
    Args:
        chapters_path: 章节目录配置文件路径，默认使用 uploads/chapter_config.json
        
    Returns:
        章节目录列表
    """
    if chapters_path is None:
        # 默认从 uploads 目录加载
        base_dir = Path(__file__).parent
        chapters_path = base_dir / 'web_system' / 'uploads' / 'chapter_config.json'
    
    if os.path.exists(chapters_path):
        with open(chapters_path, 'r', encoding='utf-8') as f:
            chapters = json.load(f)
        print(f'已加载章节目录配置：{chapters_path}')
        print(f'  共 {len(chapters)} 个一级章节')
        return chapters
    else:
        print(f'未找到章节目录配置文件：{chapters_path}')
        return []


# ==================== 样式函数 ====================

def add_heading(doc, text, level, styles=None):
    """添加标题 - 根据配置的样式动态渲染
    
    Args:
        doc: Document 对象
        text: 标题文本
        level: 标题层级 (1-4)
        styles: 样式配置字典
    """
    if styles is None:
        styles = DEFAULT_STYLES
    
    # 根据层级获取样式配置
    style_key = f'heading{level}'
    style = styles.get(style_key, DEFAULT_STYLES.get(style_key, {}))
    
    p = doc.add_paragraph()
    
    # 设置对齐方式
    alignment = style.get('alignment', 'left')
    if alignment == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif alignment == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    elif alignment == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # 设置间距
    space_before = style.get('space_before_pt', 0)
    space_after = style.get('space_after_pt', 0)
    if space_before:
        p.paragraph_format.space_before = Pt(space_before)
    if space_after:
        p.paragraph_format.space_after = Pt(space_after)
    
    # 添加文本
    run = p.add_run(text)
    
    # 应用字体样式
    font_name = style.get('font', '黑体')
    font_size = style.get('size_pt', 12)
    bold = style.get('bold', True)
    
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    return p


def add_heading_1(doc, text, styles=None):
    """添加一级标题 - 黑体 二号 18pt（兼容旧接口）"""
    return add_heading(doc, text, level=1, styles=styles)


def add_heading_2(doc, text, styles=None):
    """添加二级标题 - 黑体 三号 15pt（兼容旧接口）"""
    return add_heading(doc, text, level=2, styles=styles)


def add_heading_3(doc, text, styles=None):
    """添加三级标题 - 黑体 小三 14pt（兼容旧接口）"""
    return add_heading(doc, text, level=3, styles=styles)


def add_heading_4(doc, text, styles=None):
    """添加四级标题 - 黑体 四号 12pt（兼容旧接口）"""
    return add_heading(doc, text, level=4, styles=styles)


def add_normal(doc, text, styles=None):
    """添加正文段落 - 根据配置的样式动态渲染
    
    Args:
        doc: Document 对象
        text: 段落文本
        styles: 样式配置字典
    """
    if styles is None:
        styles = DEFAULT_STYLES
    
    style = styles.get('normal', DEFAULT_STYLES.get('normal', {}))
    
    paragraphs = text.split('\n')
    for para_text in paragraphs:
        if not para_text.strip():
            continue
        
        p = doc.add_paragraph()
        
        # 设置行距
        line_spacing = style.get('line_spacing', 1.5)
        p.paragraph_format.line_spacing = line_spacing
        
        # 设置首行缩进
        first_indent_chars = style.get('first_indent_chars', 2)
        if first_indent_chars > 0:
            p.paragraph_format.first_line_indent = Cm(first_indent_chars * 0.37)
        
        # 添加文本
        run = p.add_run(para_text)
        
        # 应用字体样式
        font_name = style.get('font', '仿宋')
        font_size = style.get('size_pt', 12)
        bold = style.get('bold', False)
        
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


# ==================== 文档结构函数 ====================

def add_title_page(doc, project_info=None, styles=None):
    """添加封面页 - 根据配置的样式动态渲染
    
    Args:
        doc: Document 对象
        project_info: 项目信息字典，包含 name, org_name 等
        styles: 样式配置字典
    """
    if styles is None:
        styles = DEFAULT_STYLES
    
    if project_info is None:
        project_info = {
            'name': '建设项目',
            'subtitle': '可行性研究报告',
            'org_name': 'XX 单位',
            'compiler': 'XX 数字科技有限公司',
            'date': None
        }
    
    # 添加空白行（约 8 行）
    for _ in range(8):
        doc.add_paragraph()
    
    # 主标题
    title = project_info.get('name', '建设项目')
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(title)
    
    # 使用封面标题样式（如果没有则使用 heading1）
    cover_style = styles.get('cover_title', styles.get('heading1', {}))
    run.font.size = Pt(cover_style.get('size_pt', 26))
    run.font.bold = cover_style.get('bold', True)
    run.font.name = cover_style.get('font', '黑体')
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    # 副标题
    subtitle = project_info.get('subtitle', '可行性研究报告')
    if subtitle:
        for _ in range(3):
            doc.add_paragraph()
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = sub_para.add_run(subtitle)
        
        cover_sub_style = styles.get('cover_subtitle', styles.get('heading2', {}))
        run.font.size = Pt(cover_sub_style.get('size_pt', 16))
        run.font.name = cover_sub_style.get('font', '楷体')
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    # 添加空白行（约 12 行）
    for _ in range(12):
        doc.add_paragraph()
    
    # 建设单位
    builder = project_info.get('org_name', 'XX 单位')
    if builder:
        p = doc.add_paragraph(f'建设单位：{builder}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.name = '楷体'
        p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    # 编制单位
    compiler = project_info.get('compiler', 'XX 数字科技有限公司')
    if compiler:
        p = doc.add_paragraph(f'编制单位：{compiler}')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.size = Pt(12)
        p.runs[0].font.name = '楷体'
        p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    # 日期
    date = project_info.get('date')
    if not date:
        date = '编制日期：2026 年 3 月'
    p = doc.add_paragraph(date)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = '楷体'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '楷体')
    
    doc.add_page_break()


def add_editor_list(doc, styles=None):
    """添加编审人员名单
    
    Args:
        doc: Document 对象
        styles: 样式配置字典
    """
    if styles is None:
        styles = DEFAULT_STYLES
    
    # 标题
    title = doc.add_paragraph('项目编审人员名单')
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.bold = True
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.name = '黑体'
    title.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    
    doc.add_paragraph()
    
    # 编审人员名单（使用占位符）
    items = [
        '项目负责：×××',
        '编制小组：×××、×××、×××',
        '勘误核稿：×××',
        '项目审定：×××'
    ]
    for item in items:
        p = doc.add_paragraph(item)
        p.paragraph_format.left_indent = Cm(2)
    
    doc.add_page_break()


def add_toc(doc, chapters, styles=None):
    """添加目录页

    Args:
        doc: Document 对象
        chapters: 章节树列表
        styles: 样式配置字典
    """
    if styles is None:
        styles = DEFAULT_STYLES

    # 目录标题（使用一级标题样式）
    add_heading(doc, '目录', level=1, styles=styles)
    doc.add_paragraph()

    def render_chapters(nodes, level=0):
        for node in nodes:
            indent = '  ' * level
            p = doc.add_paragraph(f"{indent}{node.get('number', '')} {node.get('title', '')}")
            p.paragraph_format.left_indent = Cm(0.74 * level)

            # 设置目录项样式（使用正文样式）
            normal_style = styles.get('normal', DEFAULT_STYLES.get('normal', {}))
            run = p.runs[0] if p.runs else p.add_run('')
            run.font.name = normal_style.get('font', '宋体')
            run.font.size = Pt(normal_style.get('size_pt', 10.5))
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

            # 递归渲染子节点（检查 children 是否存在且非空）
            children = node.get('children')
            if children is not None and len(children) > 0:
                render_chapters(children, level + 1)

    render_chapters(chapters)
    doc.add_page_break()


# ==================== 内容生成函数 ====================

def generate_content_for_chapter(chapter_title):
    """
    根据章节标题生成内容
    这里使用预定义的专业内容，后续可以根据需求文档动态生成
    """
    content_templates = {
        '项目概况': '良熟新苑未来社区建设项目位于杭州市余杭区良渚街道，是回迁安置小区（一期、二期合并，封闭式小区），总人口约 1.1 万人，3000 多户。项目以"邻里、健康、治理、低碳"四大场景为核心，致力于打造回迁社区未来社区建设样板。\n\n项目由良熟社区居委会负责建设，通过数字化改革和硬件设施提升，切实解决居民实际需求（监控老化、电瓶车防盗、流动人口管理等），提升社区治理能力和居民生活品质。',
        
        '项目名称': '良熟新苑未来社区建设项目',
        
        '项目建设单位': '良熟社区居委会是良熟新苑未来社区的建设单位，作为基层群众自治组织，负责社区的日常管理和服务工作。\n\n主要职能包括：\n1. 社区治理：负责社区日常管理和公共服务\n2. 物业管理：社区自己管理物业，自负盈亏\n3. 民生保障：负责居民生活保障和服务\n4. 文化建设：组织邻里活动，增强社区凝聚力',
        
        '建设目标': '良熟新苑未来社区的建设目标是：以人民为中心，以数字化改革为牵引，聚焦邻里、健康、治理、低碳四大场景，打造回迁社区未来社区样板。\n\n具体目标包括：\n1. 提升社区安全水平：解决监控老化、电瓶车盗窃等安全问题\n2. 改善居民生活品质：建设共享书房、健康服务设施\n3. 提升治理能力：实现流动人口精准管理、平安码系统\n4. 优化物业服务：探索水电自主上报功能，提升服务满意度\n5. 符合未来社区验收标准：达到浙江省未来社区建设要求',
        
        '建设规模': '良熟新苑社区总人口约 1.1 万人，3000 多户，包含一期、二期两个小区（实际合并为一个封闭小区）。\n\n建设规模包括：\n1. 智慧安防系统：高清监控、高空抛物探头、电瓶车防盗、单元门禁\n2. 智慧治理平台：平安码系统、人口数据分析、流动人口管理\n3. 邻里服务设施：共享书房、邻里中心、健康服务点\n4. 物业管理平台：水电自主上报系统、费用管理',
        
        '建设内容': '根据需求调研，良熟新苑未来社区重点建设以下内容：\n\n一、未来邻里场景\n1. 依托现有邻里中心，增设共享书房\n2. 打造 15 分钟生活圈\n3. 开展美丽庭院建设\n4. 组织邻里活动，增强社区凝聚力\n\n二、未来健康场景\n1. 建设健康服务点\n2. 配置健康监测设备\n3. 重点关注老年人健康管理\n\n三、未来治理场景\n1. 平安码系统（类似扫楼宝）\n2. 高空抛物监控探头\n3. 单元门禁系统\n4. 人员结构分析\n\n四、未来低碳场景\n1. 按照新社区低碳验收标准建设\n2. 推广绿色生活方式\n\n五、物业管理提升\n1. 探索水电自主上报功能\n2. 提升物业服务满意度',
        
        '建设周期': '项目建设周期预计 12 个月，分三个阶段实施：\n第一阶段（1-4 月）：基础硬件设施建设\n第二阶段（5-8 月）：数字化平台建设\n第三阶段（9-12 月）：场景应用完善',
        
        '项目建设的必要性': '一、解决安全问题的迫切需要\n良熟新苑作为回迁安置小区，现有监控设备老化，电瓶车盗窃问题频发，居民对安全问题的诉求强烈。\n\n二、提升居民生活品质的需要\n社区人口约 1.1 万人，居民对共享书房、健康服务、邻里活动等品质生活的需求日益增长。\n\n三、数字化治理的需要\n流动人口管理困难，需要对接区平台，实现平安码系统、人员结构分析等数字化治理功能。\n\n四、提升物业服务的需要\n物业公司由社区自己管理，自负盈亏，需要提升服务满意度。\n\n五、符合未来社区建设要求\n新社区对低碳等场景有一定验收标准，需要按照浙江省未来社区建设规范进行改造提升。',
        
        '需求分析': '一、业务需求\n1. 安全需求：高空抛物监控、电瓶车防盗、人员结构分析\n2. 治理需求：流动人口管理、平安码扫楼系统\n3. 服务需求：共享书房、健康服务、邻里活动\n4. 物业需求：水电自主上报、物业费收缴管理\n\n二、功能需求\n1. 智慧安防：高清监控、高空抛物探头、烟感报警、单元门禁\n2. 智慧治理：人口数据分析、平安码系统\n3. 智慧服务：邻里中心、共享书房、健康驿站\n4. 智慧物业：水电表自主上报、费用管理',
        
        '总体思路': '坚持以人民为中心的发展思想，以提升居民生活品质为核心，以数字化改革为牵引，打造"邻里、健康、治理、低碳"四大特色场景，建设回迁社区未来社区样板。\n\n指导思想：\n1. 需求导向：解决老百姓实际痛点，不花里胡哨\n2. 实用优先：注重功能实用性，避免形象工程\n3. 数字赋能：对接区平台，实现数据共享\n4. 共建共享：居民参与，共建共治共享',
        
        '建设原则': '1. 需求导向原则\n   从业主实际想法出发，提升小区硬件设施，解决老百姓痛点。\n\n2. 实用优先原则\n   升级为能识别车牌的高清监控，解决电瓶车盗窃等实际问题。\n\n3. 数字赋能原则\n   与区平台对接，实现平安码、人口数据分析等数字化治理功能。\n\n4. 共建共享原则\n   物业公司（社区）自负盈亏，建立可持续的运营模式。',
        
        '投资估算': '投资估算包括硬件设备购置费、软件开发费、系统集成费、工程建设费等。\n\n一、硬件设备购置费\n1. 高清监控设备（含高空抛物探头）\n2. 单元门禁系统\n3. 烟感报警设备\n4. 电瓶车防盗设备\n\n二、软件开发费\n1. 平安码系统\n2. 物业管理平台\n3. 人口数据分析系统\n\n三、系统集成费\n1. 与区平台对接\n2. 与电力局、水利局系统对接\n\n四、工程建设费\n1. 邻里中心装修\n2. 共享书房建设\n3. 健康服务点建设',
        
        '资金筹措': '项目资金通过以下渠道解决：\n1. 政府补助：未来社区建设专项资金\n2. 社区自筹：社区集体收入\n3. 物业收益：物业费、停车费等收入',
        
        '效益分析': '一、社会效益\n1. 提升社区安全水平，减少电瓶车盗窃等案件\n2. 改善居民生活品质，增强社区凝聚力\n3. 提升治理能力，实现精准化管理\n4. 打造回迁社区未来社区样板，具有示范意义\n\n二、经济效益\n1. 提升物业服务质量，提高物业费收缴率\n2. 降低管理成本，提高管理效率\n3. 带动周边商业发展，提升区域价值',
        
        '结论': '良熟新苑未来社区建设项目符合浙江省未来社区建设要求，能够切实解决居民实际需求（监控老化、电瓶车防盗、流动人口管理等），具备建设的必要性和可行性。\n\n项目以"邻里、健康、治理、低碳"四大场景为核心，通过数字化改革和硬件设施提升，将显著提升社区治理能力和居民生活品质，打造回迁社区未来社区样板。',
        
        '建议': '1. 尽快启动项目实施\n   建议尽快启动项目实施，做好组织保障、资金保障、技术保障等工作。\n\n2. 充分听取居民意见\n   项目建设要从业主实际想法出发，解决老百姓痛点。\n\n3. 加强与区平台对接\n   实现与区平台、电力局、水利局系统的对接，确保数据共享和业务协同。\n\n4. 注重可持续发展\n   物业公司（社区）自负盈亏，需要建立可持续的运营模式。',
    }
    
    # 精确匹配
    if chapter_title in content_templates:
        return content_templates[chapter_title]
    
    # 模糊匹配
    for key, content in content_templates.items():
        if key in chapter_title:
            return content
    
    # 默认内容
    return f'{chapter_title}\n\n根据项目实际情况进行编制，确保符合可行性研究报告和未来社区建设规范要求。'


# ==================== 主生成函数 ====================

def generate_plan_from_config(chapters=None, styles=None, project_info=None, 
                               chapters_path=None, styles_path=None, output_path=None):
    """
    从"目录与样式配置"中加载配置，生成 Word 文档
    
    参数:
        chapters: 章节目录列表（可选，如果提供则直接使用）
        styles: 样式配置字典（可选，如果提供则直接使用）
        project_info: 项目信息字典（可选）
        chapters_path: 章节目录配置文件路径（可选）
        styles_path: 样式配置文件路径（可选）
        output_path: 输出文件路径（可选）
    
    返回:
        生成的文档路径
    """
    # 1. 加载章节目录配置
    if chapters is None:
        chapters = load_chapter_config(chapters_path)
        if not chapters:
            print('错误：章节目录配置为空，请确保已保存目录配置')
            return None
    
    # 2. 加载样式配置
    if styles is None:
        styles = load_style_config(styles_path)
    
    print(f'使用样式配置：{list(styles.keys())}')
    
    # 3. 创建文档
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    
    # 4. 封面
    print('生成封面...')
    if project_info is None:
        project_info = {
            'name': '建设项目',
            'subtitle': '可行性研究报告',
            'org_name': 'XX 单位',
            'compiler': 'XX 数字科技有限公司'
        }
    add_title_page(doc, project_info=project_info, styles=styles)
    
    # 5. 编审名单
    print('生成编审名单...')
    add_editor_list(doc, styles=styles)
    
    # 6. 目录
    print('生成目录...')
    add_toc(doc, chapters, styles=styles)
    
    # 7. 正文 - 严格按照目录配置生成
    print('生成正文...')
    
    def generate_chapter(node, level=1):
        """递归生成章节 - 严格按照配置的层级结构"""
        node_title = node.get('title', '')
        node_number = node.get('number', '')
        node_level = node.get('level', level)  # 使用配置中的 level
        children = node.get('children')
        
        # 添加标题（使用配置的样式）
        if node_level <= 6:  # 支持到 6 级标题
            add_heading(doc, f"{node_number} {node_title}", level=node_level, styles=styles)
        
        # 判断是否有子节点
        if children is not None and len(children) > 0:
            # 有子节点，递归处理
            for child in children:
                generate_chapter(child, level=node_level + 1)
        else:
            # 叶子节点，生成内容
            content = generate_content_for_chapter(node_title)
            add_normal(doc, content, styles=styles)
    
    for chapter in chapters:
        generate_chapter(chapter, level=1)
        doc.add_page_break()
    
    # 8. 保存文档
    if output_path is None:
        timestamp = __import__('datetime').datetime.now().strftime('%Y%m%d_%H%M%S')
        output_path = os.path.join(os.path.dirname(__file__), 'outputs', f'建设方案_{timestamp}.docx')
    
    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    print(f'文档已保存至：{output_path}')
    return output_path


def generate_plan_with_data(chapters, styles, project_info=None, output_path=None):
    """
    使用传入的数据直接生成文档（不加载配置文件）
    用于从 web_system 直接调用生成
    
    参数:
        chapters: 章节目录列表
        styles: 样式配置字典
        project_info: 项目信息字典（可选）
        output_path: 输出文件路径（可选）
    
    返回:
        生成的文档路径
    """
    return generate_plan_from_config(
        chapters=chapters,
        styles=styles,
        project_info=project_info,
        output_path=output_path
    )


def main():
    """主函数 - 从配置文件加载并生成文档"""
    # 使用默认路径加载配置
    chapters = load_chapter_config()
    styles = load_style_config()
    
    if not chapters:
        print('错误：未找到章节目录配置，请先在 web 系统中保存目录配置')
        return
    
    # 设置项目信息（后续可以从需求文档中提取）
    project_info = {
        'name': '良熟新苑未来社区建设项目',
        'subtitle': '可行性研究报告',
        'org_name': '良熟社区居委会',
        'compiler': '数字科技有限公司'
    }
    
    # 生成文档
    output_path = generate_plan_from_config(
        chapters=chapters,
        styles=styles,
        project_info=project_info
    )
    
    if output_path:
        print(f'\n生成完成！文档路径：{output_path}')


if __name__ == '__main__':
    main()
