# -*- coding: utf-8 -*-
from docx import Document

doc_path = '新模板_1773310314036 (1).docx'
doc = Document(doc_path)

print(f'=== 文档：{doc_path} ===')
print(f'总段落数：{len(doc.paragraphs)}\n')

print('段落详情:')
for i, p in enumerate(doc.paragraphs):
    text = p.text.strip()
    if not text:
        continue
    style_name = p.style.name if p.style else 'None'
    print(f'{i:3d}: [{style_name:20s}] | {text[:60]}')
