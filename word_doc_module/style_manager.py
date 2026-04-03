# -*- coding: utf-8 -*-
"""
Word 样式管理模块

功能：
1. 定义文档样式配置（字体、字号、对齐等）
2. 添加带样式的标题（Heading 1-6）
3. 添加正文段落（首行缩进、行距等）
4. 创建目录
5. 设置大纲级别（支持 Word 自动目录）

依赖：python-docx
安装：pip install python-docx
"""

from typing import Dict, Optional, Any
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# =============================================================================
# 默认样式配置
# =============================================================================

DEFAULT_STYLES: Dict[str, Dict[str, Any]] = {
    'heading1': {
        'font_name': '黑体',
        'font_size': 22,  # pt
        'bold': True,
        'alignment': 'center',
        'space_before': 0,  # cm
        'space_after': 0.5,  # cm
        'color': None,
    },
    'heading2': {
        'font_name': '楷体',
        'font_size': 16,
        'bold': True,
        'alignment': 'left',
        'space_before': 0.5,
        'space_after': 0.3,
        'color': None,
    },
    'heading3': {
        'font_name': '楷体',
        'font_size': 14,
        'bold': True,
        'alignment': 'left',
        'space_before': 0.3,
        'space_after': 0.2,
        'color': None,
    },
    'heading4': {
        'font_name': '黑体',
        'font_size': 12,
        'bold': True,
        'alignment': 'left',
        'space_before': 0.2,
        'space_after': 0.1,
        'color': None,
    },
    'heading5': {
        'font_name': '黑体',
        'font_size': 11,
        'bold': True,
        'alignment': 'left',
        'space_before': 0.2,
        'space_after': 0.1,
        'color': None,
    },
    'heading6': {
        'font_name': '黑体',
        'font_size': 10,
        'bold': True,
        'alignment': 'left',
        'space_before': 0.1,
        'space_after': 0.1,
        'color': None,
    },
    'normal': {
        'font_name': '仿宋',
        'font_size': 10.5,
        'bold': False,
        'alignment': 'left',
        'line_spacing': 1.5,  # 1.5 倍行距
        'first_line_indent': 0.74,  # cm，首行缩进 2 字符（约 0.74cm）
        'space_before': 0,
        'space_after': 0,
    },
}


# =============================================================================
# 字体设置辅助函数
# =============================================================================

def set_run_font(run, font_name: str, font_size: float = None,
                 bold: bool = None, italic: bool = None):
    """设置 Run 的字体属性

    Args:
        run: python-docx Run 对象
        font_name: 字体名称（如 '黑体', '仿宋', '楷体'）
        font_size: 字号（pt）
        bold: 是否加粗
        italic: 是否斜体
    """
    # 设置西文字体
    if font_name:
        run.font.name = font_name

    # 设置中文字体（关键！）
    if font_name:
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

    # 设置字号
    if font_size:
        run.font.size = Pt(font_size)

    # 设置加粗
    if bold is not None:
        run.font.bold = bold

    # 设置斜体
    if italic is not None:
        run.font.italic = italic


def set_paragraph_format(para, alignment: str = 'left',
                         space_before: float = 0, space_after: float = 0,
                         line_spacing: float = None, first_line_indent: float = None,
                         left_indent: float = None):
    """设置段落格式

    Args:
        para: python-docx Paragraph 对象
        alignment: 对齐方式 ('left', 'center', 'right', 'justify')
        space_before: 段前间距 (cm)
        space_after: 段后间距 (cm)
        line_spacing: 行距倍数（如 1.5）
        first_line_indent: 首行缩进 (cm)
        left_indent: 左边距 (cm)
    """
    # 对齐方式
    alignment_map = {
        'left': WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right': WD_ALIGN_PARAGRAPH.RIGHT,
        'justify': WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    if alignment in alignment_map:
        para.alignment = alignment_map[alignment]

    # 段前段后间距
    para.paragraph_format.space_before = Cm(space_before)
    para.paragraph_format.space_after = Cm(space_after)

    # 行距
    if line_spacing:
        para.paragraph_format.line_spacing = line_spacing

    # 首行缩进
    if first_line_indent is not None:
        para.paragraph_format.first_line_indent = Cm(first_line_indent)

    # 左边距
    if left_indent is not None:
        para.paragraph_format.left_indent = Cm(left_indent)


# =============================================================================
# 样式管理器
# =============================================================================

class StyleManager:
    """Word 样式管理器

    用于创建和管理 Word 文档的样式。

    示例：
        manager = StyleManager()
        doc = manager.create_document()

        # 添加标题
        manager.add_heading(doc, '项目概况', level=1)
        manager.add_heading(doc, '1.1 项目背景', level=2)

        # 添加正文
        manager.add_normal_paragraph(doc, '这是正文内容...')

        doc.save('output.docx')
    """

    def __init__(self, custom_styles: Dict[str, Dict] = None):
        """初始化样式管理器

        Args:
            custom_styles: 自定义样式配置，会与默认配置合并
        """
        self.styles = DEFAULT_STYLES.copy()
        if custom_styles:
            self.styles.update(custom_styles)

    def create_document(self, template_path: str = None) -> Document:
        """创建新的 Word 文档

        Args:
            template_path: 可选的模板文件路径，使用模板创建文档

        Returns:
            Document 对象
        """
        if template_path:
            return Document(template_path)
        return Document()

    def get_style_config(self, style_name: str) -> Dict[str, Any]:
        """获取样式配置

        Args:
            style_name: 样式名称（如 'heading1', 'normal'）

        Returns:
            样式配置字典
        """
        return self.styles.get(style_name, self.styles['normal'])

    def add_heading(self, doc: Document, text: str, level: int = 1,
                    styles: Dict[str, Dict] = None) -> Document:
        """添加标题段落

        Args:
            doc: Document 对象
            text: 标题文本
            level: 标题级别（1-6）
            styles: 可选的样式配置，会覆盖默认配置

        Returns:
            添加的 Paragraph 对象
        """
        # 获取样式配置
        style_key = f'heading{level}'
        config = (styles or self.styles).get(style_key, self.styles['normal'])

        # 创建标题
        heading = doc.add_heading(text, level=level)

        # 设置字体
        for run in heading.runs:
            set_run_font(
                run,
                font_name=config.get('font_name', '黑体'),
                font_size=config.get('font_size', 16),
                bold=config.get('bold', True)
            )

        # 设置对齐
        alignment = config.get('alignment', 'left')
        set_paragraph_format(
            heading,
            alignment=alignment,
            space_before=config.get('space_before', 0),
            space_after=config.get('space_after', 0.3)
        )

        # 清除默认的加粗（因为我们已经在 run 中设置了）
        return heading

    def add_normal_paragraph(self, doc: Document, text: str,
                            indent: bool = True,
                            styles: Dict[str, Dict] = None) -> Document:
        """添加正文段落

        Args:
            doc: Document 对象
            text: 正文文本
            indent: 是否首行缩进
            styles: 可选的样式配置

        Returns:
            添加的 Paragraph 对象
        """
        config = (styles or self.styles).get('normal', self.styles['normal'])

        para = doc.add_paragraph()
        run = para.add_run(text)

        # 设置字体
        set_run_font(
            run,
            font_name=config.get('font_name', '仿宋'),
            font_size=config.get('font_size', 10.5),
            bold=config.get('bold', False)
        )

        # 设置段落格式
        set_paragraph_format(
            para,
            alignment=config.get('alignment', 'left'),
            space_before=config.get('space_before', 0),
            space_after=config.get('space_after', 0),
            line_spacing=config.get('line_spacing', 1.5),
            first_line_indent=config.get('first_line_indent', 0.74) if indent else None
        )

        return para

    def add_cover_title(self, doc: Document, title: str,
                        font_size: int = 36) -> Document:
        """添加封面标题

        Args:
            doc: Document 对象
            title: 标题文本
            font_size: 字号（默认 36pt）

        Returns:
            添加的 Paragraph 对象
        """
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(title)
        run.font.name = '黑体'
        run.font.size = Pt(font_size)
        run.font.bold = True
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        return para

    def add_page_break(self, doc: Document) -> None:
        """添加分页符"""
        doc.add_page_break()

    def add_table(self, doc: Document, rows: int, cols: int,
                  data: list = None) -> Any:
        """添加表格

        Args:
            doc: Document 对象
            rows: 行数
            cols: 列数
            data: 表格数据，二维列表

        Returns:
            Table 对象
        """
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        if data:
            for i, row_data in enumerate(data):
                if i < rows:
                    row = table.rows[i]
                    for j, cell_text in enumerate(row_data):
                        if j < cols:
                            row.cells[j].text = str(cell_text)

        return table

    def ensure_outline_levels(self, doc: Document) -> None:
        """确保文档的 Heading 样式有正确的大纲级别

        这对于 Word 自动目录（TOC）至关重要。
        python-docx 默认不会为 Heading 样式设置 outlineLvl，
        导致目录无法自动收集标题。

        Args:
            doc: Document 对象
        """
        for i in range(1, 10):
            style_name = f'Heading {i}'
            try:
                style = doc.styles[style_name]
                pPr = style._element.pPr

                # 移除已存在的 outlineLvl
                for existing in pPr.findall(qn('w:outlineLvl')):
                    pPr.remove(existing)

                # 添加新的 outlineLvl
                outlineLvl = OxmlElement('w:outlineLvl')
                outlineLvl.set(qn('w:val'), str(i - 1))
                pPr.insert(0, outlineLvl)

            except Exception:
                # 样式不存在，跳过
                pass

    def apply_styles_to_document(self, doc: Document) -> None:
        """为文档应用所有样式配置

        设置中文字体、大纲级别等。

        Args:
            doc: Document 对象
        """
        # 确保大纲级别
        self.ensure_outline_levels(doc)

        # 设置默认字体
        try:
            normal_style = doc.styles['Normal']
            normal_style.font.name = '仿宋'
            normal_style.font.size = Pt(10.5)
            normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), '仿宋')
        except Exception:
            pass


# =============================================================================
# 便捷函数
# =============================================================================

def load_style_config(config_path: str = None) -> Dict[str, Dict[str, Any]]:
    """加载样式配置

    Args:
        config_path: 配置文件路径（JSON），如果为 None 则使用默认配置

    Returns:
        样式配置字典
    """
    if config_path:
        import json
        try:
            with open(config_path, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception:
            pass

    return DEFAULT_STYLES.copy()
