# -*- coding: utf-8 -*-
"""
Word 模板扫描模块

功能：从 Word 模板文档中提取章节结构、样式信息、编号模式。
用于 AI 生成文档时获取模板的章节大纲和格式参考。

依赖：python-docx
安装：pip install python-docx
"""

import re
from typing import Dict, List, Optional, Any
from docx import Document
from docx.shared import Pt, Cm


# =============================================================================
# 编号模式识别
# =============================================================================

class NumberingPattern:
    """编号模式定义"""

    # 按优先级排列（先匹配更具体的模式）
    PATTERNS = [
        # 纯数字多级编号 X.X.X.X
        (r'^(\d+)\.(\d+)\.(\d+)\.(\d+)', 4),
        # 三级编号 X.X.X
        (r'^(\d+)\.(\d+)\.(\d+)', 3),
        # 二级编号 X.X
        (r'^(\d+)\.(\d+)', 2),
        # 中文章节标题
        (r'^第\s*[一二三四五六七八九十百千零〇\d]+\s*[章节]', 1),
        # 括号编号 (一)(1)(①)
        (r'^\(([一二三四五六七八九十\d①-⑩]+)\)', 3),
        # 中文括号章节
        (r'^【([^】]+)】', 2),
    ]

    @classmethod
    def match(cls, text: str) -> Optional[tuple]:
        """匹配编号模式

        Args:
            text: 段落文本

        Returns:
            (level, number_str) 或 None
        """
        text = text.strip()
        for pattern, level in cls.PATTERNS:
            match = re.match(pattern, text)
            if match:
                number_str = match.group(0)
                return (level, number_str)
        return None


# =============================================================================
# 章节节点
# =============================================================================

class ChapterNode:
    """章节节点"""

    def __init__(self, level: int, number: str, title: str, style: str):
        self.level = level
        self.number = number
        self.title = title
        self.style = style
        self.children: List['ChapterNode'] = []

    def to_dict(self) -> Dict[str, Any]:
        return {
            'level': self.level,
            'number': self.number,
            'title': self.title,
            'style': self.style,
            'children': [c.to_dict() for c in self.children]
        }

    def __repr__(self):
        return f"ChapterNode(level={self.level}, title='{self.title}')"


# =============================================================================
# 样式级别判断
# =============================================================================

def get_heading_level(style_name: str, text: str = '') -> int:
    """从样式名称获取标题级别

    Args:
        style_name: Word 样式名称（如 'Heading 1', 'Title'）
        text: 段落文本（用于无样式时的备用判断）

    Returns:
        级别 0-6，0 表示非标题
    """
    if not style_name:
        # 尝试从文本推断
        match = NumberingPattern.match(text)
        if match:
            return match[0]
        return 0

    style_lower = style_name.lower()

    # Word 内置标题样式
    if 'heading 1' in style_lower or style_lower == 'heading1':
        return 1
    if 'heading 2' in style_lower or style_lower == 'heading2':
        return 2
    if 'heading 3' in style_lower or style_lower == 'heading3':
        return 3
    if 'heading 4' in style_lower or style_lower == 'heading4':
        return 4
    if 'heading 5' in style_lower or style_lower == 'heading5':
        return 5
    if 'heading 6' in style_lower or style_lower == 'heading6':
        return 6

    # 其他可能表示标题的样式
    title_indicators = ['title', 'caption', 'toc']
    if any(ind in style_lower for ind in title_indicators):
        return 1

    # 从文本模式判断
    match = NumberingPattern.match(text)
    if match:
        return match[0]

    return 0


def is_likely_heading(text: str, style_name: str = '') -> bool:
    """判断是否为标题段落

    Args:
        text: 段落文本
        style_name: 样式名称

    Returns:
        是否为标题
    """
    if not text or len(text.strip()) > 200:
        return False

    # 有明确标题样式
    if 'heading' in style_name.lower():
        return True

    # 匹配编号模式
    if NumberingPattern.match(text):
        return True

    return False


# =============================================================================
# 章节树构建
# =============================================================================

def build_chapter_tree(heading_list: List[Dict]) -> List[Dict]:
    """构建章节树

    Args:
        heading_list: 平铺的标题列表，每项包含 level, number, title, style

    Returns:
        树形结构的章节列表
    """
    if not heading_list:
        return []

    root_nodes: List[ChapterNode] = []
    stack: List[ChapterNode] = []

    for item in heading_list:
        node = ChapterNode(
            level=item.get('level', 1),
            number=item.get('number', ''),
            title=item.get('title', ''),
            style=item.get('style', '')
        )

        # 找到父节点
        while stack and stack[-1].level >= node.level:
            stack.pop()

        if not stack:
            root_nodes.append(node)
        else:
            stack[-1].children.append(node)

        stack.append(node)

    return [n.to_dict() for n in root_nodes]


def generate_chapter_numbering(heading_list: List[Dict]) -> List[Dict]:
    """为标题列表生成编号

    如果标题文本中没有编号，则自动生成编号（第一章、1.1 等）

    Args:
        heading_list: 原始标题列表

    Returns:
        带编号的标题列表
    """
    result = []
    counters = [0, 0, 0, 0, 0, 0]  # 各级别计数器

    for item in heading_list:
        level = item.get('level', 1)
        text = item.get('title', '')
        style = item.get('style', '')
        existing_number = item.get('number', '')

        # 更新计数器
        counters[level - 1] += 1
        for i in range(level, 6):
            counters[i] = 0

        # 确定编号
        if existing_number:
            number = existing_number
        else:
            number = _generate_number(level, counters, text)

        result.append({
            'level': level,
            'number': number,
            'title': text,
            'style': style
        })

    return result


def _generate_number(level: int, counters: List[int], text: str) -> str:
    """生成章节编号"""
    if level == 1:
        # 中文大写数字
        chinese_nums = ['一', '二', '三', '四', '五', '六', '七', '八', '九', '十']
        num = counters[0]
        if num <= 10:
            return f"第{chinese_nums[num - 1]}章"
        else:
            return f"第{num}章"
    else:
        # 数字编号
        parts = [str(counters[i]) for i in range(level - 1)]
        return '.'.join(parts)


def flatten_chapter_tree(tree: List[Dict], result: List[Dict] = None) -> List[Dict]:
    """将章节树展平为列表"""
    if result is None:
        result = []

    for node in tree:
        result.append({
            'level': node['level'],
            'number': node['number'],
            'title': node['title'],
            'style': node['style']
        })
        if node.get('children'):
            flatten_chapter_tree(node['children'], result)

    return result


# =============================================================================
# 模板扫描器
# =============================================================================

class TemplateScanner:
    """Word 模板扫描器

    用于从 Word 模板文档中提取：
    1. 章节结构（标题树）
    2. 样式使用统计
    3. 编号模式

    示例：
        scanner = TemplateScanner()
        result = scanner.scan('template.docx')
        print(result['chapters'])
    """

    def __init__(self):
        self.styles_usage: Dict[str, int] = {}
        self.heading_styles: Dict[str, Dict] = {}

    def scan(self, file_path: str = None, file_stream=None) -> Dict[str, Any]:
        """扫描 Word 模板

        Args:
            file_path: Word 文件路径
            file_stream: 文件流（BytesIO），与 file_path 二选一

        Returns:
            {
                'success': bool,
                'chapters': List[Dict],  # 章节树
                'chapters_flat': List[Dict],  # 展平的章节列表
                'styles_usage': Dict[str, int],  # 样式使用统计
                'total_paragraphs': int,
            }
        """
        if not file_path and not file_stream:
            return {'success': False, 'error': '必须提供 file_path 或 file_stream'}

        try:
            doc = Document(file_path) if file_path else Document(file_stream)
        except Exception as e:
            return {'success': False, 'error': f'无法打开文档: {str(e)}'}

        self.styles_usage.clear()
        self.heading_styles.clear()

        paragraphs = []
        total_paragraphs = len(doc.paragraphs)

        # 遍历段落
        for para in doc.paragraphs:
            text = para.text.strip()
            style_name = para.style.name if para.style else 'Normal'

            # 统计样式使用
            self.styles_usage[style_name] = self.styles_usage.get(style_name, 0) + 1

            # 判断是否为标题
            level = get_heading_level(style_name, text)

            if level > 0 and text:
                paragraphs.append({
                    'level': level,
                    'text': text,
                    'style': style_name,
                    'number': ''
                })

        # 生成编号
        paragraphs = generate_chapter_numbering(paragraphs)

        # 构建章节树
        chapters = build_chapter_tree(paragraphs)
        chapters_flat = flatten_chapter_tree(chapters)

        return {
            'success': True,
            'chapters': chapters,
            'chapters_flat': chapters_flat,
            'styles_usage': self.styles_usage.copy(),
            'total_paragraphs': total_paragraphs,
            'heading_count': len(paragraphs)
        }

    def get_template_text(self, file_path: str = None, file_stream=None) -> str:
        """获取模板全文文本

        用于 AI 生成时的格式参考

        Args:
            file_path: Word 文件路径
            file_stream: 文件流

        Returns:
            模板文档的纯文本内容
        """
        if not file_path and not file_stream:
            return ''

        try:
            doc = Document(file_path) if file_path else Document(file_stream)
        except Exception:
            return ''

        texts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                texts.append(text)

        return '\n'.join(texts)


# =============================================================================
# 便捷函数
# =============================================================================

def scan_template_styles(file_path: str) -> Dict[str, Any]:
    """扫描模板文件，提取章节结构

    这是一个便捷函数。

    Args:
        file_path: Word 模板文件路径

    Returns:
        {
            'success': bool,
            'chapters': List[Dict],  # 章节树
            'chapters_flat': List[Dict],  # 展平的章节列表
        }

    示例：
        result = scan_template_styles('template.docx')
        if result['success']:
            for chapter in result['chapters_flat']:
                print(f"{'#' * chapter['level']} {chapter['title']}")
    """
    scanner = TemplateScanner()
    return scanner.scan(file_path)


def extract_section_from_template(template_text: str, section_title: str) -> str:
    """从模板文本中提取指定章节的内容

    用于 AI 生成时获取参考格式。

    Args:
        template_text: 模板文档全文
        section_title: 要提取的章节标题

    Returns:
        章节内容，如果未找到则返回空字符串
    """
    if not template_text or not section_title:
        return ''

    # 查找章节标题位置
    patterns = [
        rf'{re.escape(section_title)}[\s\n]*',
        rf'{re.escape(section_title.split()[-1] if " " in section_title else section_title)}[\s\n]*',
    ]

    for pattern in patterns:
        match = re.search(pattern, template_text, re.IGNORECASE)
        if match:
            start_pos = match.end()
            # 查找下一个章节标题
            next_section_match = re.search(
                r'\n\s*(?:\d+\.|第[一二三四五六七八九十]+章|【)',
                template_text[start_pos:]
            )
            if next_section_match:
                end_pos = start_pos + next_section_match.start()
                section_content = template_text[start_pos:end_pos].strip()
            else:
                section_content = template_text[start_pos:].strip()

            # 限制长度
            if len(section_content) > 1500:
                section_content = section_content[:1500] + '...'

            return section_content

    return template_text[:800] if template_text else ''
