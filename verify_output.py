# -*- coding: utf-8 -*-
"""验证生成的文档"""
from docx import Document

template_path = r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx'
output_path = r'e:\Qwen\xmjl\outputs\test_with_logs.docx'

print('=' * 80)
print('验证生成的文档')
print('=' * 80)

orig_doc = Document(template_path)
gen_doc = Document(output_path)

orig_styles = {}
for para in orig_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    orig_styles[style_name] = orig_styles.get(style_name, 0) + 1

gen_styles = {}
for para in gen_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    gen_styles[style_name] = gen_styles.get(style_name, 0) + 1

print(f'\n原始模板段落数：{len(orig_doc.paragraphs)}')
print(f'生成文档段落数：{len(gen_doc.paragraphs)}')
print(f'差异：{len(orig_doc.paragraphs) - len(gen_doc.paragraphs)}')

print(f'\n原始模板样式数：{len(orig_styles)}')
print(f'生成文档样式数：{len(gen_styles)}')

print(f'\n样式对比:')
all_styles = set(orig_styles.keys()) | set(gen_styles.keys())
for style in sorted(all_styles):
    orig = orig_styles.get(style, 0)
    gen = gen_styles.get(style, 0)
    match = '✓' if orig == gen else '✗'
    print(f'  {match} {style}: 原始={orig}, 生成={gen}')

# 检查前 20 个标题
print('\n生成文档的前 20 个标题/编号段落:')
import re
count = 0
for para in gen_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    if 'Heading' in style_name or re.match(r'^\d+\.\d+', text):
        print(f'  [{style_name}] {text[:50]}')
        count += 1
        if count >= 20:
            break

print('\n' + '=' * 80)
print(f'生成文档路径：{output_path}')
print('=' * 80)
