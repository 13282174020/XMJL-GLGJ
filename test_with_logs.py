# -*- coding: utf-8 -*-
"""测试完整的模板分析和文档生成流程 - 带详细日志"""
import sys
import logging
sys.path.insert(0, 'backend')

# 配置日志输出到控制台
logging.basicConfig(
    level=logging.INFO,
    format='%(levelname)s - %(message)s',
    handlers=[logging.StreamHandler(sys.stdout)]
)

from app.services.template_analyzer import TemplateAnalyzer
from app.utils.doc_builder import DocBuilder
import os
import time

# 模板文件路径
template_path = r'e:\Qwen\xmjl\03-临平区数字慈善系统建设方案模板.docx'
output_path = r'e:\Qwen\xmjl\outputs\test_with_logs.docx'

print('=' * 80)
print('开始测试完整的模板分析和文档生成流程')
print('=' * 80)

start_time = time.time()

# ========== 步骤 1: 分析模板 ==========
print(f'\n[步骤 1/3] 分析模板文档: {template_path}')
print('-' * 80)

analyzer = TemplateAnalyzer(template_path)
template_structure = {
    'chapter_tree': analyzer.get_chapter_tree()
}

print(f'\n[步骤 1 完成] 共找到 {len(template_structure["chapter_tree"])} 个一级章节')

# ========== 步骤 2: 构建章节内容字典 ==========
print(f'\n[步骤 2/3] 构建章节内容字典')
print('-' * 80)

# 构建章节内容字典（从模板结构中提取）
chapter_contents = {}
for chapter in template_structure['chapter_tree']:
    chapter_title = chapter.get('title', '')
    chapter_contents[chapter_title] = ''
    
    # 处理小节
    for section in chapter.get('children', []):
        section_title = section.get('title', '')
        chapter_contents[section_title] = ''
        
        # 处理小小节
        for child in section.get('children', []):
            child_title = child.get('title', '')
            chapter_contents[child_title] = ''

print(f'[步骤 2 完成] 构建了 {len(chapter_contents)} 个内容块（空内容，仅测试结构）')

# ========== 步骤 3: 生成文档 ==========
print(f'\n[步骤 3/3] 生成新文档')
print('-' * 80)

builder = DocBuilder(template_path=template_path)
builder.render_from_template_structure(template_structure, chapter_contents)

os.makedirs(os.path.dirname(output_path), exist_ok=True)
builder.save(output_path)

end_time = time.time()

print(f'\n[步骤 3 完成] 文档已保存到：{output_path}')
print(f'生成耗时：{end_time - start_time:.2f} 秒')

# ========== 验证生成的文档 ==========
print('\n' + '=' * 80)
print('验证生成的文档')
print('=' * 80)

from docx import Document

orig_doc = Document(template_path)
gen_doc = Document(output_path)

orig_styles = {}
for para in orig_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    orig_styles[style_name] = orig_styles.get(style_name, 0) + 1

gen_styles = {}
for para in gen_doc.paragraphs:
    text = para.text.strip()
    if not text:
        continue
    style_name = para.style.name if para.style else 'Normal'
    gen_styles[style_name] = gen_styles.get(style_name, 0) + 1

print(f'\n原始模板段落数：{len(orig_doc.paragraphs)}')
print(f'生成文档段落数：{len(gen_doc.paragraphs)}')

print(f'\n原始模板样式数：{len(orig_styles)}')
print(f'生成文档样式数：{len(gen_styles)}')

print(f'\n样式对比:')
all_styles = set(orig_styles.keys()) | set(gen_styles.keys())
for style in sorted(all_styles):
    orig = orig_styles.get(style, 0)
    gen = gen_styles.get(style, 0)
    match = '✓' if orig == gen else '✗'
    print(f'  {match} {style}: 原始={orig}, 生成={gen}')

print('\n' + '=' * 80)
print('测试完成')
print('=' * 80)
